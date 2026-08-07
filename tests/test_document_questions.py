import json
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from azure_search import (
    CHANGE_MANIFEST_NAME,
    _credential,
    _changed_document_ids,
    _clear_change_manifest,
    _clear_deletion_manifest,
    _deletion_document_ids,
    _document_records,
    _add_runtime_sharepoint_parent_context,
    _searchable_filename_terms,
    _document_pages,
    _chunks,
    _document_relevance_score,
    _diversify_candidate_records,
    _has_minimum_content_coverage,
    _is_script_record,
    _requests_script,
    _filter_records_for_requested_country,
    _focused_keyword_query,
    _question_without_background_action,
    _record_has_authorized_provenance,
    _record_matches_requested_version,
    _records_supported_by_index,
    _entra_credential,
    _excerpt_around_query,
    _rerank_records,
    retrieve_azure_search_evidence,
    index_directory,
)
from document_index import has_requested_action_coverage
from classification import (
    _focused_procedure_evidence,
    _grounded_document_summary,
    classify_case_by_rules,
    is_direct_document_question,
    is_underspecified_query,
    needs_extension_subject_context,
)
from document_index import tokenize
from formatting import format_user_response
from models import EvidenceSource
from models import BotDecision
from sharepoint_sync import (
    CHANGE_MANIFEST_NAME as SYNC_CHANGE_MANIFEST_NAME,
    DELETION_MANIFEST_NAME,
    _change_ids,
    sync_pdfs,
)


class DocumentQuestionTests(unittest.TestCase):
    def test_generic_queries_are_marked_for_context_before_retrieval(self):
        self.assertTrue(is_underspecified_query("¿Qué se debe revisar?"))
        self.assertTrue(is_underspecified_query("No funciona."))
        self.assertFalse(
            is_underspecified_query(
                "Después de reinstalar MSDTC, ¿qué se debe revisar en ambos servidores?"
            )
        )

    def test_filename_procedure_is_answered_when_action_and_topic_are_explicit(self):
        decision = classify_case_by_rules(
            "¿Cómo puedo arreglar vacaciones negativas con un script?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="acc.proc_arreglar_vac_negativos.sql — Documento",
                    ubicacion="https://contoso.example/vacaciones.sql",
                    fragmento=(
                        "CREATE PROCEDURE acc.proc_arreglar_vac_negativos. "
                        "Corrige vacaciones con saldo negativo."
                    ),
                )
            ],
        )

        self.assertEqual("resuelto", decision.estado)
        self.assertEqual(1, len(decision.fuentes))

    def test_incapacity_manual_is_answered_from_direct_classification_evidence(self):
        decision = classify_case_by_rules(
            "¿Cómo se clasifican las incapacidades en Evolution?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Acciones de personal.pdf — Página 38",
                    ubicacion="https://contoso.example/acciones.pdf",
                    fragmento=(
                        "Tipos de incapacidad y clasificaciones establecidas por la "
                        "compañía. Según su duración, pueden ser permanentes o temporales."
                    ),
                )
            ],
        )

        self.assertEqual("resuelto", decision.estado)

    def test_ambiguous_extension_requires_subject_context(self):
        self.assertTrue(needs_extension_subject_context("¿Qué parámetros se pueden configurar para prórroga en Evolution?"))
        self.assertFalse(needs_extension_subject_context("¿Qué parámetros se pueden configurar para prórroga de contratos en Evolution?"))

    def test_parameter_evidence_accepts_documented_extension_and_incapacity_fields(self):
        extension = [EvidenceSource(
            tipo="sharepoint", titulo="Acciones de personal.pdf — Página 18",
            ubicacion="https://contoso.example/acciones.pdf",
            fragmento="Parámetro: ProrrogaContratoDiasAtrasInicioRangoFechaFinContrato. Define los días antes de la fecha final.",
        )]
        incapacity = [EvidenceSource(
            tipo="sharepoint", titulo="Acciones de personal.pdf — Página 34",
            ubicacion="https://contoso.example/acciones.pdf",
            fragmento="El parámetro de aplicación: IncapacidadesValidaTraslapeConAcciones valida traslapes.",
        )]
        self.assertEqual("resuelto", classify_case_by_rules("¿Qué parámetros se pueden configurar para prórroga de contratos?", extension).estado)
        self.assertEqual("resuelto", classify_case_by_rules("¿Qué parámetros se pueden configurar para incapacidades?", incapacity).estado)

    def test_parameter_summary_lists_all_documented_extension_parameters(self):
        evidence = [EvidenceSource(
            tipo="sharepoint", titulo="Acciones de personal.pdf — Página 18",
            ubicacion="https://contoso.example/acciones.pdf",
            fragmento=(
                "Parámetro: ProrrogaContratoDiasAtrasInicioRangoFechaFinContrato "
                "permite especificar días antes de la fecha final. Parámetro: "
                "ProrrogaContratoDiasDespuesFinalRangoFechaFinContrato permite "
                "especificar días después de la fecha final."
            ),
        )]
        answer = _grounded_document_summary(
            "¿Qué parámetros se pueden configurar para prórroga de contratos?", evidence
        )
        self.assertIn("ProrrogaContratoDiasAtrasInicioRangoFechaFinContrato", answer)
        self.assertIn("ProrrogaContratoDiasDespuesFinalRangoFechaFinContrato", answer)

    def test_conceptual_classification_and_examples_do_not_become_navigation_steps(self):
        classification = [
            EvidenceSource(
                tipo="sharepoint", titulo="Acciones de personal.pdf — Página 38",
                ubicacion="https://contoso.example/acciones.pdf",
                fragmento="Clasificación de incapacidades. Según su duración se consideran permanentes y temporales. Según su magnitud hay incapacidades parciales y totales. Según su cualidad se separan en físicas y psíquicas.",
            ),
            EvidenceSource(
                tipo="sharepoint", titulo="Manual DB.docx",
                ubicacion="https://contoso.example/db.docx",
                fragmento="Haga clic en el módulo de acciones de personal. Seleccione opciones. Haga clic en guardar.",
            ),
        ]
        examples = [EvidenceSource(
            tipo="sharepoint", titulo="Gestion de documentos.pdf — Página 4",
            ubicacion="https://contoso.example/documentos.pdf",
            fragmento="Ejemplo de los tipos de documento que puede administrar: Formularios, Manuales, Procedimientos, Instructivos. Haga clic en el botón Nuevo. Seleccione Guardar.",
        )]
        classification_answer = _grounded_document_summary("¿Cómo se clasifican las incapacidades en Evolution?", classification)
        examples_answer = _grounded_document_summary("Dame ejemplos de tipos de documentos que se pueden administrar en Evolution", examples)
        self.assertIn("permanentes", classification_answer)
        self.assertNotIn("Haga clic", classification_answer)
        self.assertIn("Formularios", examples_answer)
        self.assertNotIn("Haga clic", examples_answer)

    def test_download_failure_does_not_turn_navigation_steps_into_a_diagnosis(self):
        evidence = [
            EvidenceSource(
                tipo="azure_ai_search",
                titulo="Gestion de documentos.pdf — Página 12",
                ubicacion="https://contoso.example/gestion.pdf",
                fragmento=(
                    "Capítulo: GESTIÓN DE DOCUMENTOS 11 A Administrar documentos gestionados, "
                    "5 D Descargue los documentos sobre los que se tiene permisos, 9"
                ),
            ),
            EvidenceSource(
                tipo="azure_ai_search",
                titulo="Gestion de documentos.pdf — Página 10",
                ubicacion="https://contoso.example/gestion.pdf",
                fragmento=(
                    "Descargue los documentos sobre los que se tiene permisos. "
                    "Haga clic en el área Portal. Seleccione el modulo Consultas. "
                    "Haga clic en la opción Documentos Gestionados. "
                    "Haga clic en el Titulo del documento que desea descargar. "
                    "Haga clic en la opción Descargar."
                ),
            ),
            EvidenceSource(
                tipo="azure_ai_search",
                titulo="Portal Consultas.pdf — Página 30",
                ubicacion="https://contoso.example/portal.pdf",
                fragmento=(
                    "Documentos Gestionados. Haga clic en el área Portal. "
                    "Seleccione el Modulo Consultas. Haga clic en la opción "
                    "Documentos Gestionados. Haga clic en el documento que desea descargar. "
                    "Haga clic en la opción Descargar. Seleccione la opción "
                    "Documentos Gestionados."
                ),
            ),
        ]

        decision = classify_case_by_rules(
            "Un usuario tiene permisos pero no logra bajar los documentos del módulo de gestión, ¿qué se debe revisar?",
            evidence,
        )

        self.assertEqual("sin_evidencia", decision.estado)
        self.assertEqual([], decision.fuentes)

    def test_answer_quality_matrix_resolves_procedure_and_diagnostic_but_abstains_unrelated(self):
        cases = [
            (
                "¿Cómo se pueden administrar documentos gestionados?",
                "Administrar documentos gestionados. Haga clic en Portal. "
                "Seleccione Consultas. Haga clic en Descargar.",
                "resuelto",
            ),
            (
                "Un usuario tiene permisos pero no logra bajar documentos, ¿qué se debe revisar?",
                "Gestión de documentos. Descargue los documentos sobre los que se tiene permisos. "
                "Si falla, revisar permisos y acceso al módulo.",
                "resuelto",
            ),
            (
                "¿Cuál es la fecha de vencimiento del certificado SSL?",
                "Guía de vacaciones y aprobación de solicitudes.",
                "sin_evidencia",
            ),
        ]

        for question, fragment, expected_state in cases:
            with self.subTest(question=question):
                decision = classify_case_by_rules(
                    question,
                    [
                        EvidenceSource(
                            tipo="sharepoint",
                            titulo="Manual de gestión",
                            ubicacion="https://contoso.example/manual.pdf",
                            fragmento=fragment,
                        )
                    ],
                )
                self.assertEqual(expected_state, decision.estado)

    def test_retrieval_uses_fields_supported_by_legacy_search_indexes(self):
        class FakeSearchClient:
            def __init__(self):
                self.calls = []

            def search(self, **kwargs):
                self.calls.append(kwargs)
                return [
                    {
                        "id": "sv-page-1",
                        "title": "Políticas de Pago SV.pdf — Página 1",
                        "source_url": "https://contoso.example/politicas-sv.pdf",
                        "source_system": "sharepoint",
                        "document_context": "El Salvador. Planilla Mensual, Planilla Quincenal y Planilla Aguinaldo.",
                        "content": "El Salvador. Dentro del proyecto se contemplan la Planilla Mensual, la Planilla Quincenal y la Planilla Aguinaldo.",
                        "content_tokens": "salvador planilla mensual quincenal aguinaldo",
                    }
                ]

        fake_search = FakeSearchClient()
        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="chat-salvador-docs",
            azure_search_api_key="not-a-real-key",
            azure_search_use_entra_id=False,
            openai_api_key="test-key",
            openai_base_url="",
            resolved_openai_base_url="https://api.openai.com/v1",
            openai_embedding_model="text-embedding-3-small",
        )

        with patch("azure_search.SearchClient", return_value=fake_search), patch(
            "azure_search._embed_texts", return_value=[[0.0, 0.0]]
        ):
            sources = retrieve_azure_search_evidence(
                "¿Cuáles son las planillas que se pagan en El Salvador?",
                config,
            )

        self.assertEqual("Políticas de Pago SV.pdf — Página 1", sources[0].titulo)
        for call in fake_search.calls:
            self.assertNotIn("document_id", call["select"])
            self.assertNotIn("document_version", call["select"])
            self.assertIn("folder_path", call["select"])

    def test_explicit_filename_prioritizes_the_exact_indexed_document(self):
        requested_name = "sp_anular_solicitud_vac.sql"

        class FakeSearchClient:
            def search(self, **kwargs):
                if kwargs.get("search_text") == requested_name:
                    return [
                        {
                            "id": "scripts-anular-vac",
                            "title": f"{requested_name} — Documento",
                            "source_url": "https://contoso.example/scripts/sp_anular_solicitud_vac.sql",
                            "source_system": "sharepoint",
                            "folder_path": "",
                            "drive_id": "drive-scripts",
                            "content": "EXEC sp_anular_solicitud_vac @id_solicitud = 123;",
                            "content_tokens": "exec sp anular solicitud vac id solicitud",
                        }
                    ]
                return [
                    {
                        "id": "related-config",
                        "title": "ConfiguracionAnulacionSolicitudVac.sql — Documento",
                        "source_url": "https://contoso.example/scripts/ConfiguracionAnulacionSolicitudVac.sql",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-scripts",
                        "content": "Configuración relacionada con solicitudes de vacaciones.",
                        "content_tokens": "configuracion anulacion solicitud vacaciones",
                    }
                ]

        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
            azure_search_use_entra_id=False,
            sharepoint_sources=(("", "drive-scripts"),),
        )

        with patch("azure_search.SearchClient", return_value=FakeSearchClient()), patch(
            "azure_search._embed_texts", return_value=[[0.0, 0.0]]
        ):
            sources = retrieve_azure_search_evidence(
                "Busca exactamente el archivo sp_anular_solicitud_vac.sql. "
                "¿Qué parámetros utiliza y qué operación realiza?",
                config,
            )

        self.assertEqual([f"{requested_name} — Documento"], [source.titulo for source in sources])

    def test_unknown_explicit_filename_does_not_fall_back_to_related_documents(self):
        class FakeSearchClient:
            def search(self, **kwargs):
                return [
                    {
                        "id": "related-readme",
                        "title": "Readme 1.19.1.13.pdf — Página 7",
                        "source_url": "https://contoso.example/readme.pdf",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-readme",
                        "content": "Validación previa a la instalación de una actualización.",
                        "content_tokens": "validacion previa instalacion actualizacion",
                    }
                ]

        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
            azure_search_use_entra_id=False,
            sharepoint_sources=(("", "drive-readme"),),
        )

        with patch("azure_search.SearchClient", return_value=FakeSearchClient()), patch(
            "azure_search._embed_texts", return_value=[[0.0, 0.0]]
        ):
            sources = retrieve_azure_search_evidence(
                "Busca exactamente el archivo procedimiento_inexistente.sql.", config
            )

        self.assertEqual([], sources)

    def test_exact_version_excludes_readmes_with_only_a_shared_prefix(self):
        class FakeSearchClient:
            def search(self, **_kwargs):
                return [
                    {
                        "id": "readme-1-19-1-0",
                        "title": "Readme 1.19.1.0.pdf — Página 1",
                        "source_url": "https://contoso.example/readme-1.19.1.0.pdf",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-readme",
                        "content": "Evolution 1.19.1.0 incluye cambios de instalación.",
                    },
                    {
                        "id": "readme-1-19-1-10",
                        "title": "Readme 1.19.1.10.pdf — Página 1",
                        "source_url": "https://contoso.example/readme-1.19.1.10.pdf",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-readme",
                        "content": "Evolution 1.19.1.10 incorpora la actualización documentada.",
                    },
                ]

        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
            azure_search_use_entra_id=False,
            sharepoint_sources=(("", "drive-readme"),),
        )
        with patch("azure_search.SearchClient", return_value=FakeSearchClient()), patch(
            "azure_search._embed_texts", return_value=[[0.0, 0.0]]
        ):
            sources = retrieve_azure_search_evidence(
                "Dame detalles sobre la versión de Evolution 1.19.1.10.", config
            )

        self.assertEqual(["Readme 1.19.1.10.pdf — Página 1"], [source.titulo for source in sources])
        self.assertFalse(
            _record_matches_requested_version(
                {"title": "Readme 1.19.1.0.pdf"}, ("1.19.1.10",)
            )
        )

    def test_versioned_readme_section_excludes_contents_and_update_documents(self):
        class FakeSearchClient:
            def search(self, **_kwargs):
                return [
                    {
                        "id": "update-page",
                        "title": "Actualización 1.19.1.11.pdf — Página 1",
                        "source_url": "https://contoso.example/update.pdf",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-readme",
                        "content": "Mejoras de Evolution 1.19.1.11.",
                    },
                    {
                        "id": "readme-contents",
                        "title": "Readme 1.19.1.11.pdf — Página 2",
                        "source_url": "https://contoso.example/readme.pdf",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-readme",
                        "content": "Evolution 1.19.1.11. Tabla de contenido y recomendaciones generales.",
                    },
                    {
                        "id": "readme-requirements",
                        "title": "Readme 1.19.1.11.pdf — Página 8",
                        "source_url": "https://contoso.example/readme.pdf",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-readme",
                        "content": "Nuevos requisitos de software: Ninguno.",
                    },
                ]

        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
            azure_search_use_entra_id=False,
            sharepoint_sources=(("", "drive-readme"),),
        )
        with patch("azure_search.SearchClient", return_value=FakeSearchClient()), patch(
            "azure_search._embed_texts", return_value=[[0.0, 0.0]]
        ):
            sources = retrieve_azure_search_evidence(
                "¿Qué nuevos requisitos de software necesita Evolution versión Readme 1.19.1.11?", config
            )

        self.assertEqual(["Readme 1.19.1.11.pdf — Página 8"], [source.titulo for source in sources])

    def test_entra_credential_is_reused_for_multiple_search_operations(self):
        config = SimpleNamespace(
            azure_search_api_key="",
            azure_search_use_entra_id=True,
        )
        _entra_credential.cache_clear()
        try:
            with patch("azure_search.DefaultAzureCredential") as credential_class:
                first = _credential(config)
                second = _credential(config)

            self.assertIs(first, second)
            credential_class.assert_called_once_with(
                exclude_interactive_browser_credential=False
            )
        finally:
            _entra_credential.cache_clear()

    def test_evidence_source_preserves_optional_document_metadata(self):
        source = EvidenceSource(
            tipo="sharepoint",
            titulo="Manual",
            ubicacion="https://contoso.example/manual.pdf",
            fragmento="Texto de prueba.",
            document_id="drive-item-123",
            document_version='"etag-2"',
            last_modified="2026-07-22T12:00:00Z",
            document_type="pdf",
            folder_path="Operaciones/Manuales",
        )

        self.assertEqual("drive-item-123", source.document_id)
        self.assertEqual("pdf", source.document_type)
        self.assertEqual("Operaciones/Manuales", source.folder_path)

    def test_deletion_manifest_is_read_and_cleared_after_indexing(self):
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            manifest = source_dir / DELETION_MANIFEST_NAME
            manifest.write_text(
                '{"deleted_document_ids": ["one", "two", "one"]}', encoding="utf-8"
            )

            self.assertEqual({"one", "two"}, _deletion_document_ids(source_dir))
            _clear_deletion_manifest(source_dir)
            self.assertEqual(set(), _deletion_document_ids(source_dir))

    def test_change_manifest_is_read_and_cleared_after_indexing(self):
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            manifest = source_dir / CHANGE_MANIFEST_NAME
            manifest.write_text(
                '{"changed_document_ids": ["one", "two", "one"]}', encoding="utf-8"
            )

            self.assertEqual({"one", "two"}, _changed_document_ids(source_dir))
            _clear_change_manifest(source_dir)
            self.assertEqual(set(), _changed_document_ids(source_dir))

    def test_pdf_records_include_stable_sharepoint_metadata(self):
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            pdf_path = source_dir / "manual.pdf"
            pdf_path.write_bytes(b"placeholder")
            pdf_path.with_suffix(".pdf.metadata.json").write_text(
                """{
                  "source_system": "sharepoint",
                  "document_id": "drive-item-123",
                  "etag": "etag-2",
                  "last_modified": "2026-07-22T12:00:00Z",
                  "web_url": "https://contoso.example/manual.pdf",
                  "folder_path": "Operaciones/Manuales",
                  "drive_id": "drive-1"
                }""",
                encoding="utf-8",
            )
            with patch(
                "azure_search._document_pages",
                return_value=[(1, "Procedimiento para instalar el hotfix de nómina.")],
            ):
                records = _document_records(source_dir)

        self.assertEqual(1, len(records))
        record = records[0]
        self.assertEqual("drive-item-123", record["document_id"])
        self.assertEqual("etag-2", record["document_version"])
        self.assertEqual("pdf", record["document_type"])
        self.assertEqual("Operaciones/Manuales", record["folder_path"])
        self.assertEqual("drive-1", record["drive_id"])
        self.assertIn("manual", record["content_tokens"])
        self.assertIn("Título del archivo", record["document_context"])
        self.assertTrue(record["content_hash"])
        self.assertTrue(record["indexed_at"])

    def test_parent_folder_is_searchable_when_file_name_is_generic(self):
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            document = source_dir / "Indicaciones.txt"
            document.write_text("Reemplace la sección indicada.", encoding="utf-8")
            document.with_suffix(".txt.metadata.json").write_text(
                json.dumps(
                    {
                        "source_system": "sharepoint",
                        "document_id": "item-indicaciones",
                        "web_url": (
                            "https://contoso.sharepoint.com/sites/Soporte/Documentos%20compartidos/"
                            "SOLUCIONES/EVOLUTION/ERROR%20DE%20FECHAS%20EN%20TIEMPOS%20NO%20TRABAJADOS/"
                            "Indicaciones.txt"
                        ),
                        "folder_path": "SOLUCIONES",
                        "drive_id": "drive-soluciones",
                    }
                ),
                encoding="utf-8",
            )
            records = _document_records(source_dir)

        self.assertEqual(1, len(records))
        self.assertIn("error", records[0]["content_tokens"])
        self.assertIn("trabajado", records[0]["content_tokens"])
        self.assertIn("ERROR DE FECHAS", records[0]["document_context"])

    def test_legacy_record_derives_parent_folder_from_sharepoint_url_at_retrieval(self):
        record = _add_runtime_sharepoint_parent_context(
            {
                "title": "Indicaciones.txt — Documento",
                "document_context": "Título del archivo: Indicaciones.txt.",
                "folder_path": "SOLUCIONES",
                "source_url": (
                    "https://contoso.sharepoint.com/sites/Soporte/Documentos%20compartidos/"
                    "SOLUCIONES/EVOLUTION/ERROR%20DE%20FECHAS%20EN%20TIEMPOS%20NO%20TRABAJADOS/"
                    "Indicaciones.txt"
                ),
            }
        )

        self.assertIn("ERROR DE FECHAS EN TIEMPOS NO TRABAJADOS", record["document_context"])

    def test_legacy_index_update_omits_additive_v2_fields(self):
        filtered = _records_supported_by_index(
            [{"id": "fragment-1", "content": "texto", "retrieval_text": "metadato"}],
            {"id", "content"},
        )

        self.assertEqual([{"id": "fragment-1", "content": "texto"}], filtered)

    def test_filename_conventions_are_exposed_as_searchable_terms(self):
        terms = _searchable_filename_terms("acc.proc_arreglar_vac_negativos.sql")

        self.assertEqual("acc proc arreglar vac negativos sql", terms)

    def test_sharepoint_description_is_indexed_and_searchable(self):
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            document = source_dir / "acc.proc_arreglar_vac_negativos.sql"
            document.write_text("CREATE PROCEDURE acc.proc_arreglar_vac_negativos AS SELECT 1", encoding="utf-8")
            document.with_suffix(".sql.metadata.json").write_text(
                json.dumps(
                    {
                        "source_system": "sharepoint",
                        "document_id": "script-vac-negativos",
                        "name": document.name,
                        "web_url": "https://contoso.example/Scripts%20de%20Apoyo/acc.proc_arreglar_vac_negativos.sql",
                        "folder_path": "Scripts de Apoyo",
                        "description": "Corrige los periodos de vacaciones que tienen saldo negativo y mueve el tiempo a nuevos periodos.",
                        "dependency": "Base de datos respaldada.",
                    }
                ),
                encoding="utf-8",
            )
            with patch("azure_search._document_pages", return_value=[(0, document.read_text(encoding="utf-8"))]):
                records = _document_records(source_dir)

        self.assertEqual(1, len(records))
        self.assertIn("saldo negativo", records[0]["document_context"])
        self.assertIn("nuevo", records[0]["content_tokens"])
        self.assertIn("periodo", records[0]["content_tokens"])

    def test_result_fragment_exposes_sharepoint_description(self):
        from azure_search import _result_fragment

        fragment = _result_fragment(
            {
                "document_context": "Descripción de la solución: Corrige vacaciones con saldo negativo.",
                "content": "CREATE PROCEDURE acc.proc_arreglar_vac_negativos",
            },
            "script para vacaciones negativas",
        )

        self.assertIn("Descripción de la solución", fragment)
        self.assertIn("saldo negativo", fragment)

    def test_sql_identifiers_keep_concepts_searchable(self):
        tokens = set(tokenize("vac_vacaciones acc.proc_arreglar_vac_negativos.sql"))

        self.assertIn("vacacion", tokens)
        self.assertIn("arreglar", tokens)
        self.assertIn("negativo", tokens)

    def test_document_context_participates_in_relevance_score(self):
        question = "¿Cómo puedo arreglar vacaciones negativas con un script?"
        record = {
            "title": "acc.proc_arreglar_vac_negativos.sql — Documento",
            "document_context": "Título del archivo: acc.proc_arreglar_vac_negativos.sql. "
            "CREATE PROCEDURE para corregir vacaciones con saldo negativo.",
            "content_tokens": "acc proc arreglar vac vacacion negativo sql",
            "content": "CREATE PROCEDURE acc.proc_arreglar_vac_negativos",
            "document_type": "sql",
        }

        self.assertGreaterEqual(_document_relevance_score(record, question), 12)

    def test_title_concept_outweighs_incidental_fragment_match(self):
        question = "¿Cómo puedo arreglar vacaciones negativas con un script?"
        named_procedure = {
            "title": "acc.proc_arreglar_vac_negativos.sql — Documento",
            "document_context": "Procedimiento para corregir vacaciones con saldo negativo.",
            "content_tokens": "arreglar vacacion negativo sql",
            "content": "CREATE PROCEDURE acc.proc_arreglar_vac_negativos",
            "document_type": "sql",
        }
        incidental_sql = {
            "title": "ConfiguracionAnulacionSolicitudVac.sql — Documento",
            "document_context": "Configuración de solicitudes de vacaciones.",
            "content_tokens": "solicitud vacacion sql",
            "content": "EXEC sp_anular_solicitud_vac",
            "document_type": "sql",
        }

        self.assertGreater(
            _document_relevance_score(named_procedure, question),
            _document_relevance_score(incidental_sql, question),
        )

    def test_script_intent_prioritizes_executable_artifact_over_readme(self):
        question = (
            "¿Hay algún script para corregir los periodos de vacaciones con saldo "
            "negativo y mover el tiempo a nuevos periodos?"
        )
        script = {
            "title": "acc.proc_arreglar_vac_negativos.sql — Documento",
            "document_context": (
                "Descripción de la solución: Corrige periodos de vacaciones con "
                "saldo negativo y mueve el tiempo a nuevos periodos."
            ),
            "content_tokens": "acc proc arreglar vacacion negativo mover periodo",
            "content": "CREATE PROCEDURE acc.proc_arreglar_vac_negativos",
            "document_type": "sql",
        }
        readme = {
            "title": "ReadME Hotfixes.pdf — Página 1",
            "document_context": "Incidencia de vacaciones en una actualización.",
            "content_tokens": "vacacion actualizacion",
            "content": "Se corrigió una incidencia de vacaciones.",
            "document_type": "pdf",
        }

        self.assertTrue(_requests_script(question))
        self.assertTrue(_is_script_record(script))
        self.assertFalse(_is_script_record(readme))
        self.assertGreater(
            _document_relevance_score(script, question),
            _document_relevance_score(readme, question),
        )

    def test_content_coverage_accepts_morphology_but_rejects_one_term_country_noise(self):
        from azure_search import _has_minimum_content_coverage

        self.assertTrue(
            _has_minimum_content_coverage(
                {
                    "title": "acc.proc_arreglar_vac_negativos.sql — Documento",
                    "document_context": "Procedimiento para corregir vacaciones con saldo negativo.",
                    "content_tokens": "arreglar vacacion negativo sql",
                    "content": "CREATE PROCEDURE acc.proc_arreglar_vac_negativos",
                    "document_type": "sql",
                },
                "¿Cómo puedo arreglar vacaciones negativas con un script?",
            )
        )
        self.assertFalse(
            _has_minimum_content_coverage(
                {
                    "title": "Readme.pdf — Página 1",
                    "document_context": "El Salvador. Incidencia sobre descuentos cíclicos.",
                    "content_tokens": "salvador descuento ciclico",
                    "content": "Incidencia técnica sin legislación vigente.",
                },
                "¿Qué descuentos legales existen en El Salvador?",
            )
        )

    def test_content_coverage_rejects_navigation_page_for_operational_question(self):
        self.assertFalse(
            _has_minimum_content_coverage(
                {
                    "title": "Gestion de documentos.pdf — Página 2",
                    "content_tokens": "gestion documentos administrar",
                    "content": (
                        "Página 2 Tabla de contenido Gestión de documentos 2 "
                        "Tipos de documentos gestionados 3 Áreas de documentos "
                        "gestionados 4 Administrar documentos gestionados 5 "
                        "Agregue una nueva versión 7 Descargue documentos 9"
                    ),
                    "document_type": "pdf",
                },
                "¿Cómo se pueden administrar documentos en Evolution?",
            )
        )

    def test_content_coverage_rejects_accented_navigation_page_for_incapacity_parameters(self):
        self.assertFalse(
            _has_minimum_content_coverage(
                {
                    "title": "Acciones de personal.pdf — Página 103",
                    "content": (
                        "Página 103 Acreditación, 67 Administración, 38 "
                        "Amonestaciones, 75 Ausencias, 68 Cambio de centro, 84 "
                        "Cambio de jornada, 95 Cambio de planilla, 97"
                    ),
                },
                "¿Qué parámetros se pueden modificar para incapacidades?",
            )
        )

    def test_version_change_question_rejects_table_of_contents_as_evidence(self):
        self.assertFalse(
            _has_minimum_content_coverage(
                {
                    "title": "Upgrade 1.24.1.1.pdf — Página 2",
                    "content_tokens": "evolution 1.24.1.1 cambio parametro aplicacion",
                    "content": (
                        "Página 2 Evolution 1.24.1.1 Tabla de contenido "
                        "Instrucciones de instalación 4, actualización de base 6, "
                        "cambios en parámetros de infraestructura 8, "
                        "mejoras de aplicación 10, correcciones 12, anexos 14"
                    ),
                    "document_type": "pdf",
                },
                "¿Qué cambios incluye Upgrade 1.24.1.1 y qué precauciones debo tomar al aplicarlo?",
            )
        )

    def test_content_coverage_does_not_mistake_sql_commas_for_navigation(self):
        self.assertTrue(
            _has_minimum_content_coverage(
                {
                    "title": "acc.proc_arreglar_vac_negativos.sql — Documento",
                    "content_tokens": "arreglar vacacion negativo sql",
                    "content": (
                        "update sal.dss_descuentos set dss_valor = 0, "
                        "dss_tipo = 1, dss_estado = 1, dss_aplica = 1, "
                        "dss_codigo = 1, dss_periodo = 1"
                    ),
                    "document_type": "sql",
                },
                "¿Cómo puedo arreglar vacaciones negativas con un script?",
            )
        )

    def test_action_coverage_recognizes_conjugated_action(self):
        self.assertTrue(
            has_requested_action_coverage(
                "¿Cómo se clasifican las incapacidades?",
                "Tipos de incapacidad y clasificaciones establecidas por la compañía.",
            )
        )

    def test_action_coverage_maps_bajar_to_manual_descargue(self):
        self.assertTrue(
            has_requested_action_coverage(
                "Un usuario tiene permisos pero no logra bajar los documentos del módulo de gestión, ¿qué se debe revisar?",
                "Gestión de documentos. Descargue los documentos sobre los que se tiene permisos.",
            )
        )

    def test_focused_keyword_query_removes_question_scaffolding(self):
        self.assertEqual(
            _focused_keyword_query(
                "¿Qué parámetros se pueden configurar para prórroga de contratos en Evolution?"
            ),
            (
                "parametro configurar prorroga contrato evolution "
                "prorrogacontratodiasatrasiniciorangofechafincontrato "
                "prorrogacontratodiasdespuesfinalrangofechafincontrato"
            ),
        )
        self.assertEqual(
            _focused_keyword_query(
                "Dame los parámetros que se relacionan con la prórroga de contratos"
            ),
            (
                "parametro prorroga contrato "
                "prorrogacontratodiasatrasiniciorangofechafincontrato "
                "prorrogacontratodiasdespuesfinalrangofechafincontrato"
            ),
        )
        self.assertIn(
            "riesgo",
            _focused_keyword_query("¿Qué parámetros se pueden modificar para incapacidades?"),
        )
        self.assertEqual(
            "precaucion deben tomar instalar actualizacion evolution preparacion respaldo configuracion",
            _focused_keyword_query(
                "¿Qué precauciones se deben tomar antes de instalar una actualización de Evolution?"
            ),
        )

    def test_preinstallation_precautions_retrieve_preparation_evidence(self):
        class FakeSearchClient:
            def __init__(self):
                self.queries = []

            def search(self, **kwargs):
                query = kwargs.get("search_text") or ""
                self.queries.append(query)
                if "preparacion" not in query:
                    return []
                return [
                    {
                        "id": "upgrade-preparation",
                        "title": "Upgrade Evolution 1.24.1.1.pdf — Página 3",
                        "source_url": "https://contoso.example/upgrade.pdf",
                        "source_system": "sharepoint",
                        "folder_path": "",
                        "drive_id": "drive-upgrades",
                        "content": (
                            "Preparación. Antes de proceder con la instalación o actualización, "
                            "realice respaldos completos de archivos de configuración y de la base de datos Evolution."
                        ),
                        "content_tokens": "preparacion instalacion actualizacion respaldo configuracion base dato evolution",
                    }
                ]

        fake_search = FakeSearchClient()
        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
            azure_search_use_entra_id=False,
            sharepoint_sources=(("", "drive-upgrades"),),
        )

        with patch("azure_search.SearchClient", return_value=fake_search), patch(
            "azure_search._embed_texts", return_value=[[0.0, 0.0]]
        ):
            sources = retrieve_azure_search_evidence(
                "¿Qué precauciones se deben tomar antes de instalar una actualización de Evolution?",
                config,
            )

        self.assertEqual(["Upgrade Evolution 1.24.1.1.pdf — Página 3"], [source.titulo for source in sources])
        self.assertTrue(any("preparacion" in query for query in fake_search.queries))

    def test_preinstallation_ranking_prefers_a_preparation_checklist(self):
        question = "¿Qué precauciones se deben tomar antes de instalar una actualización de Evolution?"
        tangential = {
            "title": "Readme de vulnerabilidades.pdf — Página 5",
            "content": "La actualización instala dependencias para mitigar vulnerabilidades.",
            "content_tokens": "actualizacion instalacion vulnerabilidad",
            "document_type": "pdf",
        }
        checklist = {
            "title": "Upgrade Evolution.pdf — Página 3",
            "content": (
                "Preparación previa: realice un respaldo y revise las recomendaciones "
                "iniciales antes de instalar la actualización."
            ),
            "content_tokens": "preparacion previa respaldo recomendacion inicial instalacion actualizacion",
            "document_type": "pdf",
        }

        self.assertGreater(
            _document_relevance_score(checklist, question),
            _document_relevance_score(tangential, question),
        )

    def test_preinstallation_query_expands_update_and_backup_variants(self):
        query = _focused_keyword_query(
            "Antes de actualizar Evolution, ¿qué respaldos y precauciones recomienda la guía de instalación?"
        )

        self.assertIn("preparacion", query)
        self.assertIn("respaldo", query)

    def test_dtc_validation_query_expands_to_actionable_manual_pages(self):
        query = _focused_keyword_query(
            "Después de reinstalar MSDTC, ¿qué debo validar en ambos servidores?"
        )

        self.assertIn("firewall", query)
        self.assertIn("component", query)
        self.assertIn("inboud", query)

    def test_dtc_validation_ranking_prefers_actionable_check_over_heading(self):
        question = "Después de reinstalar MSDTC, ¿qué debo validar en ambos servidores?"
        heading = {
            "title": "Manual DTC Verificacion.pdf — Página 3",
            "content": "Validación. Validar en ambos servidores que estos servicios están corriendo Base de datos:",
            "content_tokens": "validacion validar ambos servidor servicio base dato dtc",
            "document_type": "pdf",
        }
        actionable_check = {
            "title": "Manual DTC Verificacion.pdf — Página 5",
            "content": (
                "Validar en Component Services LOCAL DTC en ambos servidores y confirmar "
                "que la configuración sea igual."
            ),
            "content_tokens": "validar component services local dtc ambos servidor configuracion",
            "document_type": "pdf",
        }

        self.assertGreater(
            _document_relevance_score(actionable_check, question),
            _document_relevance_score(heading, question),
        )

    def test_parameter_request_rejects_a_related_fragment_without_parameter_fields(self):
        decision = classify_case_by_rules(
            "¿Qué parámetros se pueden modificar para riesgos de incapacidad?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Acciones de personal.pdf — Página 34",
                    ubicacion="https://contoso.example/acciones.pdf",
                    fragmento="Los días de subsidio dependen del riesgo de incapacidad.",
                )
            ],
        )

        self.assertEqual("sin_evidencia", decision.estado)
        self.assertEqual([], decision.fuentes)

    def test_calculation_request_rejects_an_include_flag_without_a_formula(self):
        decision = classify_case_by_rules(
            "¿Cómo se calcula el aguinaldo según la documentación disponible?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Readme.pdf — Página 40",
                    ubicacion="https://contoso.example/readme.pdf",
                    fragmento="GTISRIncluirAguinaldo indica si incluye aguinaldo: Si / No.",
                )
            ],
        )

        self.assertEqual("sin_evidencia", decision.estado)

    def test_calculation_request_rejects_a_parameter_with_aguinaldo_days(self):
        decision = classify_case_by_rules(
            "¿Cómo se calcula el aguinaldo en Evolution?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Readme 1.19.1.9.pdf — Página 38",
                    ubicacion="https://contoso.example/readme.pdf",
                    fragmento=(
                        "AguinaldoMesParaCalculo: mes en el que se calcula el aguinaldo. "
                        "AguinaldoNumeroDias: número de días de aguinaldo a que tiene "
                        "derecho un empleado."
                    ),
                )
            ],
        )

        self.assertEqual("sin_evidencia", decision.estado)

    def test_reinstallation_validation_rejects_a_diagnostic_heading_without_checks(self):
        decision = classify_case_by_rules(
            "Después de reinstalar MSDTC, ¿qué debo validar en ambos servidores?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Manual DTC Verificacion.pdf — Página 3",
                    ubicacion="https://contoso.example/dtc.pdf",
                    fragmento=(
                        "Validar en ambos servidores que estos servicios están corriendo\n"
                        "Base de datos:"
                    ),
                )
            ],
        )

        self.assertEqual("sin_evidencia", decision.estado)

    def test_post_update_validation_rejects_release_note_entries(self):
        decision = classify_case_by_rules(
            "Después de aplicar una actualización de Evolution, ¿qué validaciones operativas recomienda la documentación?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Readme.pdf — Página 8",
                    ubicacion="https://contoso.example/readme.pdf",
                    fragmento="EVO-4208 Validación para creación de esquema de evaluación.",
                )
            ],
        )

        self.assertEqual("sin_evidencia", decision.estado)

    def test_key_vault_request_rejects_a_generic_rest_api_key_reference(self):
        decision = classify_case_by_rules(
            "¿Cómo configuro una API key mediante Key Vault para una integración de Evolution?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Escenario de infraestructura.pdf — Página 14",
                    ubicacion="https://contoso.example/infrastructure.pdf",
                    fragmento="Trasladan el ID del sitio y REST API KEY.",
                )
            ],
        )

        self.assertEqual("sin_evidencia", decision.estado)

    def test_procedure_prefers_a_solution_instruction_over_an_incidental_configuration(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Instrucciones.txt — Documento",
                ubicacion="https://contoso.sharepoint.com/sites/x/Documentos%20compartidos/SOLUCIONES/Evolution/Reiniciar%20AppJob/Instrucciones.txt",
                fragmento="Importar la tarea desde el Task Scheduler y ejecutar la tarea creada.",
            ),
            EvidenceSource(
                tipo="sharepoint",
                titulo="Configuración SMTP.docx — Documento",
                ubicacion="https://contoso.sharepoint.com/sites/x/_layouts/15/Doc.aspx?sourcedoc=abc",
                fragmento="Paso 4 Reiniciar el servicio AppJob desde Task Scheduler.",
            ),
        ]

        selected = _focused_procedure_evidence(
            "¿Cuál es el procedimiento para reiniciar AppJob desde Task Scheduler?", evidence
        )

        self.assertEqual(["Instrucciones.txt — Documento"], [source.titulo for source in selected])

    def test_procedure_summary_excludes_internal_log_tail(self):
        answer = _grounded_document_summary(
            "¿Cuál es el procedimiento para reiniciar AppJob desde Task Scheduler?",
            [
                EvidenceSource(
                    tipo="sharepoint",
                    titulo="Instrucciones.txt — Documento",
                    ubicacion="https://contoso.sharepoint.com/sites/x/SOLUCIONES/AppJob/Instrucciones.txt",
                    fragmento=(
                        "Renombrar la tarea anterior. Importar la tarea desde Task Scheduler. "
                        "Seleccionar la tarea creada. Presionar Ejecutar y verificar el resultado. "
                        "2026-08-06 10:10:00 INFO Error de aplicación; "
                        "cadena de conexión: Server=internal-db; Database=Evolution."
                    ),
                )
            ],
        )

        self.assertIn("Presionar Ejecutar", answer)
        self.assertNotIn("INFO", answer)
        self.assertNotIn("cadena de conexión", answer.casefold())

    def test_content_coverage_accepts_configuration_variant_in_same_fragment(self):
        self.assertTrue(
            _has_minimum_content_coverage(
                {
                    "title": "Acciones de personal.pdf — Página 18",
                    "content_tokens": (
                        "parametro prorroga contrato configuracion "
                        "prorrogacontratodiasatrasiniciorangofechafincontrato"
                    ),
                    "content": (
                        "Parámetros para prórroga de contratos en la configuración "
                        "de los parámetros de infraestructura."
                    ),
                    "document_type": "pdf",
                },
                "¿Qué parámetros se pueden configurar para prórroga de contratos en Evolution?",
            )
        )

    def test_content_coverage_accepts_common_operator_vocabulary(self):
        self.assertTrue(
            _has_minimum_content_coverage(
                {
                    "title": "Gestion de documentos.pdf — Página 4",
                    "content_tokens": "gestion documento administrar descargar",
                    "content": (
                        "En esta opción permite acceder y descargar los documentos "
                        "administrados en el módulo Gestión de documentos, a los que se tienen permiso. "
                        "Si falla, revisar permisos y acceso al módulo."
                    ),
                    "document_type": "pdf",
                },
                "Un usuario tiene permisos pero no logra bajar los documentos del módulo de gestión, ¿qué se debe revisar?",
            )
        )
        self.assertTrue(
            _has_minimum_content_coverage(
                {
                    "title": "Ofuscación de datos.sql — Documento",
                    "content_tokens": "ofuscacion dato sensible sql server",
                    "content": "Procedimiento SQL Server para ofuscar datos sensibles.",
                    "document_type": "sql",
                },
                "¿Cómo se ofuscan datos sensibles en SQL Server?",
            )
        )

    def test_operator_vocabulary_does_not_make_unrelated_evidence_pass(self):
        self.assertFalse(
            _has_minimum_content_coverage(
                {
                    "title": "Readme licenciamiento.pdf — Página 2",
                    "content_tokens": "licencia token administrar usuario",
                    "content": "Administrar tokens de licencia del sistema.",
                    "document_type": "pdf",
                },
                "Un usuario no puede bajar los documentos del módulo de gestión.",
            )
        )

    def test_contract_temporal_wording_expands_to_renewal_window(self):
        self.assertTrue(
            _has_minimum_content_coverage(
                {
                    "title": "Acciones de personal.pdf — Página 18",
                    "content_tokens": "parametro prorroga contrato empleado rango dia",
                    "content": (
                        "Parámetros para prórroga de contratos. Controla el rango de días "
                        "para mostrar empleados. "
                        "ProrrogaContratoDiasAtrasInicioRangoFechaFinContrato."
                    ),
                    "document_type": "pdf",
                },
                "¿Qué parámetro controla el rango de días para mostrar empleados cuya prórroga de contrato está próxima?",
            )
        )

    def test_direct_document_request_accepts_imperative_wording(self):
        evidence = [
            EvidenceSource(
                tipo="azure_ai_search",
                titulo="Acciones de personal.pdf — Página 18",
                ubicacion="https://contoso.example/acciones.pdf",
                fragmento="Parámetros para prórroga de contratos en Evolution.",
            )
        ]
        self.assertTrue(
            is_direct_document_question(
                "Dame los parámetros de prórroga de contratos en Evolution.",
                evidence,
            )
        )

    def test_background_action_is_not_required_from_validation_fragment(self):
        question = (
            "Después de reinstalar MSDTC en un servidor clonado, ¿qué hay que revisar "
            "en ambos servidores para confirmar la comunicación DTC?"
        )
        self.assertEqual(
            _question_without_background_action(question),
            "¿qué hay que revisar en ambos servidores para confirmar la comunicación DTC?",
        )
        self.assertTrue(
            _has_minimum_content_coverage(
                {
                    "title": "Manual DTC Verificacion.pdf — Página 3",
                    "content_tokens": "validar ambos servidor servicio dtc comunicacion",
                    "content": "Validar en ambos servidores que estos servicios están corriendo.",
                    "document_type": "pdf",
                },
                question,
            )
        )
        self.assertFalse(
            has_requested_action_coverage(
                "¿Cómo se clasifican las incapacidades?",
                "Las incapacidades se registran en las mismas tablas de nómina.",
            )
        )
        self.assertFalse(
            _has_minimum_content_coverage(
                {
                    "title": "Carrera y Sucesión.pdf — Página 16",
                    "content_tokens": "carrera sucesion cuadrante modificar parametro infraestructura",
                    "content": "Los cuadrantes se pueden modificar mediante un parámetro de infraestructura.",
                },
                "¿Qué parámetros se pueden modificar para incapacidades?",
            )
        )
        self.assertFalse(
            _has_minimum_content_coverage(
                {
                    "title": "Configuracion de infraestructura.pdf — Página 44",
                    "document_context": "El documento completo contiene referencias a incapacidades.",
                    "content_tokens": "configuracion infraestructura parametro tablero indicador",
                    "content": "Parámetros del tablero: código, nombre, tipo y origen de datos.",
                },
                "¿Qué parámetros se pueden configurar para incapacidades en Evolution?",
            )
        )

    def test_pdf_records_ignore_sharepoint_files_outside_sync_state(self):
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            (source_dir / ".libras-sharepoint-sync-state.json").write_text(
                '{"documents": {"allowed-id": {"filename": "allowed.pdf"}}}',
                encoding="utf-8",
            )
            for filename, document_id in (("allowed.pdf", "allowed-id"), ("stale.pdf", "stale-id")):
                pdf_path = source_dir / filename
                pdf_path.write_bytes(b"placeholder")
                pdf_path.with_suffix(".pdf.metadata.json").write_text(
                    json.dumps(
                        {
                            "source_system": "sharepoint",
                            "document_id": document_id,
                            "web_url": f"https://contoso.example/{filename}",
                        }
                    ),
                    encoding="utf-8",
                )
            with patch(
                "azure_search._document_pages",
                return_value=[(1, "Contenido controlado.")],
            ):
                records = _document_records(source_dir)

        self.assertEqual({"allowed-id"}, {record["document_id"] for record in records})

    def test_readable_office_files_are_extracted(self):
        from docx import Document
        from openpyxl import Workbook

        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            docx_path = source_dir / "manual.docx"
            document = Document()
            document.add_paragraph("Paso documentado para instalar el módulo.")
            document.save(docx_path)

            xlsx_path = source_dir / "parametros.xlsx"
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = "Configuración"
            worksheet.append(["Parametro", "Valor"])
            worksheet.append(["Timeout", "30"])
            workbook.save(xlsx_path)

            docx_pages = _document_pages(docx_path)
            xlsx_pages = _document_pages(xlsx_path)

        self.assertIn("Paso documentado", docx_pages[0][1])
        self.assertIn("Configuración", xlsx_pages[0][1])
        self.assertIn("Timeout | 30", xlsx_pages[0][1])

    def test_chunks_bound_long_lines_for_embedding_limits(self):
        chunks = list(_chunks("x" * 13_000, size=450, max_characters=6_000))

        self.assertEqual([6_000, 6_000, 1_000], [len(chunk) for chunk in chunks])

    def test_empty_documents_are_skipped(self):
        with TemporaryDirectory() as directory:
            empty_pdf = Path(directory) / "empty.pdf"
            empty_pdf.touch()

            self.assertEqual([], _document_pages(empty_pdf))

    def test_sync_reuses_unchanged_pdf_and_preserves_pending_deletion(self):
        class FakeSharePointClient:
            drive_id = "drive-1"

            def __init__(self):
                self.files = [
                    {
                        "id": "document-1",
                        "name": "manual.pdf",
                        "eTag": '"etag-1"',
                        "webUrl": "https://contoso.example/manual.pdf",
                        "lastModifiedDateTime": "2026-07-22T12:00:00Z",
                    }
                ]
                self.download_calls = 0

            def list_pdfs(self):
                return self.files

            def download(self, _item, destination):
                self.download_calls += 1
                destination.write_bytes(b"placeholder PDF")

        fake_client = FakeSharePointClient()
        config = SimpleNamespace(
            sharepoint_site_id="site-1",
            sharepoint_folder_path="Operaciones",
        )

        with TemporaryDirectory() as directory, patch(
            "sharepoint_sync.create_sharepoint_client", return_value=fake_client
        ):
            source_dir = Path(directory)
            sync_pdfs(config, source_dir)
            sync_pdfs(config, source_dir)
            self.assertEqual(1, fake_client.download_calls)
            self.assertEqual({"document-1"}, _change_ids(source_dir / SYNC_CHANGE_MANIFEST_NAME))

            fake_client.files[0]["name"] = "manual-renombrado.pdf"
            sync_pdfs(config, source_dir)
            self.assertEqual(2, fake_client.download_calls)
            self.assertFalse((source_dir / "document-1_manual.pdf").exists())
            self.assertTrue((source_dir / "document-1_manual-renombrado.pdf").exists())

            fake_client.files = []
            sync_pdfs(config, source_dir)
            self.assertEqual({"document-1"}, _deletion_document_ids(source_dir))
            self.assertEqual(set(), _change_ids(source_dir / SYNC_CHANGE_MANIFEST_NAME))
            self.assertFalse((source_dir / "document-1_manual-renombrado.pdf").exists())

            # If Azure ingestion has not yet acknowledged the deletion, a
            # subsequent sync must retain the tombstone instead of losing it.
            sync_pdfs(config, source_dir)
            self.assertEqual({"document-1"}, _deletion_document_ids(source_dir))

    def test_index_directory_only_embeds_pending_sharepoint_changes(self):
        class FakeSearchClient:
            def __init__(self):
                self.uploaded_records = []

            def search(self, **_kwargs):
                return []

            def merge_or_upload_documents(self, documents):
                self.uploaded_records.extend(documents)
                return [
                    SimpleNamespace(succeeded=True, key=document["id"])
                    for document in documents
                ]

            def delete_documents(self, documents):
                return [
                    SimpleNamespace(succeeded=True, key=document["id"])
                    for document in documents
                ]

        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
        )
        fake_search = FakeSearchClient()
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            for document_id in ("document-1", "document-2"):
                pdf_path = source_dir / f"{document_id}.pdf"
                pdf_path.write_bytes(b"placeholder")
                pdf_path.with_suffix(".pdf.metadata.json").write_text(
                    (
                        '{"source_system": "sharepoint", '
                        f'"document_id": "{document_id}", "etag": "etag-1"}}'
                    ),
                    encoding="utf-8",
                )
            (source_dir / CHANGE_MANIFEST_NAME).write_text(
                '{"changed_document_ids": ["document-1"]}', encoding="utf-8"
            )

            with (
                patch("azure_search.SearchClient", return_value=fake_search),
                patch(
                    "azure_search._document_pages",
                    return_value=[(1, "Contenido aprobado del manual.")],
                ),
                patch("azure_search._attach_embeddings") as attach_embeddings,
            ):
                uploaded = index_directory(source_dir, config)

            self.assertEqual(1, uploaded)
            embedded_records = attach_embeddings.call_args.args[0]
            self.assertEqual(["document-1"], [record["document_id"] for record in embedded_records])
            self.assertEqual(
                ["document-1"],
                [record["document_id"] for record in fake_search.uploaded_records],
            )
            self.assertEqual(set(), _changed_document_ids(source_dir))

    def test_create_index_reindexes_every_pdf_and_acknowledges_pending_changes(self):
        class FakeSearchClient:
            def search(self, **_kwargs):
                return []

            def merge_or_upload_documents(self, documents):
                return [
                    SimpleNamespace(succeeded=True, key=document["id"])
                    for document in documents
                ]

            def delete_documents(self, documents):
                return [
                    SimpleNamespace(succeeded=True, key=document["id"])
                    for document in documents
                ]

        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
        )
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            for document_id in ("document-1", "document-2"):
                pdf_path = source_dir / f"{document_id}.pdf"
                pdf_path.write_bytes(b"placeholder")
                pdf_path.with_suffix(".pdf.metadata.json").write_text(
                    f'{{"document_id": "{document_id}"}}', encoding="utf-8"
                )
            (source_dir / CHANGE_MANIFEST_NAME).write_text(
                '{"changed_document_ids": ["document-1"]}', encoding="utf-8"
            )

            with (
                patch("azure_search.SearchClient", return_value=FakeSearchClient()),
                patch("azure_search.ensure_index"),
                patch(
                    "azure_search._document_pages",
                    return_value=[(1, "Contenido aprobado del manual.")],
                ),
                patch("azure_search._attach_embeddings") as attach_embeddings,
            ):
                uploaded = index_directory(source_dir, config, create_index=True)

            self.assertEqual(2, uploaded)
            self.assertEqual(
                {"document-1", "document-2"},
                {record["document_id"] for record in attach_embeddings.call_args.args[0]},
            )
            self.assertEqual(set(), _changed_document_ids(source_dir))

    def test_changed_existing_chunk_merges_without_reembedding(self):
        class FakeSearchClient:
            def __init__(self):
                self.merged_records = []

            def search(self, **_kwargs):
                return [{"id": "existing-chunk"}]

            def merge_documents(self, documents):
                self.merged_records.extend(documents)
                return [
                    SimpleNamespace(succeeded=True, key=document["id"])
                    for document in documents
                ]

            def delete_documents(self, documents):
                return [
                    SimpleNamespace(succeeded=True, key=document["id"])
                    for document in documents
                ]

        config = SimpleNamespace(
            azure_search_configured=True,
            azure_search_endpoint="https://search.example",
            azure_search_index_name="libras-docs",
            azure_search_api_key="not-a-real-key",
        )
        record = {
            "id": "existing-chunk",
            "document_id": "document-1",
            "title": "Indicaciones.txt — Documento",
            "content": "Solución actualizada.",
            "document_context": "Carpetas de origen: SOLUCIONES / EVOLUTION.",
            "content_tokens": "solucion evolution",
        }
        with TemporaryDirectory() as directory:
            source_dir = Path(directory)
            (source_dir / CHANGE_MANIFEST_NAME).write_text(
                '{"changed_document_ids": ["document-1"]}', encoding="utf-8"
            )
            fake_search = FakeSearchClient()
            with (
                patch("azure_search.SearchClient", return_value=fake_search),
                patch("azure_search._document_records", return_value=[record]),
                patch("azure_search._attach_embeddings") as attach_embeddings,
            ):
                uploaded = index_directory(source_dir, config)

        self.assertEqual(1, uploaded)
        self.assertEqual([record], fake_search.merged_records)
        attach_embeddings.assert_not_called()

    def test_excerpt_prefers_matching_sentence_over_chunk_start(self):
        content = (
            "Página 3 Información general de la planilla y parámetros administrativos. "
            "Los descuentos se gestionan conforme a la configuración vigente. "
            "Bono Decreto: Bono 37-2001, es de Q 250.00. Para empleados de nuevo "
            "ingreso es proporcional al tiempo laborado en el período."
        )

        excerpt = _excerpt_around_query(
            content,
            "¿Cuál es el valor del Bono Decreto 37-2001 para un empleado de nuevo ingreso?",
        )

        self.assertIn("Q 250.00", excerpt)
        self.assertIn("nuevo ingreso", excerpt)

    def test_excerpt_keeps_separate_facts_from_the_same_page(self):
        content = (
            "El aguinaldo equivale a quince días de salario después de un año continuo. "
            "Esta prestación se paga en diciembre. "
            "El aguinaldo está exento de renta hasta 30 UMA."
        )

        excerpt = _excerpt_around_query(
            content,
            "¿A cuántos días de salario equivale el aguinaldo y hasta cuánto está exento de renta?",
        )

        self.assertIn("quince días", excerpt)
        self.assertIn("30 UMA", excerpt)

    def test_calculation_question_keeps_the_formula(self):
        content = (
            "El aguinaldo se calcula proporcionalmente al tiempo trabajado. "
            "La empresa realiza el pago en diciembre. "
            "Ejemplo: días de aguinaldo = 15 * días trabajados / días del año. "
            "La fórmula de pago es salario diario * días de aguinaldo."
        )

        excerpt = _excerpt_around_query(
            content,
            "¿Cómo se calcula proporcionalmente el aguinaldo?",
        )

        self.assertIn("fórmula de pago", excerpt)

    def test_direct_document_question_is_resolved(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Políticas de Pago MEXICO — Página 14",
                ubicacion="https://contoso.example/politicas-mexico.pdf",
                fragmento=(
                    "El aguinaldo equivale a quince días de salario después de un año "
                    "de trabajo continuo y está exento de renta hasta 30 UMA."
                ),
            )
        ]

        decision = classify_case_by_rules(
            "En México, ¿a cuántos días de salario equivale el aguinaldo después de un año continuo y hasta cuánto está exento de renta?",
            evidence,
        )

        self.assertEqual("resuelto", decision.estado)
        self.assertFalse(decision.requiere_escalamiento)

    def test_version_question_uses_the_retrieved_text_instead_of_a_generic_incident_message(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Readme 1.19.1.10.pdf — Página 1",
                ubicacion="https://contoso.example/readme-1.19.1.10.pdf",
                fragmento="La versión incorpora el ajuste de instalación y requiere actualizar AppSettings.config.",
            )
        ]

        decision = classify_case_by_rules(
            "Dame detalles sobre la versión de Evolution 1.19.1.10.", evidence
        )

        self.assertEqual("resuelto", decision.estado)
        self.assertIn("1.19.1.10", decision.resumen)
        self.assertIn("actualizar AppSettings.config", decision.resumen)
        self.assertNotIn("describe el problema", decision.resumen)

    def test_versioned_software_requirements_returns_none_when_documented(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Readme 1.19.1.11.pdf — Página 8",
                ubicacion="https://contoso.example/readme-1.19.1.11.pdf",
                fragmento="Nuevos requisitos de software Ninguno.",
            )
        ]

        decision = classify_case_by_rules(
            "¿Qué nuevos requisitos de software necesita Evolution versión Readme 1.19.1.11?",
            evidence,
        )

        self.assertEqual("Ninguno.", decision.resumen)

    def test_versioned_software_requirements_rejects_a_changelog_without_requirements(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Readme 1.19.1.11.pdf — Página 1",
                ubicacion="https://contoso.example/readme-1.19.1.11.pdf",
                fragmento="Mejoras de la versión: se agregó el detalle de plazas y dependencias.",
            )
        ]

        decision = classify_case_by_rules(
            "Para la versión 1.19.1.11, ¿se necesita algún requisito nuevo de software?",
            evidence,
        )

        self.assertEqual("sin_evidencia", decision.estado)
        self.assertEqual([], decision.fuentes)

    def test_document_version_procedure_rejects_document_type_as_tangential_evidence(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Readme 1.19.1.13.pdf — Página 7",
                ubicacion="https://contoso.example/readme-1.19.1.13.pdf",
                fragmento="Los valores esperados son CrystalReports, WordTemplate y DocumentoGestionado.",
            )
        ]

        decision = classify_case_by_rules(
            "¿Me puedes explicar de forma breve cómo crear una nueva versión de un documento gestionado?",
            evidence,
        )

        self.assertEqual("sin_evidencia", decision.estado)
        self.assertEqual([], decision.fuentes)

    def test_parameter_list_request_rejects_a_fragment_without_parameter_details(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Acciones de personal.pdf — Página 34",
                ubicacion="https://contoso.example/acciones-personal.pdf",
                fragmento=(
                    "Los días de descuento o subsidio se configuran a partir de los "
                    "riesgos de incapacidad según el Seguro Social."
                ),
            )
        ]

        decision = classify_case_by_rules(
            "Necesito configurar riesgos de incapacidad. ¿Qué parámetros se pueden modificar?",
            evidence,
        )

        self.assertEqual("sin_evidencia", decision.estado)
        self.assertEqual([], decision.fuentes)

    def test_sql_evidence_is_summarized_without_dumping_executable_code(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Ofuscación de datos.sql",
                ubicacion="https://contoso.example/ofuscacion.sql",
                fragmento="UPDATE exp_expedientes SET exp_nombre = @rnd;",
                document_type="sql",
            )
        ]

        decision = classify_case_by_rules("¿Cómo se ofuscan datos sensibles en SQL?", evidence)

        self.assertEqual("resuelto", decision.estado)
        self.assertIn("script técnico", decision.resumen)
        self.assertNotIn("UPDATE", decision.resumen)

    def test_sql_description_is_used_in_grounded_summary(self):
        from classification import _grounded_document_summary

        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="acc.proc_arreglar_vac_negativos.sql — Documento",
                ubicacion="https://contoso.example/vacaciones.sql",
                fragmento="Descripción de la solución: Corrige periodos de vacaciones con saldo negativo.",
                document_type="sql",
                descripcion="Corrige los periodos de vacaciones que tienen saldo negativo y mueve el tiempo a nuevos periodos",
            )
        ]

        summary = _grounded_document_summary(
            "algún script para corregir vacaciones con saldo negativo", evidence
        )

        self.assertIn("acc.proc_arreglar_vac_negativos.sql", summary)
        self.assertIn("saldo negativo", summary)
        self.assertNotIn("CREATE PROCEDURE", summary)

    def test_diagnostic_request_combines_complete_checks_from_multiple_sources(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Manual DTC Verificacion.pdf — Página 3",
                ubicacion="https://contoso.example/dtc.pdf",
                fragmento=(
                    "Valide en ambos servidores que el servicio Distributed Transaction "
                    "Coordinator esté en ejecución."
                ),
            ),
            EvidenceSource(
                tipo="sharepoint",
                titulo="Manual DTC Verificacion.pdf — Página 4",
                ubicacion="https://contoso.example/dtc.pdf",
                fragmento=(
                    "Verifique que el firewall permita la comunicación DTC entre ambos servidores."
                ),
            ),
        ]

        decision = classify_case_by_rules(
            "¿Qué servicios o validaciones debo revisar para que la comunicación DTC quede funcionando?",
            evidence,
        )

        self.assertEqual("resuelto", decision.estado)
        self.assertIn("1. Valide en ambos servidores", decision.resumen)
        self.assertIn("2. Verifique que el firewall", decision.resumen)

    def test_diagnostic_list_request_rejects_an_incomplete_heading(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Manual DTC Verificacion.pdf — Página 3",
                ubicacion="https://contoso.example/dtc.pdf",
                fragmento="Validar en ambos servidores que estos servicios están corriendo: Base de datos:",
            )
        ]

        decision = classify_case_by_rules(
            "¿Qué servicios o validaciones debo revisar para que la comunicación DTC quede funcionando?",
            evidence,
        )

        self.assertEqual("sin_evidencia", decision.estado)
        self.assertEqual([], decision.fuentes)

    def test_important_numeric_terms_are_searchable(self):
        self.assertIn("37", tokenize("Bono Decreto 37-2001"))

    def test_plural_terms_normalize_without_altering_acronyms(self):
        tokens = tokenize("Empleados ISSS")
        self.assertIn("empleado", tokens)
        self.assertIn("isss", tokens)

    def test_inflected_terms_normalize_for_retrieval(self):
        self.assertIn("proporcional", tokenize("proporcionalmente proporcionales"))

    def test_camel_case_application_fields_are_searchable_by_concept(self):
        tokens = tokenize("BaseCalculoISSS")

        self.assertTrue({"base", "calculo", "isss"}.issubset(tokens))

    def test_specific_planilla_phrase_outranks_liquidation_page(self):
        question = "En la planilla quincenal, ¿cómo se aplican el ISSS, AFP e impuesto sobre la renta?"
        quincenal_page = {
            "title": "Políticas de Pago SV — Página 2",
            "content": (
                "Planilla Quincenal. Los descuentos de ley ISSS, AFP e Impuesto sobre "
                "la Renta serán aplicados en cada quincena con ajuste mensual."
            ),
        }
        liquidation_page = {
            "title": "Políticas de Pago SV — Página 14",
            "content": "En liquidación se aplican AFP, ISSS e Impuesto sobre la Renta.",
        }

        self.assertGreater(
            _document_relevance_score(quincenal_page, question),
            _document_relevance_score(liquidation_page, question),
        )

        ranked = _rerank_records([liquidation_page, quincenal_page], question)
        self.assertEqual(quincenal_page, ranked[0][1])

    def test_generic_coverage_prioritizes_calculation_page(self):
        question = "¿Cómo se aplica el ISR quincenal y qué descuentos se restan de su base?"
        tax_page = {
            "title": "Políticas de Pago SV — Página 6",
            "content": "ISR BaseCalculoRenta. Tabla de Renta Quincenal. Descuentos AFP e ISSS.",
        }
        monthly_page = {
            "title": "Políticas de Pago SV — Página 8",
            "content": "Planilla mensual: ISR, AFP e ISSS. Se usa el mismo agrupador de la planilla quincenal.",
        }

        ranked = _rerank_records([monthly_page, tax_page], question)
        self.assertEqual(tax_page, ranked[0][1])

    def test_vector_rank_breaks_a_lexical_tie(self):
        question = "¿Cómo se calcula el aguinaldo proporcional?"
        first_vector_result = {
            "title": "México — Página 9",
            "content": "El aguinaldo proporcional se calcula según los días laborados.",
            "_vector_rank": 1,
        }
        later_vector_result = {
            "title": "Guatemala — Página 7",
            "content": "El aguinaldo proporcional se calcula según los días laborados.",
            "_vector_rank": 20,
        }

        ranked = _rerank_records([later_vector_result, first_vector_result], question)

        self.assertEqual(first_vector_result, ranked[0][1])

    def test_semantic_reranking_accepts_a_supported_action_paraphrase(self):
        record = {
            "title": "Gestión de documentos.pdf — Página 1",
            "content": "Evolution permite gestionar documentos asociados a colaboradores.",
            "@search.reranker_score": 2.3,
        }

        self.assertTrue(
            _has_minimum_content_coverage(
                record,
                "¿Cómo se pueden administrar documentos en Evolution?",
                semantic_enabled=True,
            )
        )

    def test_candidate_diversification_limits_repeated_chunks_of_one_document(self):
        records = [
            {"id": f"manual-{number}", "document_id": "manual", "_vector_rank": number}
            for number in range(1, 6)
        ] + [
            {"id": "guide-1", "document_id": "guide", "_vector_rank": 6},
        ]

        diversified = _diversify_candidate_records(records)

        self.assertEqual(3, sum(record["document_id"] == "manual" for record in diversified))
        self.assertIn("guide", {record["document_id"] for record in diversified})

    def test_country_specific_request_does_not_mix_evidence(self):
        guatemala_record = {
            "title": "Guatemala — Página 1",
            "document_context": "Procedimiento autorizado para Guatemala.",
            "content": "Guatemala. Ajuste de sesión.",
        }
        salvador_record = {
            "title": "El Salvador — Página 1",
            "document_context": "Procedimiento autorizado para El Salvador.",
            "content": "El Salvador. Ajuste de sesión.",
        }

        filtered = _filter_records_for_requested_country(
            [salvador_record, guatemala_record],
            "¿Cómo se realiza el ajuste de sesión en Guatemala?",
        )

        self.assertEqual([guatemala_record], filtered)

    def test_country_specific_request_without_matching_evidence_returns_no_records(self):
        generic_record = {
            "title": "Manual general",
            "document_context": "Procedimiento técnico sin país.",
            "content": "Ajuste de sesión.",
        }

        filtered = _filter_records_for_requested_country(
            [generic_record], "¿Cómo se realiza el ajuste de sesión en El Salvador?"
        )

        self.assertEqual([], filtered)

    def test_country_request_rejects_a_multi_country_contact_footer(self):
        shared_footer_record = {
            "title": "Cambios generales de Evolution",
            "document_context": "GUATEMALA oficina central. EL SALVADOR oficina regional.",
            "content": "Cambios de interfaz de Evolution.",
        }

        filtered = _filter_records_for_requested_country(
            [shared_footer_record],
            "¿Qué cambios existen para El Salvador en Evolution?",
        )

        self.assertEqual([], filtered)

    def test_tangential_single_term_hit_is_not_sufficient_evidence(self):
        vacation_policy = {
            "title": "Políticas de pago",
            "content": "Las vacaciones pendientes se procesan al retiro del empleado.",
        }

        self.assertFalse(
            _has_minimum_content_coverage(
                vacation_policy,
                "¿Cuál es el procedimiento oficial de Recursos Humanos para aprobar vacaciones?",
            )
        )
        self.assertTrue(
            _has_minimum_content_coverage(
                vacation_policy,
                "¿Cómo se procesan las vacaciones pendientes?",
            )
        )

    def test_vacation_guide_without_approval_action_is_not_evidence_of_approval_procedure(self):
        technical_guide = {
            "title": "Guía de temas técnicos Evolution",
            "content": (
                "Creación y modificación de vacaciones. Para crear un periodo de "
                "vacaciones, asigne el año inicial y valide el empleo."
            ),
        }
        approval_procedure = {
            "title": "Procedimiento de aprobación de vacaciones",
            "content": (
                "El jefe aprueba la solicitud de vacaciones y Recursos Humanos "
                "confirma la aprobación antes de registrar el movimiento."
            ),
        }
        question = "¿Cuál es el procedimiento oficial para aprobar vacaciones?"

        self.assertFalse(_has_minimum_content_coverage(technical_guide, question))
        self.assertTrue(_has_minimum_content_coverage(approval_procedure, question))

    def test_related_vacation_evidence_without_approval_action_is_not_classified_as_resolved(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Guía de temas técnicos Evolution",
                ubicacion="https://contoso.example/guia.docx",
                fragmento="Creación y modificación de vacaciones para el empleo.",
            )
        ]

        decision = classify_case_by_rules(
            "¿Cuál es el procedimiento oficial para aprobar vacaciones?", evidence
        )

        self.assertEqual("sin_evidencia", decision.estado)
        self.assertEqual([], decision.fuentes)

    def test_out_of_scope_question_requires_three_specific_terms(self):
        human_resources_capacity = {
            "title": "Requerimientos de Evolution",
            "content": "Usuarios concurrentes del departamento de Recursos Humanos.",
        }

        self.assertFalse(
            _has_minimum_content_coverage(
                human_resources_capacity,
                "¿Cuál es el procedimiento oficial de Recursos Humanos para aprobar vacaciones?",
            )
        )

    def test_only_https_sharepoint_records_are_authorized_evidence(self):
        self.assertTrue(
            _record_has_authorized_provenance(
                {"source_system": "sharepoint", "source_url": "https://contoso.example/manual.pdf"}
            )
        )
        self.assertFalse(
            _record_has_authorized_provenance(
                {"source_system": "web", "source_url": "https://unapproved.example/manual.pdf"}
            )
        )
        self.assertFalse(
            _record_has_authorized_provenance(
                {"source_system": "sharepoint", "source_url": "http://contoso.example/manual.pdf"}
            )
        )
        approved_folder = ("SOLUCIONES",)
        self.assertTrue(
            _record_has_authorized_provenance(
                {
                    "source_system": "sharepoint",
                    "source_url": "https://contoso.example/manual.pdf",
                    "folder_path": "SOLUCIONES",
                },
                approved_folder,
            )
        )
        self.assertFalse(
            _record_has_authorized_provenance(
                {
                    "source_system": "sharepoint",
                    "source_url": "https://contoso.example/manual.pdf",
                    "folder_path": "ReadME Hotfixes",
                },
                approved_folder,
            )
        )

    def test_multi_library_provenance_accepts_only_approved_drive_and_path(self):
        approved_sources = (("", "drive-readme"), ("SOLUCIONES", "drive-documents"))
        base_record = {
            "source_system": "sharepoint",
            "source_url": "https://contoso.example/manual.pdf",
        }

        self.assertTrue(
            _record_has_authorized_provenance(
                {**base_record, "drive_id": "drive-readme", "folder_path": ""},
                approved_sources,
            )
        )
        self.assertTrue(
            _record_has_authorized_provenance(
                {
                    **base_record,
                    "drive_id": "drive-documents",
                    "folder_path": "SOLUCIONES",
                },
                approved_sources,
            )
        )
        self.assertFalse(
            _record_has_authorized_provenance(
                {**base_record, "drive_id": "drive-other", "folder_path": ""},
                approved_sources,
            )
        )

    def test_legacy_provenance_fails_closed_for_unscoped_sharepoint_urls(self):
        approved_sources = (("", "drive-readme"), ("SOLUCIONES", "drive-documents"))
        labels = ("ReadME Hotfixes", "Documentos/SOLUCIONES")

        self.assertTrue(
            _record_has_authorized_provenance(
                {
                    "source_system": "sharepoint",
                    "source_url": "https://contoso.sharepoint.com/sites/x/ReadME%20Hotfixes/Readme.pdf",
                },
                approved_sources,
                labels,
            )
        )
        self.assertFalse(
            _record_has_authorized_provenance(
                {
                    "source_system": "sharepoint",
                    "source_url": "https://contoso.sharepoint.com/sites/x/_layouts/15/Doc.aspx?sourcedoc=abc",
                },
                approved_sources,
                labels,
            )
        )

    def test_excerpt_prefers_a_numbered_procedure_step_over_prior_configuration(self):
        excerpt = _excerpt_around_query(
            "Host: smtp.office365.com Password: secreto de ejemplo Paso 3 Se guarda la información "
            "del archivo Paso 4 Reiniciar el servicio AppJob desde Task Scheduler.",
            "El servicio AppJob está detenido, ¿qué indica la guía para reiniciarlo?",
            limit=500,
        )

        self.assertIn("Reiniciar el servicio AppJob", excerpt)
        self.assertNotIn("smtp.office365.com", excerpt)

    def test_unmatched_question_remains_without_evidence(self):
        evidence = [
            EvidenceSource(
                tipo="sharepoint",
                titulo="Políticas de Pago MEXICO — Página 14",
                ubicacion="https://contoso.example/politicas-mexico.pdf",
                fragmento="El aguinaldo está exento de renta hasta 30 UMA.",
            )
        ]

        decision = classify_case_by_rules(
            "¿Cuál es la fecha de vencimiento del certificado SSL?",
            evidence,
        )

        self.assertEqual("sin_evidencia", decision.estado)

    def test_teams_response_shows_only_answer_and_brief_source(self):
        response = format_user_response(
            BotDecision(
                estado="resuelto",
                confianza="alta",
                resumen="El aguinaldo equivale a quince días de salario.",
                fuentes=[
                    EvidenceSource(
                        tipo="azure_ai_search",
                        titulo="Políticas de Pago MEXICO — Página 14",
                        ubicacion="C:/datos/politicas.pdf",
                        fragmento="Texto de evidencia interno.",
                    )
                ],
                siguiente_accion="No aplica.",
            )
        )

        self.assertIn("El aguinaldo equivale", response)
        self.assertIn("Políticas de Pago MEXICO — Página 14", response)
        self.assertIn("Azure AI Search", response)
        self.assertNotIn("Estado", response)
        self.assertNotIn("Confianza", response)
        self.assertNotIn("Ruta de investigacion", response)

    def test_response_includes_verifiable_source_link(self):
        response = format_user_response(
            BotDecision(
                estado="resuelto",
                confianza="alta",
                resumen="El procedimiento está documentado.",
                fuentes=[
                    EvidenceSource(
                        tipo="sharepoint",
                        titulo="Procedimiento de actualización",
                        ubicacion="https://contoso.example/procedimiento.pdf",
                        fragmento="Pasos aprobados.",
                    )
                ],
            )
        )

        self.assertIn(
            "Enlace: [Ver documento: Procedimiento de actualización]"
            "(https://contoso.example/procedimiento.pdf)",
            response,
        )
        self.assertNotIn("Enlace: https://contoso.example/procedimiento.pdf", response)

    def test_sharepoint_solution_includes_its_related_files_folder_link(self):
        response = format_user_response(
            BotDecision(
                estado="resuelto",
                confianza="alta",
                resumen="Renombre el archivo y ejecútelo como PowerShell.",
                fuentes=[
                    EvidenceSource(
                        tipo="sharepoint",
                        titulo="Instrucciones.txt — Documento",
                        ubicacion=(
                            "https://aseinfocorp.sharepoint.com/sites/Soportealcliente/"
                            "Documentos%20compartidos/SOLUCIONES/EVOLUTION/Reiniciar%20appjob/"
                            "Instrucciones.txt"
                        ),
                        fragmento="Renombrar el archivo a Reiniciar appjob.ps1.",
                    )
                ],
            ),
            config=SimpleNamespace(
                sharepoint_folder_ctid="0x0120009FAB9B5A94350F489104FB62DC2E926D"
            ),
        )

        self.assertIn(
            "Enlace: [Ver documento: Instrucciones.txt]"
            "(https://aseinfocorp.sharepoint.com",
            response,
        )
        self.assertIn("Archivos relacionados: [Abrir carpeta relacionada](https://aseinfocorp.sharepoint.com", response)
        self.assertIn(
            "Forms/AllItems.aspx?FolderCTID=0x0120009FAB9B5A94350F489104FB62DC2E926D"
            "&id=%2Fsites%2FSoportealcliente%2FDocumentos%20compartidos%2FSOLUCIONES%2F"
            "EVOLUTION%2FReiniciar%20appjob",
            response,
        )

    def test_no_evidence_response_does_not_show_tangential_source(self):
        response = format_user_response(
            BotDecision(
                estado="sin_evidencia",
                confianza="baja",
                resumen="No se encontró evidencia suficiente.",
                fuentes=[
                    EvidenceSource(
                        tipo="documento",
                        titulo="Documento no relacionado",
                        ubicacion="docs/changelog.md",
                        fragmento="Texto no relacionado.",
                    )
                ],
            )
        )

        self.assertEqual("No se encontró evidencia suficiente.", response)

    def test_local_fallback_is_not_labeled_as_azure(self):
        response = format_user_response(
            BotDecision(
                estado="resuelto",
                confianza="media",
                resumen="Respuesta desde el respaldo local.",
                fuentes=[
                    EvidenceSource(
                        tipo="documento",
                        titulo="Documento local",
                        ubicacion="docs/prueba.md",
                        fragmento="Texto de respaldo.",
                    )
                ],
            )
        )

        self.assertIn("Base documental local", response)
        self.assertNotIn("Azure AI Search", response)


if __name__ == "__main__":
    unittest.main()
