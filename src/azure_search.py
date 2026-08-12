"""Proveedor e ingesta documental para Azure AI Search.

La búsqueda y la carga usan la misma clave del servicio durante el MVP. En un
entorno compartido se deben separar una clave de consulta y una identidad con
el rol ``Search Index Data Contributor`` para la ingesta.
"""

from __future__ import annotations

import hashlib
import json
import math
import re
from urllib.parse import unquote, urlparse
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from typing import Iterable

from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import ResourceNotFoundError
from azure.identity import DefaultAzureCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    HnswAlgorithmConfiguration,
    SearchFieldDataType,
    SearchField,
    SearchIndex,
    SearchableField,
    SemanticConfiguration,
    SemanticField,
    SemanticPrioritizedFields,
    SemanticSearch,
    SimpleField,
    VectorSearch,
    VectorSearchProfile,
)
from azure.search.documents.models import VectorizedQuery
from openai import OpenAI
from pypdf import PdfReader

from document_index import has_requested_action_coverage, tokenize
from evidence_verifier import verify_semantic_evidence
from logging_utils import get_logger
from models import EvidenceSource, RetrievalTrace
from query_plan import (
    QueryPlan,
    build_query_plan,
    concept_keys,
    covered_requirements,
    requirement_is_covered,
)


logger = get_logger()


# Solo se ingieren formatos con texto recuperable. Los binarios de la carpeta
# (imágenes, vídeos, ejecutables y comprimidos) permanecen fuera del índice.
SUPPORTED_EXTENSIONS = {
    ".aspx",
    ".bat",
    ".csv",
    ".docx",
    ".json",
    ".pdf",
    ".ps1",
    ".rdlc",
    ".sql",
    ".txt",
    ".xlsx",
    ".xml",
}
CONTENT_FIELD = "content"
CONTEXT_FIELD = "document_context"
CONTENT_VECTOR_FIELD = "content_vector"
RETRIEVAL_TEXT_FIELD = "retrieval_text"
RETRIEVAL_CONCEPTS_FIELD = "retrieval_concepts"
SEARCH_TIMEOUT_SECONDS = 10
MAX_CANDIDATES = 30
# The first candidate pool must be wider than the final answer set. Otherwise
# a large spreadsheet, SQL dump or manual can occupy all nearest-neighbour
# slots and hide a relevant smaller document before reranking can inspect it.
CANDIDATE_POOL_SIZE = 100
MAX_CANDIDATES_PER_DOCUMENT = 3
# Bound the merged result of the lexical, focused and vector passes before any
# evaluator sees it. The smaller reranking pool is intentionally separate from
# Azure's retrieval pool so secondary candidates remain available for review.
MAX_MERGED_CANDIDATES = 60
RERANK_POOL_SIZE = 20
_CANDIDATE_RANK_FIELDS = (
    "_keyword_rank",
    "_focused_keyword_rank",
    "_release_readme_rank",
    "_vector_rank",
    "_prefix_rank",
    "_script_rank",
)
SEMANTIC_ACTION_MIN_SCORE = 2.0
_UPPERCASE_ACRONYM = re.compile(r"\b[A-Z][A-Z0-9]{1,7}\b")
_DOCUMENT_INJECTION_PATTERN = re.compile(
    r"(?:ignore\s+(?:all\s+)?(?:previous|prior|above)\s+instructions?"
    r"|disregard\s+(?:all\s+)?instructions?"
    r"|reveal\s+(?:the\s+)?(?:system|developer)\s+prompt"
    r"|system\s+message\s*:\s*"
    r"|developer\s+message\s*:\s*"
    r"|you\s+are\s+(?:chatgpt|an\s+ai\s+assistant))",
    re.IGNORECASE,
)
_PREINSTALLATION_OPERATION_CONCEPTS = {"instal", "actualiz"}
_PREINSTALLATION_CUE_CONCEPTS = {"antes", "precaucion", "previo"}
_PREINSTALLATION_EVIDENCE_CONCEPTS = {"recomend", "inicial", "respald", "prepar", "previo"}
_LEGACY_PREINSTALLATION_OPERATION_PREFIXES = ("instal", "actualiz")
_LEGACY_PREINSTALLATION_CUE_PREFIXES = (
    "antes",
    "precaucion",
    "previa",
    "previo",
    "respald",
    "prepar",
)
_LEGACY_PREINSTALLATION_EVIDENCE_TOKENS = {
    "preparacion",
    "respaldo",
    "recomendacion",
    "inicial",
    "previa",
    "previo",
}
_RELEASE_GUIDANCE_TITLE_MARKERS = ("readme", "release", "hotfix")
_DOCUMENT_ACCESS_PERMISSION_PREFIXES = ("permis", "autoriz", "acces", "credencial")
_DOCUMENT_ACCESS_DOWNLOAD_PREFIXES = ("baj", "descarg")
_DOCUMENT_ACCESS_DIAGNOSTIC_PREFIXES = (
    "revis", "verific", "valid", "configur", "error", "fall", "proble", "soluc",
)
_DTC_VALIDATION_EVIDENCE_TOKENS = {
    "firewall",
    "component",
    "inboud",
    "outboud",
    "regla",
}
SCRIPT_REQUEST_PATTERN = re.compile(
    r"\bscript\w*\b|\bprocedimiento\s+almacenado\b|\barchivo\s+(?:sql|ps1|bat)\b",
    re.IGNORECASE,
)
# Query-side vocabulary bridges for common operator language.  These do not
# alter indexed content; they let an existing index match the wording users
# actually type without requiring an immediate reindex.
QUERY_SYNONYM_GROUPS = (
    frozenset({"bajar", "bajan", "baje", "descargar", "descarga", "descargue", "descargan"}),
    # Keep the common conjugations together.  The legacy retrieval path builds
    # its focused lexical query from the tokens typed by the operator, so
    # recognising only the infinitive made "cómo se administran documentos"
    # behave differently from "cómo se pueden administrar documentos".
    frozenset({
        "gestionar", "gestion", "gestiona", "gestionan", "gestione", "gestionando",
        "administrar", "administrado", "administra", "administran", "administre",
        "administrando",
    }),
    frozenset({"ofuscar", "ofuscacion", "ofuscado", "ofuscan"}),
)
DELETION_MANIFEST_NAME = ".libras-sharepoint-deletions.json"
CHANGE_MANIFEST_NAME = ".libras-sharepoint-changes.json"
SYNC_STATE_NAME = ".libras-sharepoint-sync-state.json"
# ``folder_path`` and ``drive_id`` are part of the provenance boundary:
# production retrieval must reject records outside the approved SharePoint
# libraries/folders.
SEARCH_SELECT_FIELDS = [
    "id",
    "title",
    "source_url",
    "source_system",
    "folder_path",
    "drive_id",
    "document_type",
    CONTEXT_FIELD,
    CONTENT_FIELD,
    "content_tokens",
    "chunk_number",
]
V2_SEARCH_SELECT_FIELDS = [
    *SEARCH_SELECT_FIELDS,
    "document_id",
    "document_version",
    "last_modified",
    "document_type",
    RETRIEVAL_TEXT_FIELD,
    RETRIEVAL_CONCEPTS_FIELD,
    "product",
    "module",
    "operation",
    "artifact_role",
    "version",
    "country",
    "quality_status",
    "evidence_kind",
]
V2_ONLY_INDEX_FIELDS = {
    RETRIEVAL_TEXT_FIELD,
    RETRIEVAL_CONCEPTS_FIELD,
    "product",
    "module",
    "operation",
    "artifact_role",
    "version",
    "country",
    "quality_status",
    "evidence_kind",
}
EXCLUDED_QUALITY_STATUSES = {"obsoleto", "duplicado", "fuera_de_alcance"}
ARTIFACT_ROLE_BY_EXTENSION = {
    ".sql": "script",
    ".ps1": "script",
    ".bat": "script",
    ".pdf": "manual",
    ".docx": "manual",
    ".xlsx": "reporte",
    ".csv": "reporte",
    ".rdlc": "reporte",
    ".json": "configuracion",
    ".xml": "configuracion",
    ".aspx": "configuracion",
    ".txt": "procedimiento",
}
COUNTRY_TOKENS = {
    "guatemala": {"guatemala", "guatemalteco", "guatemalteca"},
    "el_salvador": {"salvador", "salvadoreno", "salvadoreño", "salvadorena", "salvadoreña"},
}
# Terms that make a country-scoped query dependent on authoritative legal or
# payroll evidence. They are domain anchors, not aliases for any one question.
COUNTRY_SENSITIVE_ANCHORS = {
    "legal",
    "legislacion",
    "descuento",
    "aguinaldo",
    "impuesto",
    "renta",
    "nomina",
}
OPERATIONAL_QUERY_TOKENS = {
    "administrar",
    "aplicar",
    "arreglar",
    "calcular",
    "clasificar",
    "configurar",
    "crear",
    "eliminar",
    "gestionar",
    "instalar",
    "modificar",
    "pagar",
    "parametro",
    "procesar",
    "restaurar",
    "actualizar",
}
NAVIGATION_QUERY_TOKENS = {"indice", "tabla", "seccion", "contenido", "capitulo"}
DIAGNOSTIC_ACTION_TOKENS = {"revisar", "validar", "verificar", "confirmar"}
GENERIC_QUERY_TOKENS = {
    "cambio",
    "como",
    "configur",
    "document",
    "documentacion",
    "dame",
    "existir",
    "hacer",
    "informacion",
    "oficial",
    "paso",
    "puede",
    "pueden",
    "podria",
    "podrian",
    "procedimiento",
    "realizar",
    "tema",
}
_EXPLICIT_FILENAME_PATTERN = re.compile(
    r"(?<![\w-])([\w.-]+(?:" + "|".join(re.escape(extension) for extension in SUPPORTED_EXTENSIONS) + r"))(?![\w-])",
    re.IGNORECASE,
)
# A period that begins a file extension or ends a sentence is not part of the
# version. Only a following decimal component would make it a longer version.
_VERSION_PATTERN = re.compile(r"(?<![\d.])(\d+(?:\.\d+){2,})(?!\d|\.\d)")
_BACKGROUND_ACTION_CLAUSE = re.compile(
    r"(?:^|[.?!]\s*)(?:despu[eé]s)\s+de\s+[^,;:.!?]{1,220}[,;:]\s*",
    re.IGNORECASE,
)


def _question_without_background_action(user_message: str) -> str:
    """Remove a leading temporal circumstance from the evidence requirement.

    In ``Después de reiniciar X, ¿qué debo revisar?``, restarting X explains
    when the question happens; the evidence sought is the verification step.
    Keeping the circumstance as a mandatory per-page term rejects valid
    multi-page procedures whose validation is documented separately.
    """
    return _BACKGROUND_ACTION_CLAUSE.sub(" ", user_message or "").strip()


@lru_cache(maxsize=1)
def _entra_credential():
    """Reuse Entra token state instead of prompting for every local query."""
    return DefaultAzureCredential(exclude_interactive_browser_credential=False)


def _credential(config):
    """Use the explicitly selected identity mode, avoiding stale keys."""
    if getattr(config, "azure_search_use_entra_id", False):
        return _entra_credential()
    if getattr(config, "azure_search_api_key", ""):
        return AzureKeyCredential(config.azure_search_api_key)
    raise RuntimeError("Falta AZURE_SEARCH_API_KEY o AZURE_SEARCH_USE_ENTRA_ID=true.")


def _embedding_client(config) -> OpenAI:
    return OpenAI(
        api_key=config.openai_api_key,
        base_url=getattr(
            config,
            "resolved_openai_base_url",
            getattr(config, "openai_base_url", "") or "https://api.openai.com/v1",
        ),
    )


def _embed_texts(texts: list[str], config, client=None) -> list[list[float]]:
    embedding_client = client or _embedding_client(config)
    response = embedding_client.embeddings.create(
        model=config.openai_embedding_model,
        input=texts,
        dimensions=config.openai_embedding_dimensions,
    )
    return [item.embedding for item in response.data]


def _attach_embeddings(records: list[dict], config) -> None:
    """Add one embedding per chunk before it is sent to Azure AI Search."""
    # Keep the request well below the provider's aggregate token limit when
    # indexing code-heavy files with many tokens per character.
    for offset in range(0, len(records), 20):
        batch = records[offset : offset + 20]
        embeddings = _embed_texts(
            [
                "\n".join(
                    str(record.get(field) or "")
                    for field in ("title", RETRIEVAL_TEXT_FIELD, CONTEXT_FIELD, CONTENT_FIELD)
                )
                for record in batch
            ],
            config,
        )
        if len(embeddings) != len(batch):
            raise RuntimeError("No se recibió un embedding para cada fragmento.")
        for record, embedding in zip(batch, embeddings):
            record[CONTENT_VECTOR_FIELD] = embedding


def _clean_text(text: str, limit: int = 500) -> str:
    compact = " ".join(text.split())
    if len(compact) <= limit:
        return compact
    return f"{compact[:limit].rsplit(' ', 1)[0]}..."


def _searchable_filename_terms(title: str) -> str:
    """Expose filename conventions as searchable words without changing facts."""
    separated = re.sub(r"(?<=[a-záéíóúñ])(?=[A-ZÁÉÍÓÚÑ])", " ", title or "")
    separated = re.sub(r"[_\-.]+", " ", separated)
    return " ".join(separated.split())


def _sharepoint_parent_context(source_url: str, folder_path: str = "") -> str:
    """Return the factual parent-folder trail for retrieval only.

    SharePoint procedures are frequently stored in files named
    ``Indicaciones.txt`` or ``Script.sql``.  The diagnosis lives in the parent
    folder name, so preserve that path as searchable metadata without treating
    it as a procedural instruction or adding it to the answer excerpt.
    """
    path = unquote(urlparse(str(source_url or "")).path)
    segments = [segment.strip() for segment in path.split("/") if segment.strip()]
    if segments:
        segments.pop()  # file name
    normalized_segments = [segment.casefold() for segment in segments]
    if "documentos compartidos" in normalized_segments:
        start = normalized_segments.index("documentos compartidos") + 1
        segments = segments[start:]
    elif "sites" in normalized_segments:
        start = normalized_segments.index("sites") + 2  # site collection + site name
        segments = segments[start:]
    if not segments and folder_path:
        segments = [folder_path]
    return " / ".join(segments)


def _excerpt_around_query(text: str, query: str, limit: int = 1_000) -> str:
    """Return the most relevant part of a chunk instead of its first characters."""
    compact = " ".join(text.split())
    query_tokens = set(tokenize(query))
    # Short chunks can still contain two unrelated headings. Keep the old
    # whole-chunk behaviour for ordinary prose, but force sentence selection
    # when the question has a distinctive technical anchor (jquery, a table or
    # procedure identifier, etc.).
    generic_excerpt_tokens = {
        "version", "actualizacion", "mejora", "documentacion", "estructura",
        "tabla", "campo", "campos", "columna", "relacion", "informacion",
        "utiliza", "utilizo", "utilizan", "utilizar",
    }
    distinctive_tokens = {
        token for token in query_tokens
        if len(token) >= 4 and token not in generic_excerpt_tokens
    }
    if (
        len(compact) <= limit
        and not re.search(r"\bPaso\s+\d+\b", compact)
        and not distinctive_tokens
    ):
        return compact
    if not query_tokens:
        return _clean_text(compact, limit)

    sentences = [
        sentence.strip()
        # Procedural manuals often flatten numbered steps into one long line
        # when extracted from Word. Treat each "Paso N" as a local evidence
        # unit so a request such as "reiniciar AppJob" does not expose the
        # preceding SMTP configuration simply because it shares a chunk.
        for sentence in re.split(r"(?<=[.!?])\s+|(?=\bPaso\s+\d+\b)", compact)
        if sentence.strip()
    ]
    if not sentences:
        return _clean_text(compact, limit)

    calculation_terms = {"calcula", "formula", "proporcional"}
    calculation_question = bool(query_tokens.intersection(calculation_terms))
    scores = []
    for sentence in sentences:
        sentence_tokens = set(tokenize(sentence))
        score = len(query_tokens.intersection(sentence_tokens))
        # Specific technical anchors must outweigh generic words such as
        # ``version``. This prevents a Crystal Reports heading from winning
        # over the later sentence that actually mentions jQuery.
        score += len(distinctive_tokens.intersection(sentence_tokens)) * 6
        if calculation_question and sentence_tokens.intersection({"formula", "ejemplo"}):
            score += 5
        scores.append(score)
    best_index = max(range(len(sentences)), key=scores.__getitem__)
    if scores[best_index] == 0:
        return _clean_text(compact, limit)

    # When a distinctive technical anchor is present, generic-only sentences
    # (for example a Crystal Reports paragraph matching only ``version``) are
    # not evidence for the requested topic. Keep the anchor sentence and, at
    # most, its immediate neighbour when it contributes additional context.
    anchor_indexes = [
        index
        for index, sentence in enumerate(sentences)
        if any(
            (
                _token_matches_query_concept(term, set(tokenize(sentence)))
                or term.lstrip("j") in set(tokenize(sentence))
            )
            for term in distinctive_tokens
        )
    ]
    if distinctive_tokens and anchor_indexes:
        focused_indexes = set(anchor_indexes)
        sentences = [sentences[index] for index in sorted(focused_indexes)]
        scores = [scores[index] for index in sorted(focused_indexes)]
        best_index = max(range(len(sentences)), key=scores.__getitem__)

    # A policy question can require two distant facts on the same page, for
    # example a benefit amount and its tax exemption. Keep the strongest
    # matching sentences rather than stopping after the first local passage.
    selected_indexes: set[int] = set()
    selected_length = 0
    for index in sorted(range(len(sentences)), key=scores.__getitem__, reverse=True):
        if scores[index] == 0:
            break
        sentence = sentences[index]
        separator_length = 1 if selected_indexes else 0
        if selected_indexes and selected_length + separator_length + len(sentence) > limit:
            continue
        selected_indexes.add(index)
        selected_length += separator_length + len(sentence)

    if not selected_indexes:
        selected_indexes.add(best_index)

    selected = " ".join(sentences[index] for index in sorted(selected_indexes))
    return _clean_text(selected, limit)


def _result_fragment(result: dict, user_message: str) -> str:
    description = _record_description(result)
    captions = result.get("@search.captions") or []
    if captions:
        caption = captions[0]
        text = caption.get("text") if isinstance(caption, dict) else getattr(caption, "text", "")
        if text:
            fragment = _excerpt_around_query(str(text), user_message)
        else:
            fragment = ""
    else:
        fragment = _excerpt_around_query(str(result.get(CONTENT_FIELD, "")), user_message)
    if description:
        prefix = f"Descripción de la solución: {description}"
        return f"{prefix}\n{fragment}" if fragment else prefix
    return fragment


def _record_description(record: dict) -> str:
    """Read the SharePoint description embedded in legacy document context."""
    direct = record.get("description") or record.get("detalle") or record.get("descripcion")
    if direct:
        return " ".join(str(direct).split())[:900]
    context = str(record.get(CONTEXT_FIELD) or "")
    match = re.search(
        r"Descripción de la solución:\s*(.*?)(?:\.\s*Dependencia:\s*|$)",
        context,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return " ".join(match.group(1).split())[:900] if match else ""


def _requested_file_names(user_message: str) -> tuple[str, ...]:
    """Return explicit, supported filenames mentioned in a user question."""
    return tuple(
        dict.fromkeys(match.group(1).casefold() for match in _EXPLICIT_FILENAME_PATTERN.finditer(user_message))
    )


def _record_matches_file_name(record: dict, file_names: tuple[str, ...]) -> bool:
    """Match a requested file against the indexed title, not a related topic."""
    title = str(record.get("title") or "").casefold()
    return any(file_name in title for file_name in file_names)


def _requested_versions(user_message: str) -> tuple[str, ...]:
    """Return exact dotted versions explicitly supplied by the user."""
    return tuple(
        dict.fromkeys(match.group(1).casefold() for match in _VERSION_PATTERN.finditer(user_message or ""))
    )


def _record_matches_requested_version(record: dict, versions: tuple[str, ...]) -> bool:
    """Avoid treating a shared version prefix as an exact version match."""
    if not versions:
        return True
    searchable_text = " ".join(
        str(record.get(field) or "")
        for field in ("title", CONTEXT_FIELD, CONTENT_FIELD)
    ).casefold()
    found_versions = {match.group(1).casefold() for match in _VERSION_PATTERN.finditer(searchable_text)}
    return any(version in found_versions for version in versions)


_STRUCTURAL_QUERY_TERMS = {
    "tabla", "tablas", "estructura", "estructuras", "campo", "campos",
    "columna", "columnas", "relacion", "relaciones", "script", "procedimiento",
}
_STRICT_VERSION_QUERY_TERMS = {
    "readme", "actualizacion", "actualizaciones", "release", "hotfix", "requisitos",
}


def _is_structural_version_query(user_message: str) -> bool:
    """Return whether a version is contextual to a technical lookup."""
    tokens = set(tokenize(user_message))
    return bool(tokens.intersection(_STRUCTURAL_QUERY_TERMS)) and not bool(
        tokens.intersection(_STRICT_VERSION_QUERY_TERMS)
    )


def _technical_anchor_tokens(user_message: str) -> set[str]:
    """Extract non-generic technical anchors used to keep direct evidence."""
    generic = GENERIC_QUERY_TOKENS | _STRUCTURAL_QUERY_TERMS | {
        "version", "versions", "evolution", "dime", "cual", "cuales",
        "utiliza", "utilizo", "usado", "usada", "indica", "necesito",
    }
    return {
        token for token in tokenize(user_message)
        if len(token) >= 3 and token not in generic and not token.isdigit()
    }


def _record_matches_technical_anchor(record: dict, user_message: str) -> bool:
    anchors = _technical_anchor_tokens(user_message)
    if not anchors:
        return False
    record_text = (
        f"{record.get('title', '')} {record.get(CONTEXT_FIELD, '')} "
        f"{record.get(CONTENT_FIELD, '')} {record.get('content_tokens', '')}"
    )
    if "ira" in set(tokenize(user_message)) and set(tokenize(user_message)).intersection(
        _STRUCTURAL_QUERY_TERMS
    ):
        compact_record = re.sub(r"[^a-z0-9]", "", record_text.casefold())
        return "irainstanciasrutasaut" in compact_record or (
            "iracodrau" in compact_record and "iracodigoentidad" in compact_record
        )
    # Compound identifiers are the strongest signal in structural questions.
    # Do not let a generic Readme page match merely because it contains words
    # such as "estructura" or "tabla"; the identifier itself must be present.
    compound_ids = re.findall(r"(?<![\w])([A-Za-z][\w]*(?:_[\w]+)+)(?![\w])", user_message or "")
    if compound_ids:
        compact_record = re.sub(r"[^a-z0-9]", "", record_text.casefold())
        for identifier in compound_ids:
            compact_identifier = re.sub(r"[^a-z0-9]", "", identifier.casefold())
            if compact_identifier in compact_record:
                return True
        return False
    record_tokens = tokenize(record_text)
    technical_context = {
        "tabla", "tablas", "campo", "campos", "columna", "columnas",
        "relacion", "relaciones", "estructura", "almacena", "almacenan",
        "procedimiento", "script", "flujo", "flujos", "ruta", "rutas",
    }
    for index, record_token in enumerate(record_tokens):
        if not any(_token_matches_query_concept(anchor, {record_token}) for anchor in anchors):
            continue
        window = set(record_tokens[max(0, index - 24) : index + 25])
        if window.intersection(technical_context):
            return True
    return False


def _question_without_versions(user_message: str) -> str:
    without_version = re.sub(
        r"\bversi[oó]n\s*(?:\d+(?:\.\d+){2,})?\b", " ", user_message or "", flags=re.IGNORECASE
    )
    return _VERSION_PATTERN.sub(" ", without_version)


def _technical_anchor_query(user_message: str) -> str:
    return " ".join(sorted(_technical_anchor_tokens(user_message)))


def _is_anchor_version_lookup(user_message: str) -> bool:
    tokens = set(tokenize(user_message))
    if not tokens.intersection({"version", "versiones"}):
        return False
    return bool(_technical_anchor_tokens(user_message))


def _record_answers_anchor_version(record: dict, user_message: str) -> bool:
    anchors = _technical_anchor_tokens(user_message)
    if not anchors:
        return True
    content = str(record.get(CONTENT_FIELD) or "")
    units = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ_]+|\d+(?:\.\d+)+|\d+", content)
    if not units:
        return False
    unit_tokens = [set(tokenize(unit)) for unit in units]
    version_like: set[int] = set()
    for index, unit in enumerate(units):
        normalized = unit.casefold()
        if re.fullmatch(r"\d+(?:\.\d+)+", normalized):
            version_like.add(index)
        elif normalized in {"sp", "version", "versiones"} and index + 1 < len(units):
            if re.fullmatch(r"\d+(?:\.\d+)*", units[index + 1]):
                version_like.update({index, index + 1})
    if not version_like:
        return False
    for index, tokens in enumerate(unit_tokens):
        if not any(
            any(
                _token_matches_query_concept(anchor, {token})
                or anchor == re.sub(r"[^a-z0-9]", "", units[index].casefold())
                or anchor.lstrip("j") == re.sub(r"[^a-z0-9]", "", units[index].casefold()).lstrip("j")
                for token in tokens
            )
            for anchor in anchors
        ):
            continue
        if any(abs(index - version_index) <= 24 for version_index in version_like):
            return True
    return False


def _source_release_version(title: str) -> tuple[int, ...] | None:
    match = _VERSION_PATTERN.search(str(title or ""))
    if not match:
        return None
    return tuple(int(part) for part in match.group(1).split("."))


def _deduplicate_equivalent_sources(
    sources: list[EvidenceSource], user_message: str
) -> list[EvidenceSource]:
    """Collapse near-identical release fragments before citation.

    Successive Readmes can repeat the same historical paragraph. For an open
    version lookup, retain the earliest release copy so the cited title does
    not point at a later document that merely repeats the change.
    """
    groups: list[list[EvidenceSource]] = []
    anchor_terms = _technical_anchor_tokens(user_message)

    def comparable_text(source: EvidenceSource) -> str:
        text = str(source.fragmento or "").casefold()
        if not anchor_terms:
            return text
        sentences = re.split(r"(?<=[.!?])\s+", text)
        focused = [
            sentence
            for sentence in sentences
            if any(_token_matches_query_concept(term, set(tokenize(sentence))) for term in anchor_terms)
        ]
        return " ".join(focused) or text

    for source in sources:
        source_tokens = set(
            re.sub(_VERSION_PATTERN, "", comparable_text(source)).split()
        )
        for group in groups:
            representative_tokens = set(
                re.sub(_VERSION_PATTERN, "", comparable_text(group[0])).split()
            )
            union = source_tokens | representative_tokens
            overlap = len(source_tokens & representative_tokens) / len(union) if union else 1
            if overlap >= (0.65 if anchor_terms else 0.90):
                group.append(source)
                break
        else:
            groups.append([source])

    selected: list[EvidenceSource] = []
    anchor_lookup = _is_anchor_version_lookup(user_message)
    for group in groups:
        if anchor_lookup:
            versioned = [item for item in group if _source_release_version(item.titulo)]
            if versioned:
                selected.append(
                    min(versioned, key=lambda item: _source_release_version(item.titulo))
                )
                continue
        selected.append(group[0])
    return selected


def _requests_readme(user_message: str) -> bool:
    return bool(re.search(r"\breadme\b", user_message or "", re.IGNORECASE))


def _readme_versions(records: Iterable[dict]) -> tuple[str, ...]:
    """Return distinct versions represented by versioned Readme titles."""
    versions: set[str] = set()
    for record in records:
        title = str(record.get("title") or "")
        if "readme" not in title.casefold():
            continue
        versions.update(
            match.group(1).casefold()
            for match in _VERSION_PATTERN.finditer(title)
        )
    return tuple(sorted(versions))


def _is_release_guidance_question(user_message: str) -> bool:
    """Recognize pre-installation release guidance without requiring its filename.

    Operators normally ask for the precaution, not for a ``Readme`` by name.
    The signal remains deliberately narrow: it needs an install/update action
    and an explicit before/preparation/precaution cue.
    """
    tokens = set(tokenize(user_message))
    return bool(
        any(token.startswith(_LEGACY_PREINSTALLATION_OPERATION_PREFIXES) for token in tokens)
        and any(token.startswith(_LEGACY_PREINSTALLATION_CUE_PREFIXES) for token in tokens)
    )


def is_release_guidance_question(user_message: str) -> bool:
    """Public policy signal for versioned installation guidance."""
    return _is_release_guidance_question(user_message)


def _has_direct_document_access_failure_coverage(record: dict) -> bool:
    """Require local troubleshooting evidence, not merely download instructions."""
    title = str(record.get("title") or "").casefold()
    content = str(record.get(CONTENT_FIELD) or "")
    # Release notes frequently mention a historical permission/download change
    # in the same line. They are not operational troubleshooting procedures
    # unless the fragment itself describes the failure and its diagnostic path.
    if "readme" in title or "changelog" in title or "actualiz" in title:
        normalized_content = " ".join(content.casefold().split())
        if not re.search(
            r"(?:no\s+(?:puede|logra)|error|falla|problema).{0,80}"
            r"(?:descarg|baj|document)",
            normalized_content,
        ):
            return False
    # Titles, document metadata and token indexes are retrieval hints, not
    # local evidence. Each prose unit must independently express the link.
    units = re.split(r"(?<=[.!?;])\s+|\n+", content)
    for unit in units:
        record_tokens = tokenize(unit)
        for start in range(len(record_tokens)):
            window = record_tokens[start : start + 28]
            has_documents = any(token.startswith("document") for token in window)
            has_permission = any(token.startswith(_DOCUMENT_ACCESS_PERMISSION_PREFIXES) for token in window)
            has_download = any(token.startswith(_DOCUMENT_ACCESS_DOWNLOAD_PREFIXES) for token in window)
            has_diagnostic = any(
                token.startswith(_DOCUMENT_ACCESS_DIAGNOSTIC_PREFIXES) for token in window
            )
            if has_documents and has_permission and has_download and has_diagnostic:
                return True
    return False


def _is_document_access_failure_question(user_message: str) -> bool:
    """Identify a specific permission-plus-download failure, not generic document management."""
    tokens = set(tokenize(user_message))
    return bool(
        any(token.startswith("document") for token in tokens)
        and any(token.startswith(_DOCUMENT_ACCESS_PERMISSION_PREFIXES) for token in tokens)
        and any(token.startswith(_DOCUMENT_ACCESS_DOWNLOAD_PREFIXES) for token in tokens)
    )


def _requests_script(user_message: str) -> bool:
    """Detect when the user is asking for an executable/script artifact."""
    return bool(SCRIPT_REQUEST_PATTERN.search(user_message or ""))


def _is_script_record(record: dict) -> bool:
    """Identify executable artifacts even when their filename is generic."""
    document_type = str(record.get("document_type") or "").casefold().lstrip(".")
    if document_type in {"sql", "ps1", "bat", "cmd", "code", "script"}:
        return True
    title = str(record.get("title") or "").casefold()
    return any(
        f"{extension} —" in title
        or title.endswith(extension)
        for extension in (".sql", ".ps1", ".bat", ".cmd")
    )


def _script_search_query(user_message: str) -> str:
    """Keep script-intent retrieval focused on the requested operation."""
    ignored = GENERIC_QUERY_TOKENS | {
        "script",
        "scripts",
        "procedimiento",
        "almacenado",
        "archivo",
        "documentacion",
        "indica",
        "indicar",
        "usa",
        "usar",
        "utiliza",
        "utilizar",
        "hay",
        "algun",
        "alguna",
        "alguno",
        "que",
        "para",
        "necesito",
    }
    terms = [
        token
        for token in tokenize(user_message)
        if len(token) >= 4 and token not in ignored
    ]
    return " ".join(dict.fromkeys(terms))


def _requested_section_pattern(user_message: str) -> re.Pattern | None:
    """Recognize a document section explicitly requested by the user."""
    normalized = " ".join((user_message or "").casefold().split())
    if re.search(r"nuevos?\s+requisitos?\s+de\s+software", normalized):
        return re.compile(r"nuevos?\s+requisitos?\s+de\s+software", re.IGNORECASE)
    return None


def _query_phrases(user_message: str) -> set[str]:
    # Question scaffolding (for example, ``qué se pueden``) does not identify
    # a manual section. Phrase scoring should describe the subject the user is
    # looking for, so headings such as "Parámetros para prórroga de contratos"
    # are not displaced by a page that merely resembles the full sentence.
    tokens = [
        token
        for token in tokenize(_question_without_background_action(user_message))
        if token not in GENERIC_QUERY_TOKENS
    ]
    return {
        " ".join(tokens[start : start + phrase_size])
        for phrase_size in (2, 3)
        for start in range(len(tokens) - phrase_size + 1)
    }


def _focused_keyword_query(user_message: str) -> str:
    """Compress question wording to its searchable domain concepts.

    This is intentionally based on stop-like query terms, never on document
    titles or individual support cases.  It complements the natural-language
    query with a heading-friendly keyword pass.
    """
    # In ``parámetros que se relacionan con X``, the relative clause describes
    # the requested parameters; it is not a request for database relations.
    # Remove that clause only after ``que se`` so a direct question such as
    # ``¿cómo se relacionan las tablas?`` keeps its technical verb.
    compact_source = _question_without_background_action(user_message)
    compact_source = re.sub(
        r"\bque\s+se\s+[a-záéíóúñ]+(?:an|en)\b",
        " ",
        user_message.casefold(),
    )
    focused_tokens = [
        token
        for token in tokenize(compact_source)
        if token not in GENERIC_QUERY_TOKENS
    ]
    # In the HR manual, configurable incapacity parameters are documented
    # under the operational section ``Riesgos de incapacidad``. Add that
    # conservative heading bridge only when both concepts are explicit in the
    # user's question; it must not broaden unrelated incapacity searches.
    focused_token_set = set(focused_tokens)
    # Administration questions are commonly answered in the management
    # section under the nouns ``tipos`` and ``documentos gestionados`` rather
    # than the operator's verb. This generic intent expansion keeps the
    # retrieval pass from selecting only the manual cover page.
    if (
        focused_token_set.intersection({"administrar", "gestionar", "gestion"})
        and any(token.startswith("document") for token in focused_token_set)
    ):
        focused_tokens.extend(
            token for token in ("gestion", "documento", "gestionado", "tipo")
            if token not in focused_token_set
        )
    # Operators often call ``ira_instancias_rutas_aut`` simply "la tabla IRA".
    # Keep that documented identifier in the lexical pass whenever the query
    # is clearly structural; this is a technical-anchor expansion, not a
    # document-specific answer rule.
    if "ira" in focused_token_set and focused_token_set.intersection(_STRUCTURAL_QUERY_TERMS):
        focused_tokens.extend(
            token
            for token in (
                "ira_instancias_rutas_aut",
                "ira_codrau",
                "ira_codigo_entidad",
            )
            if token not in focused_token_set
        )
    if {"incapacidad", "parametro"}.issubset(focused_token_set):
        focused_tokens.extend(
            token
            for token in (
                "riesgo",
                "incapacidad",
                "incapacidadesvalidatraslapeconacciones",
                "subsidio",
                "descuento",
            )
            if token not in focused_token_set
        )
    # La página de prórroga no repite el nombre de Evolution en cada campo:
    # utiliza los identificadores CamelCase de los parámetros. Anclamos la
    # pasada léxica solo cuando el operador pide explícitamente parámetros de
    # prórroga de contratos; así no se amplían búsquedas de contratos ajenas.
    if {"prorroga", "contrato", "parametro"}.issubset(focused_token_set):
        focused_tokens.extend(
            token
            for token in (
                "prorrogacontratodiasatrasiniciorangofechafincontrato",
                "prorrogacontratodiasdespuesfinalrangofechafincontrato",
            )
            if token not in focused_token_set
        )
    # Upgrade manuals usually place pre-installation precautions under
    # ``Preparación`` and ``Respaldo`` rather than under the operator's word
    # ``precauciones``. Keep this bridge limited to an explicit pre-install
    # update request so ordinary change-log lookups are not broadened.
    has_preinstallation_operation = any(
        token.startswith(("instal", "actualiz")) for token in focused_token_set
    )
    has_preinstallation_cue = any(
        token.startswith(("precaucion", "antes", "prev", "respald", "prepar"))
        for token in focused_token_set
    )
    if has_preinstallation_operation and has_preinstallation_cue:
        focused_tokens.extend(
            token
            for token in ("preparacion", "respaldo", "configuracion")
            if token not in focused_token_set
        )
    # The DTC validation manual places the actionable checks on pages after
    # its generic "Validación" heading. Expand only an explicit DTC/MSDTC
    # validation request so the focused pass retrieves those existing pages.
    if _is_dtc_validation_question(user_message):
        focused_tokens.extend(
            token
            for token in ("firewall", "component", "inboud", "outboud", "regla")
            if token not in focused_token_set
        )
    return " ".join(focused_tokens)


def _is_dtc_validation_question(user_message: str) -> bool:
    """Detect explicit DTC/MSDTC validation questions without broadening search."""
    tokens = set(tokenize(user_message))
    return bool(
        tokens.intersection({"dtc", "msdtc"})
        and any(token.startswith(("valid", "verif", "revis", "confirm")) for token in tokens)
    )


def _query_synonym_tokens(tokens: Iterable[str]) -> set[str]:
    """Return query terms plus conservative, generic vocabulary bridges."""
    expanded = set(tokens)
    for group in QUERY_SYNONYM_GROUPS:
        if expanded.intersection(group):
            expanded.update(group)
    # In this operational vocabulary, "contrato próximo" normally refers to
    # the prórroga/renewal window documented by the HR module. Keep this bridge
    # conditional so ``próximo`` remains an ordinary temporal word elsewhere.
    if "contrato" in expanded and expanded.intersection({"proxima", "proximo", "proxim"}):
        expanded.update({"prorroga", "prorrog"})
    return expanded


def _token_matches_query_concept(token: str, record_tokens: set[str]) -> bool:
    """Match a token literally, by morphology, or through a query synonym."""
    if token in record_tokens:
        return True
    for group in QUERY_SYNONYM_GROUPS:
        if token in group and record_tokens.intersection(group):
            return True
    return any(
        len(token) >= 5
        and len(candidate) >= 5
        and (token.startswith(candidate[:5]) or candidate.startswith(token[:5]))
        for candidate in record_tokens
    )


def _alias_augmented_source_text(source_text: str) -> str:
    """Make indexed synonym hits visible to the existing action checker."""
    tokens = set(tokenize(source_text))
    aliases = set()
    for group in QUERY_SYNONYM_GROUPS:
        if tokens.intersection(group):
            aliases.update(group)
    return f"{source_text} {' '.join(sorted(aliases))}"


def _document_relevance_score(
    record: dict,
    user_message: str,
    token_weights: dict[str, float] | None = None,
    phrase_weights: dict[str, float] | None = None,
) -> float:
    """Score evidence with generic token coverage, phrase matches and structure."""
    query_tokens = sorted(_query_synonym_tokens(tokenize(user_message)))
    document_tokens = tokenize(
        f"{record.get('title', '')} {record.get(CONTEXT_FIELD, '')} "
        f"{record.get('content_tokens', '')} {record.get(CONTENT_FIELD, '')} "
        f"{record.get('document_type', '')}"
    )
    if not query_tokens or not document_tokens:
        return 0.0

    document_token_set = set(document_tokens)
    token_overlap = set(query_tokens).intersection(document_token_set)
    coverage_score = sum((token_weights or {}).get(token, 1) for token in token_overlap)
    synonym_action_score = sum(
        140
        for group in QUERY_SYNONYM_GROUPS
        if set(tokenize(user_message)).intersection(group)
        and document_token_set.intersection(group)
    )
    # A concept present in the document title is a stronger document-level
    # signal than the same word buried in an arbitrary page fragment. This is
    # generic metadata weighting: it improves lookup of named procedures,
    # manuals and modules without maintaining question-specific aliases.
    title_tokens = set(tokenize(record.get("title", "")))
    title_overlap = set(query_tokens).intersection(title_tokens)
    title_score = len(title_overlap) * 10
    technical_structure_score = 0
    if (
        set(tokenize(user_message)).intersection(_STRUCTURAL_QUERY_TERMS)
        and _record_matches_technical_anchor(record, user_message)
    ):
        # A named table/technical identifier is stronger than incidental
        # version or module words in a changelog page.
        technical_structure_score = 3_000
        if _is_structural_version_query(user_message):
            technical_structure_score += 500 if not _is_script_record(record) else -500
    # Support sites commonly store a solution below a descriptive folder but
    # give the actual file a generic name (``Indicaciones.txt``, ``Custom``).
    # Treat a strong parent-folder match as document-level metadata, much like
    # a descriptive filename. Requiring three concepts prevents a broad folder
    # such as ``SOLUCIONES`` from outranking direct evidence on its own.
    parent_context = _sharepoint_parent_context(
        str(record.get("source_url") or ""), str(record.get("folder_path") or "")
    )
    parent_overlap = set(query_tokens).intersection(set(tokenize(parent_context)))
    parent_context_score = (
        300 + ((len(parent_overlap) - 3) * 100)
        if len(parent_overlap) >= 3
        else 0
    )
    document_text = " ".join(document_tokens)
    phrase_matches = {
        phrase for phrase in _query_phrases(user_message) if phrase in document_text
    }
    phrase_score = sum(
        (phrase_weights or {}).get(phrase, 4) for phrase in phrase_matches
    )
    # A phrase at the beginning of the page normally represents its heading.
    # That is stronger evidence than the same phrase appearing later in an
    # introductory table or a broad module overview.  Use only the fragment,
    # not the document-wide context, so a distant section cannot borrow this
    # boost.
    fragment_opening = " ".join(tokenize(record.get(CONTENT_FIELD, ""))[:18])
    heading_phrase_count = sum(
        1 for phrase in phrase_matches if phrase in fragment_opening
    )
    # Keep nouns such as ``parámetro`` but remove verbs such as ``configurar``
    # when looking for a section heading. A user normally asks in a sentence;
    # manuals normally name the section with its subject concepts.
    heading_query_tokens = [
        token
        for token in query_tokens
        if token not in GENERIC_QUERY_TOKENS
        and token not in (OPERATIONAL_QUERY_TOKENS - {"parametro"})
    ]
    heading_subject_phrases = {
        " ".join(heading_query_tokens[start : start + phrase_size])
        for phrase_size in (2, 3)
        for start in range(len(heading_query_tokens) - phrase_size + 1)
    }
    opening_text = " ".join(tokenize(record.get(CONTENT_FIELD, "")[:220]))
    heading_subject_phrase_count = sum(
        1 for phrase in heading_subject_phrases if phrase in opening_text
    )

    semantic_score = float(record.get("@search.reranker_score") or 0)
    azure_score = float(record.get("@search.score") or 0)
    # When the user explicitly asks for a script, an executable artifact is
    # the answer-bearing object. README pages may mention the same business
    # terms, but must not outrank the SQL/PowerShell file whose description
    # explains what it actually does.
    script_score = 0
    if _requests_script(user_message):
        script_score = 420 if _is_script_record(record) else -260
    # A pre-installation safety question is not answered by a page that merely
    # mentions an update or a vulnerability. Prefer a local checklist that
    # documents preparation, backup or recommendations. This is a generic
    # evidence shape and does not name a product, version or document.
    query_token_set = set(tokenize(_question_without_background_action(user_message)))
    document_token_set = set(document_tokens)
    is_preinstallation_question = (
        any(
            token.startswith(_LEGACY_PREINSTALLATION_OPERATION_PREFIXES)
            for token in query_token_set
        )
        and any(
            token.startswith(_LEGACY_PREINSTALLATION_CUE_PREFIXES)
            for token in query_token_set
        )
    )
    preinstallation_score = (
        len(document_token_set.intersection(_LEGACY_PREINSTALLATION_EVIDENCE_TOKENS)) * 45
        if is_preinstallation_question
        else 0
    )
    # Release notes are the authoritative artifact for precautions before an
    # installation/update. This is a ranking signal, not a provenance or
    # version exception: an authorized direct Upgrade guide can still answer
    # when no Readme/release candidate exists.
    release_guidance_score = 0
    if _is_release_guidance_question(user_message):
        title = str(record.get("title") or "").casefold()
        if (
            any(marker in title for marker in _RELEASE_GUIDANCE_TITLE_MARKERS)
            and document_token_set.intersection(_LEGACY_PREINSTALLATION_EVIDENCE_TOKENS)
        ):
            release_guidance_score = 360 + (90 * sum(
                (
                    any(token.startswith(prefix) for token in document_token_set for prefix in ("prepar", "previo", "antes")),
                    any(token.startswith(prefix) for token in document_token_set for prefix in ("respal", "backup")),
                    any(token.startswith(prefix) for token in document_token_set for prefix in ("instal", "actualiz", "aplic")),
                    any(token.startswith(prefix) for token in document_token_set for prefix in ("recomend", "advertenc", "precauc")),
                )
            ))
    is_dtc_validation_question = _is_dtc_validation_question(user_message)
    # A section heading may repeat all question terms while containing no
    # actionable check. Give equal preference to pages with a concrete DTC
    # validation control, so complementary firewall and Component Services
    # checks can both survive the final relevance threshold.
    dtc_validation_score = (
        250
        if is_dtc_validation_question
        and document_token_set.intersection(_DTC_VALIDATION_EVIDENCE_TOKENS)
        else 0
    )
    # Distinguish document-management guidance from a Portal download path.
    # Both may mention managed documents, but an operator asking how to
    # administer them needs the management manual, not the consultation flow.
    management_score = 0
    query_set = set(tokenize(user_message))
    if "document" in query_set and query_set.intersection({"administr", "gestionar", "gestion"}):
        title_text = " ".join(tokenize(record.get("title", "")))
        content_text = " ".join(tokenize(record.get(CONTENT_FIELD, "")))
        if "gestion" in title_text and "document" in title_text:
            management_score += 650
            if any(marker in content_text for marker in ("haga clic", "seleccione", "crear", "editar", "nuevo", "documento gestionado")):
                management_score += 420
        elif "portal" in title_text and "descarg" in content_text:
            management_score -= 180
    # Coverage across the question's concepts matters more than one isolated
    # exact phrase. This prevents a page that merely lists a decree number
    # from outranking the page that explains its calculation.
    return (
        (coverage_score * 6)
        + synonym_action_score
        + title_score
        + technical_structure_score
        + parent_context_score
        # A phrase that follows the user's wording is stronger evidence than
        # isolated token overlap. This is especially important for manuals
        # whose section headings describe an operation directly.
        + (phrase_score * 5)
        + (heading_phrase_count * 90)
        + (heading_subject_phrase_count * 140)
        # Vector similarity finds paraphrases; lexical rank favors pages that
        # explicitly contain the terms requested. Neither source wins alone.
        + max(0, MAX_CANDIDATES - int(record.get("_vector_rank", MAX_CANDIDATES))) * 3.0
        + max(0, MAX_CANDIDATES - int(record.get("_keyword_rank", MAX_CANDIDATES))) * 0.2
        # A compact lexical query favors a section heading made of the user's
        # meaningful concepts over filler words from the natural question.
        + max(0, MAX_CANDIDATES - int(record.get("_focused_keyword_rank", MAX_CANDIDATES))) * 1.5
        # This bounded pass exists only for pre-installation release guidance.
        + max(0, MAX_CANDIDATES - int(record.get("_release_readme_rank", MAX_CANDIDATES))) * 1.5
        # Prefix lookup is reserved for concrete error reports whose subject
        # can be embedded in a CamelCase path. Its result is still checked by
        # normal coverage/provenance rules, but deserves ranking credit once
        # those checks pass.
        + max(0, MAX_CANDIDATES - int(record.get("_prefix_rank", MAX_CANDIDATES))) * 2.0
        # Semantic reranking has already compared the full question with the
        # candidate passage. Give it material weight when enabled, rather than
        # letting a one-token lexical advantage suppress a direct paraphrase.
        + (semantic_score * 25)
        + (azure_score / 1_000)
        - (int(record.get("_missing_anchor_count", 0)) * 8)
        + script_score
        + preinstallation_score
        + release_guidance_score
        + dtc_validation_score
        + management_score
    )


def _requested_country(user_message: str) -> str | None:
    query_tokens = set(tokenize(user_message))
    for country, country_tokens in COUNTRY_TOKENS.items():
        if query_tokens.intersection(country_tokens):
            return country
    return None


_LEGACY_SHAREPOINT_URL_MARKERS = {
    "readme hotfixes": "/readme hotfixes/",
    "documentos/soluciones": "/documentos compartidos/soluciones/",
    "manuales": "/manuales/",
    "scripts de apoyo": "/scripts de apoyo/",
}


def _legacy_url_has_authorized_source(
    source_url: str, allowed_source_labels: tuple = ()
) -> bool:
    """Fail closed for legacy records that have no library metadata.

    Early production records were indexed before ``drive_id`` and
    ``folder_path`` were stored.  Their SharePoint URLs still expose a stable
    library path for the four technical-pilot sources.  URLs such as
    ``_layouts/15/Doc.aspx`` omit that path and therefore must not be trusted.
    """
    normalized_url = unquote(str(source_url or "")).casefold()
    allowed_markers = {
        _LEGACY_SHAREPOINT_URL_MARKERS[label.casefold()]
        for label in allowed_source_labels
        if label.casefold() in _LEGACY_SHAREPOINT_URL_MARKERS
    }
    return any(marker in normalized_url for marker in allowed_markers)


def _record_has_authorized_provenance(
    record: dict, allowed_sources: tuple = (), allowed_source_labels: tuple = ()
) -> bool:
    """Accept only HTTPS SharePoint items from configured libraries/folders.

    ``allowed_sources`` is normally a tuple of ``(folder_path, drive_id)``
    pairs from ``Config.sharepoint_sources``. A tuple of folder strings is
    still accepted for compatibility with older callers and tests. Legacy
    records without either metadata field are accepted only when their URL
    proves they belong to an explicitly labelled pilot library.
    """
    has_sharepoint_provenance = (
        record.get("source_system") == "sharepoint"
        and str(record.get("source_url") or "").startswith("https://")
    )
    if not has_sharepoint_provenance:
        return False
    if not allowed_sources:
        return True
    record_folder_path = str(record.get("folder_path") or "").strip("/")
    record_drive_id = str(record.get("drive_id") or "").strip()
    if all(isinstance(source, (tuple, list)) and len(source) == 2 for source in allowed_sources):
        if not record_folder_path and not record_drive_id:
            return _legacy_url_has_authorized_source(
                record.get("source_url") or "", allowed_source_labels
            )
        return any(
            record_folder_path == str(folder_path or "").strip("/")
            and record_drive_id == str(drive_id or "").strip()
            for folder_path, drive_id in allowed_sources
        )
    # Legacy single-library configurations only had folder_path available.
    return record_folder_path in {
        str(path or "").strip("/") for path in allowed_sources if path is not None
    }


def _record_contains_document_injection(record: dict) -> bool:
    """Detect only strong instruction-like markers in indexed content.

    This is a defense-in-depth gate, not a relevance classifier. It fails
    closed for explicit instruction overrides while avoiding generic words
    such as ``prompt`` or ``system`` that may occur in ordinary manuals.
    """
    searchable_text = " ".join(
        str(record.get(field) or "")
        for field in ("title", CONTEXT_FIELD, CONTENT_FIELD)
    )
    return bool(_DOCUMENT_INJECTION_PATTERN.search(searchable_text))


def _filter_records_for_requested_country(
    records: list[dict], user_message: str
) -> list[dict]:
    """Do not substitute evidence from another country for an explicit request."""
    country = _requested_country(user_message)
    if country is None:
        return records

    country_tokens = COUNTRY_TOKENS[country]
    return [
        record
        for record in records
        if _record_matches_only_requested_country(record, country_tokens)
    ]


def _record_matches_only_requested_country(record: dict, country_tokens: set[str]) -> bool:
    record_tokens = set(
        tokenize(
            f"{record.get('title', '')} {record.get(CONTEXT_FIELD, '')} "
            f"{record.get(CONTENT_FIELD, '')}"
        )
    )
    if not country_tokens.intersection(record_tokens):
        return False

    other_country_tokens = set().union(
        *(tokens for tokens in COUNTRY_TOKENS.values() if tokens != country_tokens)
    )
    # Contact footers often list several countries.  They do not establish
    # that the technical procedure applies to any one of those countries.
    return not other_country_tokens.intersection(record_tokens)


def _has_minimum_content_coverage(
    record: dict, user_message: str, semantic_enabled: bool = False
) -> bool:
    """Reject a tangential hit that shares only one broad term with the query."""
    country_tokens = set().union(*COUNTRY_TOKENS.values())
    coverage_message = _question_without_background_action(user_message)
    # Synonym expansion is for retrieval and concept matching, not extra
    # evidence requirements. Requiring every expanded alias would inflate the
    # coverage threshold and let one broad document pass by accident.
    required_tokens = set(tokenize(coverage_message)).difference(
        GENERIC_QUERY_TOKENS, country_tokens
    )
    if not required_tokens:
        return True

    record_tokens = set(
        tokenize(
            f"{record.get('title', '')} {record.get('content_tokens', '')} "
            f"{record.get(CONTEXT_FIELD, '')} "
            f"{record.get(CONTENT_FIELD, '')} "
            f"{record.get('document_type', '')}"
        )
    )
    # A cover/table of contents can contain the requested version, subject and
    # action while only pointing to the actual answer elsewhere in the file.
    # Keep navigation useful for an explicit request for the index itself, but
    # never pass it as evidence for a substantive technical question.
    fragment = str(record.get(CONTENT_FIELD, ""))
    normalized_fragment = " ".join(fragment.lower().split())
    page_references = re.findall(r",\s*\d{1,3}(?=\s|$)", normalized_fragment)
    is_navigation_fragment = (
        "tabla de contenido" in normalized_fragment
        or (
            len(page_references) >= 6
            and any(
                marker in normalized_fragment
                for marker in ("pagina", "página", "indice", "contenido", "seccion", "capitulo")
            )
        )
    )
    query_navigation = set(tokenize(coverage_message)).intersection(NAVIGATION_QUERY_TOKENS)
    if is_navigation_fragment and not query_navigation:
        return False
    # Accept ordinary gender/number variants (``negativa``/``negativo``)
    # without maintaining aliases for individual questions or documents.
    required_matches = {
        token for token in required_tokens if _token_matches_query_concept(token, record_tokens)
    }
    operational_tokens = OPERATIONAL_QUERY_TOKENS | {
        "bajar", "descargar", "descarga", "ofuscar", "ofuscacion", "ofuscan",
        "gestionar", "administrar", "prorrogar",
        "controla", "define", "indica", "muestra", "determina",
    }
    subject_tokens = required_tokens.difference(operational_tokens)
    subject_matches = subject_tokens.intersection(required_matches)
    if _requested_country(user_message) is not None:
        # A country name can appear in boilerplate. Require two independent
        # concepts for a country-scoped query so a generic Readme cannot be
        # presented as legal/payroll evidence because of one incidental term.
        minimum_matches = (
            2
            if len(required_tokens) >= 3
            and set(tokenize(user_message)).intersection(COUNTRY_SENSITIVE_ANCHORS)
            else 1
        )
    elif len(required_tokens) >= 4:
        minimum_matches = 3
    else:
        minimum_matches = 2 if len(required_tokens) >= 3 else 1
    action_question = " ".join(
        token
        for token in tokenize(coverage_message)
        if token not in DIAGNOSTIC_ACTION_TOKENS
    )
    action_is_covered = has_requested_action_coverage(
        action_question,
        _alias_augmented_source_text(
            f"{record.get('title', '')} {record.get(CONTENT_FIELD, '')}"
        ),
    )
    management_question = (
        set(tokenize(user_message)).intersection({"administrar", "gestionar", "gestion"})
        and any(token.startswith("document") for token in tokenize(user_message))
    )
    if management_question:
        management_action_tokens = set(tokenize(fragment)).intersection(
            {"administrar", "gestionar", "crear", "editar", "eliminar", "subir", "nuevo", "seleccionar", "tipo"}
        )
        action_is_covered = action_is_covered or bool(
            "documento" in record_tokens and management_action_tokens
        )
    if _is_release_guidance_question(user_message) and _is_release_guidance_record(record):
        # "Tomar precauciones" is commonly documented as preparation,
        # backups and update instructions rather than with the literal verb.
        action_is_covered = action_is_covered or _has_release_guidance_coverage(record)
    # ``qué se debe revisar`` is troubleshooting scaffolding, not a demand
    # that the cited manual literally contain the verb ``revisar``. Keep the
    # stronger subject/coverage checks; only relax this generic diagnostic verb.
    diagnostic_only = set(tokenize(user_message)).intersection(DIAGNOSTIC_ACTION_TOKENS)
    if diagnostic_only and not action_question:
        action_is_covered = bool(subject_matches)
    # A semantic reranker can validate ordinary paraphrases such as
    # ``administrar documentos`` and ``gestión de documentos``. It is only an
    # alternative to literal action matching when semantic search is explicitly
    # enabled and its score is strong; lexical concept coverage remains
    # mandatory.
    semantic_action_is_covered = (
        semantic_enabled
        and float(record.get("@search.reranker_score") or 0) >= SEMANTIC_ACTION_MIN_SCORE
    )
    if (
        semantic_enabled
        and _requested_country(user_message) is not None
        and float(record.get("@search.reranker_score") or 0) < SEMANTIC_ACTION_MIN_SCORE
    ):
        # Country names often occur in boilerplate or contact sections. A weak
        # semantic match must not turn that incidental mention into evidence
        # for a legal or payroll question.
        return False
    # An operation (for example, ``modificar``) is not evidence unless the
    # subject requested by the user is present in the same chunk. This avoids
    # returning a generic infrastructure page for a module-specific question.
    if subject_tokens and not subject_matches:
        return False
    if subject_matches and required_matches.intersection(OPERATIONAL_QUERY_TOKENS):
        direct_tokens = tokenize(
            f"{record.get('title', '')} {record.get('content_tokens', '')} "
            f"{record.get(CONTENT_FIELD, '')}"
        )
        window_size = 48
        operational_matches = required_matches.intersection(operational_tokens)
        def window_covers(token: str, window: list[str]) -> bool:
            return _token_matches_query_concept(token, set(window))
        compact_match = any(
            any(window_covers(token, direct_tokens[start : start + window_size]) for token in subject_matches)
            and all(
                window_covers(token, direct_tokens[start : start + window_size])
                for token in operational_matches
            )
            for start in range(len(direct_tokens))
        )
        if not compact_match:
            return False
        # A table of contents or alphabetical index can contain every query
        # term, but it only points to an answer elsewhere in the document.  Do
        # not present navigation pages as evidence for an operational question.
        # This is deliberately structural (rather than a list of document
        # names), so it applies to new manuals as they are indexed.
    return len(required_matches) >= minimum_matches and (
        action_is_covered or semantic_action_is_covered
    )


def _diversify_candidate_records(records: list[dict]) -> list[dict]:
    """Keep a broad candidate pool without letting one document monopolize it."""
    def candidate_rank(record: dict) -> int:
        ranks = [
            int(record[field])
            for field in (
                "_vector_rank", "_keyword_rank", "_focused_keyword_rank",
                "_release_readme_rank", "_prefix_rank",
            )
            if record.get(field) is not None
        ]
        return min(ranks) if ranks else CANDIDATE_POOL_SIZE

    per_document_count: dict[str, int] = {}
    diversified: list[dict] = []
    for record in sorted(records, key=candidate_rank):
        document_key = str(record.get("document_id") or record.get("title") or record.get("id"))
        count = per_document_count.get(document_key, 0)
        if count >= MAX_CANDIDATES_PER_DOCUMENT:
            continue
        per_document_count[document_key] = count + 1
        diversified.append(record)
    return diversified


def _is_release_guidance_record(record: dict) -> bool:
    """Identify a release/readme page that contains preparation guidance."""
    title = str(record.get("title") or "").casefold()
    if not any(marker in title for marker in _RELEASE_GUIDANCE_TITLE_MARKERS):
        return False
    document_tokens = set(
        tokenize(
            " ".join(
                str(record.get(field) or "")
                for field in ("title", CONTEXT_FIELD, CONTENT_FIELD, "content_tokens")
            )
        )
    )
    return bool(document_tokens.intersection(_LEGACY_PREINSTALLATION_EVIDENCE_TOKENS))


def _has_release_guidance_coverage(record: dict) -> bool:
    """Recognize preparation guidance expressed without the query's exact verb."""
    tokens = set(tokenize(str(record.get(CONTENT_FIELD) or "")))
    families = (
        any(token.startswith(prefix) for token in tokens for prefix in ("prepar", "previo", "antes")),
        any(token.startswith(prefix) for token in tokens for prefix in ("respal", "backup")),
        any(token.startswith(prefix) for token in tokens for prefix in ("instal", "actualiz", "aplic")),
        any(token.startswith(prefix) for token in tokens for prefix in ("recomend", "advertenc", "precauc")),
    )
    return sum(families) >= 3


def _limit_candidate_pool(
    records: list[dict], limit: int, prioritize_release: bool = False
) -> list[dict]:
    """Apply a global bound while preserving explicit and high-ranked hits."""
    if len(records) <= limit:
        return records

    def rank(record: dict) -> int:
        ranks = [
            int(record[field])
            for field in (
                "_vector_rank", "_keyword_rank", "_focused_keyword_rank",
                "_release_readme_rank", "_prefix_rank",
            )
            if record.get(field) is not None
        ]
        return min(ranks) if ranks else CANDIDATE_POOL_SIZE

    def sort_key(record: dict):
        # For pre-installation guidance, preserve release/readme pages with
        # local preparation evidence before generic manuals. This keeps a
        # direct release page in the bounded pool even when Azure's broad
        # lexical pass returns many unrelated historical pages.
        release_priority = (
            0 if prioritize_release and _is_release_guidance_record(record) else 1
        )
        release_rank = int(record.get("_release_readme_rank", CANDIDATE_POOL_SIZE))
        return (
            not bool(record.get("_explicit_filename_match")),
            release_priority,
            release_rank if prioritize_release else CANDIDATE_POOL_SIZE,
            rank(record),
        )

    return sorted(records, key=sort_key)[:limit]


def _candidate_diversity_stats(records: list[dict]) -> dict[str, int]:
    """Return non-sensitive diversity metrics for a bounded candidate pool."""
    counts: dict[str, int] = {}
    for record in records:
        key = str(record.get("document_id") or record.get("title") or record.get("id") or "")
        counts[key] = counts.get(key, 0) + 1
    return {
        "unique_documents": len(counts),
        "max_fragments_per_document": max(counts.values(), default=0),
    }


def _merge_candidate_record(existing: dict, incoming: dict) -> dict:
    """Merge Azure payloads without losing a rank assigned by an earlier pass."""
    preserved_ranks = {
        field: existing[field]
        for field in _CANDIDATE_RANK_FIELDS
        if existing.get(field) is not None
    }
    existing.update(incoming)
    existing.update(preserved_ranks)
    return existing


def _rerank_records(records: list[dict], user_message: str) -> list[tuple[float, dict]]:
    """Rerank Azure candidates with query coverage and corpus-relative weights."""
    query_tokens = set(tokenize(user_message))
    token_document_frequency = {token: 0 for token in query_tokens}
    phrases = _query_phrases(user_message)
    phrase_document_frequency = {phrase: 0 for phrase in phrases}
    for record in records:
        document_text = " ".join(
            tokenize(f"{record.get('title', '')} {record.get(CONTENT_FIELD, '')}")
        )
        document_token_set = set(document_text.split())
        for token in query_tokens:
            if token in document_token_set:
                token_document_frequency[token] += 1
        for phrase in phrases:
            if phrase in document_text:
                phrase_document_frequency[phrase] += 1

    candidate_count = max(len(records), 1)
    token_weights = {
        token: 1 + math.log((candidate_count + 1) / (frequency + 1))
        for token, frequency in token_document_frequency.items()
    }
    phrase_weights = {
        phrase: 2 + (4 * math.log((candidate_count + 1) / (frequency + 1)))
        for phrase, frequency in phrase_document_frequency.items()
        if frequency
    }
    # Query terms that appear in only part of the candidate set are useful
    # disambiguators (country, product, module, acronym). Penalize a result
    # that omits them, without keeping a vocabulary of special cases.
    anchor_tokens = {
        token
        for token, frequency in token_document_frequency.items()
        if 0 < frequency <= candidate_count * 0.45 and len(token) > 3
    }
    for record in records:
        record_tokens = set(
            tokenize(
                f"{record.get('title', '')} {record.get(CONTEXT_FIELD, '')} "
                f"{record.get(CONTENT_FIELD, '')}"
            )
        )
        record["_missing_anchor_count"] = len(anchor_tokens.difference(record_tokens))
    ranked_records = [
        (_document_relevance_score(record, user_message, token_weights, phrase_weights), record)
        for record in records
    ]
    ranked_records.sort(key=lambda item: item[0], reverse=True)
    return ranked_records


def _retrieve_legacy_azure_search_evidence(
    user_message: str, config, client=None, limit: int = 3, diagnostics: dict | None = None
) -> list[EvidenceSource]:
    """Retrieve vector candidates and normalize them to bot evidence."""
    if not getattr(config, "azure_search_configured", False):
        return []

    search_client = SearchClient(
        endpoint=config.azure_search_endpoint,
        index_name=config.azure_search_index_name,
        credential=_credential(config),
    )
    search_args = {
        "top": CANDIDATE_POOL_SIZE,
        "select": SEARCH_SELECT_FIELDS,
        "connection_timeout": SEARCH_TIMEOUT_SECONDS,
        "read_timeout": SEARCH_TIMEOUT_SECONDS,
    }
    requested_file_names = _requested_file_names(user_message)
    requested_versions = _requested_versions(user_message)
    requested_section = _requested_section_pattern(user_message)
    keyword_search_args = dict(search_args)
    if getattr(config, "azure_search_use_semantic", False):
        keyword_search_args.update(
            {
                "query_type": "semantic",
                "semantic_configuration_name": config.azure_search_semantic_configuration,
                "query_caption": "extractive",
            }
        )

    # Prefer the bounded lexical passes when they already establish direct
    # evidence. Embedding generation plus vector search can consume most of
    # the handler's retrieval budget even for questions whose wording is
    # represented verbatim (or by the query-side vocabulary aliases) in the
    # index. Keep vector retrieval as a semantic fallback for paraphrases.
    pass_counts: dict[str, int] = {}
    try:
        records_by_id: dict[str, dict] = {}
        keyword_count = 0
        for rank, result in enumerate(
            search_client.search(
                search_text=user_message,
                search_fields=["title", CONTENT_FIELD, "content_tokens"],
                **keyword_search_args,
            ),
            start=1,
        ):
            keyword_count += 1
            record = dict(result)
            record_id = str(record.get("id", ""))
            existing = records_by_id.get(record_id)
            if existing is None:
                existing = record
                records_by_id[record_id] = existing
            else:
                _merge_candidate_record(existing, record)
            existing["_keyword_rank"] = rank
        pass_counts["keyword"] = keyword_count
        focused_query = _focused_keyword_query(user_message)
        focused_query = " ".join(
            sorted(_query_synonym_tokens(tokenize(focused_query)))
        )
        if focused_query:
            focused_count = 0
            for rank, result in enumerate(
                search_client.search(
                    search_text=focused_query,
                    search_fields=["title", CONTENT_FIELD, "content_tokens"],
                    **keyword_search_args,
                ),
                start=1,
                ):
                focused_count += 1
                record = dict(result)
                record_id = str(record.get("id", ""))
                existing = records_by_id.get(record_id)
                if existing is None:
                    existing = record
                    records_by_id[record_id] = existing
                else:
                    _merge_candidate_record(existing, record)
                existing["_focused_keyword_rank"] = rank
            pass_counts["focused_keyword"] = focused_count

        # Database operators frequently shorten a named table to its acronym.
        # Search the normalized identifier tokens with an AND pass so the
        # technical manual is discoverable even when Azure tokenizes the
        # underscore-separated name differently from the natural-language
        # question.
        focused_terms = tokenize(focused_query)
        if "ira" in focused_terms and set(focused_terms).intersection(_STRUCTURAL_QUERY_TERMS):
            identifier_query = "ira instancias rutas aut"
            identifier_count = 0
            for rank, result in enumerate(
                search_client.search(
                    search_text=identifier_query,
                    search_fields=["title", CONTENT_FIELD, "content_tokens"],
                    **search_args,
                ),
                start=1,
            ):
                identifier_count += 1
                record = dict(result)
                record_id = str(record.get("id", ""))
                existing = records_by_id.get(record_id)
                if existing is None:
                    existing = record
                    records_by_id[record_id] = existing
                else:
                    _merge_candidate_record(existing, record)
                existing["_technical_anchor_rank"] = rank
            pass_counts["technical_anchor"] = identifier_count

        # A pre-installation question often omits the word "Readme", although
        # that release artifact is where its precautions are documented. Add a
        # bounded lexical pass so a generic Upgrade manual cannot hide the
        # Readme before the 60-candidate union and deterministic reranking.
        if _is_release_guidance_question(user_message):
            release_query = " ".join(
                dict.fromkeys(("readme", *tokenize(focused_query)))
            )
            release_count = 0
            for rank, result in enumerate(
                search_client.search(
                    search_text=release_query,
                    search_fields=["title", CONTENT_FIELD, "content_tokens"],
                    **keyword_search_args,
                ),
                start=1,
            ):
                release_count += 1
                record = dict(result)
                record_id = str(record.get("id", ""))
                existing = records_by_id.get(record_id)
                if existing is None:
                    existing = record
                    records_by_id[record_id] = existing
                else:
                    _merge_candidate_record(existing, record)
                existing["_release_readme_rank"] = rank
            pass_counts["release_readme"] = release_count

        # Older technical documents often use a CamelCase application path
        # (for example ``TiempoNoTrabajado``) while operators describe it as
        # separate words. For a concrete error report, add one bounded prefix
        # pass over its two trailing subject terms. This is not a broad
        # fallback: it only runs when the message has enough diagnostic detail
        # and lets the trusted SharePoint parent-folder context break ties.
        focused_terms = tokenize(focused_query)
        if "error" in focused_terms and len(focused_terms) >= 3:
            subject_terms = [term for term in focused_terms if term != "error"][-2:]
            if len(subject_terms) == 2 and all(len(term) >= 4 for term in subject_terms):
                prefix_query = " AND ".join(f"{term}*" for term in subject_terms)
                try:
                    prefix_count = 0
                    for rank, result in enumerate(
                        search_client.search(
                            search_text=prefix_query,
                            search_fields=["title", CONTENT_FIELD, "content_tokens"],
                            query_type="full",
                            search_mode="all",
                            **search_args,
                        ),
                        start=1,
                    ):
                        prefix_count += 1
                        record = dict(result)
                        record_id = str(record.get("id", ""))
                        existing = records_by_id.get(record_id)
                        if existing is None:
                            existing = record
                            records_by_id[record_id] = existing
                        else:
                            _merge_candidate_record(existing, record)
                        existing["_prefix_rank"] = rank
                    pass_counts["prefix"] = prefix_count
                except Exception:
                    # The ordinary lexical pass remains valid when a legacy
                    # service does not accept full-query prefix syntax.
                    pass

        # Do not pay for an embedding call when the lexical candidates already
        # include a chunk that passes the same deterministic evidence gate used
        # below. This is especially important in Teams, where a 12-second
        # retrieval timeout otherwise turns an answerable question into a
        # misleading "sin evidencia" response.
        keyword_records = list(records_by_id.values())
        has_lexical_evidence = any(
            _has_minimum_content_coverage(
                record,
                user_message,
                semantic_enabled=getattr(config, "azure_search_use_semantic", False),
            )
            for record in keyword_records
        )
        if not has_lexical_evidence:
            try:
                query_embedding = _embed_texts([user_message], config, client=client)[0]
                vector_query = VectorizedQuery(
                    vector=query_embedding,
                    k_nearest_neighbors=CANDIDATE_POOL_SIZE,
                    fields=CONTENT_VECTOR_FIELD,
                )
                vector_count = 0
                for rank, result in enumerate(
                    search_client.search(
                        search_text=None,
                        vector_queries=[vector_query],
                        **search_args,
                    ),
                    start=1,
                ):
                    vector_count += 1
                    record = dict(result)
                    record_id = str(record.get("id", ""))
                    existing = records_by_id.get(record_id)
                    if existing is None:
                        existing = record
                        records_by_id[record_id] = existing
                    else:
                        _merge_candidate_record(existing, record)
                    existing["_vector_rank"] = rank
                pass_counts["vector"] = vector_count
            except Exception:
                # Lexical retrieval remains a valid, bounded fallback when the
                # embedding provider or vector field is temporarily slow.
                pass
        candidate_records = [
            _add_runtime_sharepoint_parent_context(record)
            for record in _diversify_candidate_records(list(records_by_id.values()))
        ]
        if _requests_script(user_message):
            # Search the descriptive metadata explicitly.  A script can have
            # a generic body or filename, while SharePoint's ``Detalle`` is
            # the field that names its operational purpose.
            script_query = _script_search_query(user_message)
            if script_query:
                script_search_args = dict(search_args)
                script_search_args["select"] = [
                    *SEARCH_SELECT_FIELDS,
                    "artifact_role",
                ]
                try:
                    script_count = 0
                    for rank, result in enumerate(
                        search_client.search(
                            search_text=script_query,
                            search_fields=["title", CONTEXT_FIELD, "content_tokens"],
                            **script_search_args,
                        ),
                        start=1,
                    ):
                        script_count += 1
                        record = dict(result)
                        if not _is_script_record(record):
                            continue
                        record_id = str(record.get("id", ""))
                        existing = records_by_id.get(record_id)
                        if existing is None:
                            existing = record
                            records_by_id[record_id] = existing
                        else:
                            _merge_candidate_record(existing, record)
                        existing["_script_rank"] = rank
                    pass_counts["script"] = script_count
                    candidate_records = [
                        _add_runtime_sharepoint_parent_context(record)
                        for record in _diversify_candidate_records(
                            list(records_by_id.values())
                        )
                    ]
                except Exception:
                    # The ordinary lexical/vector candidates remain a safe
                    # fallback if an older index rejects the metadata pass.
                    pass
    except Exception:
        # Embeddings can be unavailable to a read-only diagnostic session.
        # Preserve semantic keyword retrieval in that case; only fall back to
        # plain keyword search if the index itself lacks semantic support.
        try:
            fallback_results = search_client.search(
                search_text=user_message,
                search_fields=["title", CONTENT_FIELD, "content_tokens"],
                **keyword_search_args,
            )
        except Exception:
            fallback_results = search_client.search(
                search_text=user_message,
                search_fields=["title", CONTENT_FIELD, "content_tokens"],
                **search_args,
            )
        candidate_records = [dict(result) for result in fallback_results]

    # A filename is an unambiguous document request. Resolve it through the
    # title field before applying semantic relevance so a related script cannot
    # displace the document the user explicitly named.
    if requested_file_names:
        records_by_id = {
            str(record.get("id", "")): record for record in candidate_records
        }
        try:
            explicit_count = 0
            for file_name in requested_file_names:
                for result in search_client.search(
                    search_text=file_name,
                    search_fields=["title"],
                    **search_args,
                ):
                    explicit_count += 1
                    record = dict(result)
                    if not _record_matches_file_name(record, (file_name,)):
                        continue
                    record["_explicit_filename_match"] = True
                    records_by_id[str(record.get("id", ""))] = record
            pass_counts["explicit_filename"] = explicit_count
            candidate_records = list(records_by_id.values())
        except Exception:
            # The regular keyword/vector candidates still provide a safe
            # fallback when a title-only query is temporarily unavailable.
            pass

    raw_candidate_count = len(candidate_records)
    candidate_records = [
        _add_runtime_sharepoint_parent_context(record) for record in candidate_records
    ]
    if diagnostics is not None:
        diagnostics["candidate_count"] = raw_candidate_count
        diagnostics["stage_counts"] = {
            "azure_union": raw_candidate_count,
            **{f"azure_{name}": count for name, count in pass_counts.items()},
        }
    allowed_sources = getattr(config, "sharepoint_sources", None)
    if allowed_sources is None:
        allowed_sources = tuple(getattr(config, "sharepoint_folder_paths", ()) or ())
    allowed_source_labels = tuple(getattr(config, "sharepoint_source_labels", ()) or ())
    authorized_records = [
        record
        for record in candidate_records
        if _record_has_authorized_provenance(
            record, allowed_sources, allowed_source_labels
        )
    ]
    rejected_reasons = diagnostics.setdefault("rejected_reasons", {}) if diagnostics is not None else {}
    if diagnostics is not None and len(authorized_records) < len(candidate_records):
        rejected_reasons["provenance"] = len(candidate_records) - len(authorized_records)
    injection_records = [
        record for record in authorized_records if _record_contains_document_injection(record)
    ]
    if injection_records:
        authorized_records = [
            record for record in authorized_records if not _record_contains_document_injection(record)
        ]
        if diagnostics is not None:
            rejected_reasons["document_injection"] = len(injection_records)
    if diagnostics is not None:
        diagnostics["stage_counts"]["authorized"] = len(authorized_records)
    country_scoped_records = _filter_records_for_requested_country(authorized_records, user_message)
    if diagnostics is not None and len(country_scoped_records) < len(authorized_records):
        rejected_reasons["country"] = len(authorized_records) - len(country_scoped_records)
    if diagnostics is not None:
        diagnostics["stage_counts"]["country_scoped"] = len(country_scoped_records)
    version_fallback_ids: set[str] = set()
    exact_version_records = [
        record
        for record in country_scoped_records
        if _record_matches_requested_version(record, requested_versions)
    ]
    if requested_versions and _is_structural_version_query(user_message):
        exact_technical_records = [
            record
            for record in exact_version_records
            if _record_matches_technical_anchor(record, user_message)
        ]
        if exact_technical_records:
            version_scoped_records = exact_technical_records
        else:
            technical_records = [
                record
                for record in country_scoped_records
                if _record_matches_technical_anchor(record, user_message)
            ]
            if technical_records:
                version_scoped_records = technical_records
                version_fallback_ids = {
                    str(record.get("id") or "") for record in technical_records
                }
            else:
                version_scoped_records = []
    else:
        version_scoped_records = exact_version_records
    # When the user supplied an exact version, the version guard above is the
    # authoritative constraint.  The proximity check is only for open-ended
    # questions such as "qué versión de jQuery", where a page can otherwise
    # win on the generic word "versión" alone.
    if _is_anchor_version_lookup(user_message) and not requested_versions:
        before_anchor_version = len(version_scoped_records)
        version_scoped_records = [
            record
            for record in version_scoped_records
            if _record_answers_anchor_version(record, user_message)
        ]
        if diagnostics is not None and len(version_scoped_records) < before_anchor_version:
            rejected_reasons["anchor_version_mismatch"] = (
                before_anchor_version - len(version_scoped_records)
            )
    if diagnostics is not None and requested_versions and len(version_scoped_records) < len(country_scoped_records):
        rejected_reasons["version_strict"] = len(country_scoped_records) - len(version_scoped_records)
    if diagnostics is not None and version_fallback_ids:
        rejected_reasons["version_fallback"] = len(version_fallback_ids)
    if requested_versions and _requests_readme(user_message):
        readme_records = [
            record
            for record in version_scoped_records
            if "readme" in str(record.get("title") or "").casefold()
        ]
        if readme_records:
            version_scoped_records = readme_records
    # Installation/update questions must not silently choose one release when
    # Azure returns multiple incompatible versioned Readmes. This policy is
    # intentionally limited to release operations; ordinary technical and
    # non-release questions keep the existing retrieval behavior.
    ambiguous_readme_versions = ()
    if not requested_versions and _is_release_guidance_question(user_message):
        ambiguous_readme_versions = _readme_versions(version_scoped_records)
        if len(ambiguous_readme_versions) > 1:
            if diagnostics is not None:
                diagnostics["requires_version_context"] = True
                diagnostics["ambiguous_readme_versions"] = len(ambiguous_readme_versions)
                diagnostics["rejected_reasons"]["ambiguous_release_version"] = len(
                    ambiguous_readme_versions
                )
            return []
    if diagnostics is not None:
        diagnostics["stage_counts"]["version_scoped"] = len(version_scoped_records)
        if _is_release_guidance_question(user_message):
            diagnostics["stage_counts"]["release_readme_candidates"] = sum(
                1
                for record in version_scoped_records
                if any(
                    marker in str(record.get("title") or "").casefold()
                    for marker in _RELEASE_GUIDANCE_TITLE_MARKERS
                )
            )
    if requested_section:
        section_records = [
            record
            for record in version_scoped_records
            if requested_section.search(
                f"{record.get('title', '')} {record.get(CONTEXT_FIELD, '')} "
                f"{record.get(CONTENT_FIELD, '')}"
            )
        ]
        # A section name is an explicit constraint, but preserve the normal
        # fallback if the index lacks that heading verbatim.
        if section_records:
            version_scoped_records = section_records
    if _requests_script(user_message):
        # Prefer executable artifacts over README incident histories.  Keep a
        # coverage-checked subset when possible, but retain the script-only
        # fallback so a short script can still be found through its title or
        # SharePoint description.
        script_records = [
            record for record in version_scoped_records if _is_script_record(record)
        ]
        if script_records:
            direct_script_records = [
                record
                for record in script_records
                if _has_minimum_content_coverage(
                    record,
                    user_message,
                    semantic_enabled=getattr(config, "azure_search_use_semantic", False),
                )
            ]
            version_scoped_records = direct_script_records or script_records
    explicit_file_records = [
        record
        for record in version_scoped_records
        if record.get("_explicit_filename_match")
        and _record_matches_file_name(record, requested_file_names)
    ]
    if requested_file_names:
        # An explicit filename is a strict lookup. Returning a related file
        # when it does not exist would falsely imply that the requested
        # document was found.
        if not explicit_file_records:
            return []
        candidate_records = explicit_file_records
    else:
        coverage_message = (
            _technical_anchor_query(user_message)
            if version_fallback_ids
            else user_message
        )
        if _requests_script(user_message) and any(
            _is_script_record(record) for record in version_scoped_records
        ):
            # The script-intent pass already restricted the candidate set to
            # executable artifacts. Do not require the literal word
            # ``script`` inside the file's business description; operators
            # normally describe what the script fixes, not its file type.
            candidate_records = version_scoped_records
        else:
            candidate_records = list(version_scoped_records)

    # Apply the global bound only after provenance, country, version, filename,
    # technical-anchor and script restrictions. A noisy Azure result must not
    # evict an authorized record before the security boundary is evaluated.
    merged_pool_limit = int(
        getattr(config, "retrieval_merged_pool_limit", MAX_MERGED_CANDIDATES)
    )
    rerank_pool_limit = int(
        getattr(config, "retrieval_rerank_pool_limit", RERANK_POOL_SIZE)
    )
    candidate_records = _limit_candidate_pool(
        candidate_records,
        merged_pool_limit,
        prioritize_release=_is_release_guidance_question(user_message),
    )
    if diagnostics is not None:
        diagnostics["stage_counts"]["bounded_pool"] = len(candidate_records)
        diagnostics["stage_counts"].update(_candidate_diversity_stats(candidate_records))

    if not (_requests_script(user_message) and any(
        _is_script_record(record) for record in version_scoped_records
    )) and not explicit_file_records:
        if _is_document_access_failure_question(user_message):
            direct_access_records = [
                record
                for record in candidate_records
                if _has_direct_document_access_failure_coverage(record)
            ]
            if diagnostics is not None and len(direct_access_records) < len(candidate_records):
                rejected_reasons["weak_document_access_coverage"] = (
                    len(candidate_records) - len(direct_access_records)
                )
            candidate_records = direct_access_records
        candidate_records = [
            record
            for record in candidate_records
            if _has_minimum_content_coverage(
                record,
                coverage_message,
                semantic_enabled=getattr(config, "azure_search_use_semantic", False),
            )
        ]
    if diagnostics is not None and len(candidate_records) < len(version_scoped_records):
        rejected_reasons["insufficient_direct_coverage"] = (
            len(version_scoped_records) - len(candidate_records)
        )
    if diagnostics is not None:
        diagnostics["stage_counts"]["eligible"] = len(candidate_records)
    rerank_records = _limit_candidate_pool(candidate_records, rerank_pool_limit)
    if diagnostics is not None:
        diagnostics["stage_counts"]["rerank_pool"] = len(rerank_records)
        if len(rerank_records) < len(candidate_records):
            rejected_reasons["rerank_pool_limit"] = len(candidate_records) - len(rerank_records)
    ranked_records = _rerank_records(rerank_records, user_message)
    if not ranked_records:
        if diagnostics is not None and candidate_records:
            rejected_reasons["relevance"] = len(candidate_records)
        return []

    # Avoid sending tangential pages to generation when one page has a much
    # stronger match. Multiple pages are retained when they are similarly
    # relevant, which still supports answers that span a section boundary.
    best_score = ranked_records[0][0]
    if best_score < 8 and not explicit_file_records:
        return []
    # Procedural and facet questions often span adjacent pages: one page names
    # the operation while the next lists its parameters or validation checks.
    # Keep a slightly wider, still bounded relevance band for those generic
    # question shapes so composition can recover the complete procedure. The
    # deterministic evidence gate below remains mandatory for every fragment.
    query_tokens = set(tokenize(user_message))
    multi_fragment_shape = bool(
        query_tokens.intersection({
            "parametro", "paso", "procedimiento", "estructura", "revis",
            "valid", "verific", "confirm", "hace", "como",
        })
    )
    relevance_floor = (
        0.65
        if multi_fragment_shape
        else (0.70 if _is_dtc_validation_question(user_message) else 0.80)
    )
    relevant_records = [
        item for item in ranked_records if item[0] >= best_score * relevance_floor
    ][:limit]
    if diagnostics is not None and len(relevant_records) < len(ranked_records):
        rejected_reasons["relevance_floor"] = len(ranked_records) - len(relevant_records)
    if diagnostics is not None:
        diagnostics["stage_counts"]["final_relevant"] = len(relevant_records)
    sources: list[EvidenceSource] = []
    for _, record in relevant_records:
        fragment = _result_fragment(record, user_message)
        if not fragment:
            continue
        source_system = record.get("source_system", "azure_ai_search")
        sources.append(
            EvidenceSource(
                tipo="sharepoint" if source_system == "sharepoint" else "azure_ai_search",
                titulo=record.get("title") or "Documento sin título",
                ubicacion=record.get("source_url") or "Azure AI Search",
                fragmento=fragment,
                source_system=source_system,
                document_id=str(record.get("document_id") or ""),
                document_version=str(record.get("document_version") or ""),
                last_modified=str(record.get("last_modified") or ""),
                document_type=str(record.get("document_type") or ""),
                folder_path=str(record.get("folder_path") or ""),
                descripcion=_record_description(record),
                version_confirmed=(
                    False
                    if str(record.get("id") or "") in version_fallback_ids
                    else (True if requested_versions else None)
                ),
                fallback_reason=(
                    "version_no_confirmada"
                    if str(record.get("id") or "") in version_fallback_ids
                    else ""
                ),
            )
        )
    return _deduplicate_equivalent_sources(sources, user_message)


def _v2_add_results(
    records_by_id: dict[str, dict], results: Iterable[dict], rank_field: str
) -> None:
    for rank, result in enumerate(results, start=1):
        record = dict(result)
        record_id = str(record.get("id") or "")
        if not record_id:
            continue
        existing = records_by_id.get(record_id)
        if existing is None:
            existing = record
            records_by_id[record_id] = existing
        else:
            existing.update(record)
        existing[rank_field] = min(int(existing.get(rank_field, rank)), rank)


def _v2_search_candidates(
    plan: QueryPlan, config, client=None
) -> tuple[list[dict], dict[str, int]]:
    """Collect vector, semantic and normalized lexical candidates before ranking."""
    search_client = SearchClient(
        endpoint=config.azure_search_endpoint,
        index_name=config.azure_search_index_name,
        credential=_credential(config),
    )
    search_args = {
        "top": CANDIDATE_POOL_SIZE,
        "select": V2_SEARCH_SELECT_FIELDS,
        "connection_timeout": SEARCH_TIMEOUT_SECONDS,
        "read_timeout": SEARCH_TIMEOUT_SECONDS,
    }
    records_by_id: dict[str, dict] = {}
    rejected: dict[str, int] = {}
    try:
        embedding = _embed_texts([plan.raw_message], config, client=client)[0]
        _v2_add_results(
            records_by_id,
            search_client.search(
                search_text=None,
                vector_queries=[
                    VectorizedQuery(
                        vector=embedding,
                        k_nearest_neighbors=CANDIDATE_POOL_SIZE,
                        fields=CONTENT_VECTOR_FIELD,
                    )
                ],
                **search_args,
            ),
            "_vector_rank",
        )
    except Exception:
        rejected["vector_unavailable"] = 1

    for query_index, query in enumerate(plan.retrieval_queries):
        try:
            if query_index == 0 and getattr(config, "azure_search_use_semantic", False):
                results = search_client.search(
                    search_text=query,
                    search_fields=["title", RETRIEVAL_TEXT_FIELD, CONTENT_FIELD, "content_tokens"],
                    query_type="semantic",
                    semantic_configuration_name=config.azure_search_semantic_configuration,
                    query_caption="extractive",
                    **search_args,
                )
                rank_field = "_semantic_rank"
            else:
                results = search_client.search(
                    search_text=query,
                    search_fields=["title", RETRIEVAL_TEXT_FIELD, RETRIEVAL_CONCEPTS_FIELD, CONTENT_FIELD, "content_tokens"],
                    search_mode="any",
                    **search_args,
                )
                rank_field = "_lexical_rank"
            _v2_add_results(records_by_id, results, rank_field)
        except Exception:
            rejected["lexical_query_failed"] = rejected.get("lexical_query_failed", 0) + 1
    # Keep every unique candidate until evidence validation. Applying the
    # per-document cap here can discard the one factual page of a manual when
    # vector similarity has already supplied several other pages from it.
    return list(records_by_id.values()), rejected


def _v2_record_allowed(
    record: dict, plan: QueryPlan, allowed_sources, allowed_source_labels=()
) -> tuple[bool, str]:
    if not _record_has_authorized_provenance(
        record, allowed_sources, allowed_source_labels
    ):
        return False, "provenance"
    if _record_contains_document_injection(record):
        return False, "document_injection"
    if str(record.get("quality_status") or "pendiente").casefold() in EXCLUDED_QUALITY_STATUSES:
        return False, "quality_status"
    if str(record.get("evidence_kind") or "primary").casefold() in {"navigation", "reference"}:
        return False, "evidence_kind"
    if plan.version and not _record_matches_requested_version(record, (plan.version,)):
        return False, "version"
    if plan.artifact_role and str(record.get("artifact_role") or "").casefold() != plan.artifact_role:
        return False, "artifact_role"
    return True, ""


def _v2_direct_text(record: dict) -> str:
    """Use the local fragment plus reviewed, indexed document facts as evidence."""
    # Keep structural facts and the first local evidence sentence in the same
    # evidence window. A filename such as ``proc.sql`` is not prose, so
    # separating it with periods would make a question asking for a *script*
    # fail even when the SQL body directly covers the requested operation.
    return " ".join(
        value
        for value in (
            f"El artefacto de tipo {record.get('artifact_role') or 'documento'} "
            f"titulado {record.get('title') or ''} "
            f"con operación {record.get('operation') or ''} contiene el fragmento",
            str(record.get(CONTENT_FIELD) or ""),
        )
        if value
    )


def _v2_validated_evidence(record: dict, plan: QueryPlan) -> tuple[str, tuple[str, ...]]:
    """Return only local fragments that directly cover individual requirements."""
    quality_status = str(record.get("quality_status") or "").casefold()
    reviewed_operation = str(record.get("operation") or "") if quality_status == "aprobado" else ""
    artifact_context = (
        str(record.get("title") or "")
        if str(record.get("artifact_role") or "").casefold() == "script"
        else ""
    )
    artifact_role = str(record.get("artifact_role") or "")
    fragments: list[str] = []
    covered: list[str] = []
    for requirement in plan.requirements:
        fragment = _result_fragment(record, requirement.text)
        if not fragment:
            continue
        supporting_text = " ".join(
            value
            for value in (
                f"El artefacto de tipo {artifact_role or 'documento'} "
                f"titulado {artifact_context} con operación {reviewed_operation} contiene",
                fragment,
            )
            if value
        )
        if not requirement_is_covered(requirement, supporting_text):
            continue
        context_lines = []
        if artifact_context:
            context_lines.append(f"Artefacto de código: {artifact_context}")
        if artifact_role:
            context_lines.append(f"Tipo de artefacto: {artifact_role}")
        if reviewed_operation:
            context_lines.append(f"Operación revisada: {reviewed_operation}")
        if context_lines:
            context_prefix = "\n".join(context_lines)
            fragment = f"{context_prefix}\n{fragment}"
        if fragment not in fragments:
            fragments.append(fragment)
        covered.append(requirement.identifier)
    return "\n\n".join(fragments), tuple(covered)


def _v2_semantic_candidate_payload(record: dict, plan: QueryPlan) -> dict:
    """Keep semantic verification bounded to local, already-retrieved text."""
    fragments = [
        {
            "requirement_id": requirement.identifier,
            "fragment": _result_fragment(record, requirement.text)[:1_000],
        }
        for requirement in plan.requirements
    ]
    return {
        # candidate_id is assigned by the bounded caller; never serialize an
        # Azure/document identifier to the model.
        "candidate_id": str(record["_verifier_candidate_id"]),
        "fragments": fragments,
    }


def _v2_semantic_evidence(
    record: dict, plan: QueryPlan, coverage: tuple[str, ...]
) -> tuple[str, tuple[str, ...]]:
    requirement_by_id = {requirement.identifier: requirement for requirement in plan.requirements}
    fragments: list[str] = []
    for requirement_id in coverage:
        requirement = requirement_by_id.get(requirement_id)
        if requirement is None:
            continue
        fragment = _result_fragment(record, requirement.text)
        if fragment and fragment not in fragments:
            fragments.append(fragment)
    if not fragments:
        return "", ()
    if str(record.get("artifact_role") or "").casefold() == "script":
        fragments.insert(0, f"Artefacto de código: {record.get('title') or ''}")
    return "\n\n".join(fragments), coverage


def _v2_semantic_coverage_is_anchored(
    record: dict, plan: QueryPlan, coverage: tuple[str, ...]
) -> bool:
    """Keep semantic paraphrase from crossing to a different technical subject."""
    requirement_by_id = {requirement.identifier: requirement for requirement in plan.requirements}
    is_code = str(record.get("artifact_role") or "").casefold() == "script"
    for requirement_id in coverage:
        requirement = requirement_by_id.get(requirement_id)
        if requirement is None or requirement.text.startswith("el calificador"):
            return False
        fragment = _result_fragment(record, requirement.text)
        if not fragment:
            return False
        # Semantic validation can bridge paraphrased actions, but it cannot
        # use a heading that names the requested topic and a distant sentence
        # that names a generic parameter as if they described the same fact.
        # Keep the topical anchors in one local prose window. Code is the
        # deliberate exception: identifiers, condition and operation normally
        # occupy separate lines of the same bounded declaration.
        units = [
            unit.strip()
            for unit in re.split(r"(?<=[.!?;])\s+|\n+", fragment)
            if unit.strip()
        ]
        windows = list(units)
        if is_code:
            windows.append(
                " ".join(
                    (str(record.get("title") or ""), str(record.get("operation") or ""), fragment)
                )
            )
        non_action_anchors = set(requirement.concepts).difference(requirement.actions)
        if non_action_anchors and not any(
            non_action_anchors.issubset(set(concept_keys(window))) for window in windows
        ):
            return False
        # The verifier may recognize wording variants, but it must not turn a
        # passage that merely mentions the subject into an answer for a named
        # operation such as classifying, configuring or exporting it.
        if requirement.actions and not any(
            has_requested_action_coverage(requirement.text, window) for window in windows
        ):
            return False
        # A code artifact may be semantically summarized, but its named
        # operation must still appear in the title, reviewed operation or code.
        if is_code and requirement.actions and not set(requirement.actions).intersection(
            set(concept_keys(" ".join(windows)))
        ):
            return False
    return True


def _v2_semantic_verifier_records(records: list[dict], plan: QueryPlan, limit: int = 12) -> list[dict]:
    """Give the bounded verifier distinct documents, not many chunks of one."""
    selected: list[dict] = []
    seen_documents: set[str] = set()
    for record in sorted(records, key=lambda item: _v2_score(item, (), plan), reverse=True):
        document_key = str(record.get("document_id") or record.get("id") or "")
        if document_key in seen_documents:
            continue
        seen_documents.add(document_key)
        selected.append({**record, "_verifier_candidate_id": f"c{len(selected) + 1:02d}"})
        if len(selected) >= limit:
            break
    return selected


def _v2_score(record: dict, coverage: tuple[str, ...], plan: QueryPlan) -> float:
    document_concepts = set(
        concept_keys(
            " ".join(
                str(record.get(field) or "")
                for field in ("title", RETRIEVAL_TEXT_FIELD, RETRIEVAL_CONCEPTS_FIELD, CONTENT_FIELD)
            )
        )
    )
    plan_concepts = {
        concept for requirement in plan.requirements for concept in requirement.concepts
    }
    title_concepts = set(concept_keys(str(record.get("title") or "")))
    raw_query_concepts = set(concept_keys(plan.raw_message))
    is_preinstallation_question = (
        _PREINSTALLATION_OPERATION_CONCEPTS.issubset(raw_query_concepts)
        and bool(raw_query_concepts.intersection(_PREINSTALLATION_CUE_CONCEPTS))
    )
    preinstallation_evidence_score = (
        len(document_concepts.intersection(_PREINSTALLATION_EVIDENCE_CONCEPTS)) * 18
        if is_preinstallation_question
        else 0
    )
    # An all-uppercase identifier such as DTC is an explicit technical anchor,
    # not a synonym. Prefer a source whose title names that identifier without
    # mapping it to any particular document or product.
    title_folded = str(record.get("title") or "").casefold()
    title_acronym_score = min(
        70,
        sum(
            35
            for acronym in set(_UPPERCASE_ACRONYM.findall(plan.raw_message))
            if acronym.casefold() in title_folded
        ),
    )
    semantic_score = float(record.get("@search.reranker_score") or 0)
    return (
        len(coverage) * 100
        + len(plan_concepts.intersection(document_concepts)) * 8
        + len(plan_concepts.intersection(title_concepts)) * 12
        + semantic_score * 20
        + max(0, CANDIDATE_POOL_SIZE - int(record.get("_vector_rank", CANDIDATE_POOL_SIZE))) * 0.5
        + max(0, CANDIDATE_POOL_SIZE - int(record.get("_semantic_rank", CANDIDATE_POOL_SIZE))) * 0.3
        + max(0, CANDIDATE_POOL_SIZE - int(record.get("_lexical_rank", CANDIDATE_POOL_SIZE))) * 0.2
        + preinstallation_evidence_score
        + title_acronym_score
    )


def _retrieve_v2_azure_search_evidence(
    user_message: str, config, client=None
) -> RetrievalTrace:
    plan = build_query_plan(user_message)
    candidate_records, rejected = _v2_search_candidates(plan, config, client=client)
    raw_candidate_count = len(candidate_records)
    candidate_records = _filter_records_for_requested_country(candidate_records, user_message)
    allowed_sources = getattr(config, "sharepoint_sources", None)
    if allowed_sources is None:
        allowed_sources = tuple(getattr(config, "sharepoint_folder_paths", ()) or ())
    allowed_source_labels = tuple(getattr(config, "sharepoint_source_labels", ()) or ())
    direct_records: list[tuple[float, dict, tuple[str, ...], str]] = []
    semantic_candidates: list[dict] = []
    for record in candidate_records:
        allowed, reason = _v2_record_allowed(
            record, plan, allowed_sources, allowed_source_labels
        )
        if not allowed:
            rejected[reason] = rejected.get(reason, 0) + 1
            continue
        coverage = covered_requirements(plan, _v2_direct_text(record))
        if not coverage:
            semantic_candidates.append(record)
            continue
        direct_records.append((_v2_score(record, coverage, plan), record, coverage, "deterministic"))

    if getattr(config, "use_llm_evidence_verifier", False) and semantic_candidates:
        verifier_records = _v2_semantic_verifier_records(semantic_candidates, plan)
        try:
            payloads = [_v2_semantic_candidate_payload(record, plan) for record in verifier_records]
            semantic_coverage = verify_semantic_evidence(
                plan,
                payloads,
                client=client,
                model=getattr(config, "evidence_verifier_model_name", config.openai_intent_model_name),
            )
            for record in verifier_records:
                coverage = semantic_coverage.get(str(record["_verifier_candidate_id"]), ())
                if coverage and _v2_semantic_coverage_is_anchored(record, plan, coverage):
                    direct_records.append((_v2_score(record, coverage, plan), record, coverage, "semantic"))
        except Exception:
            logger.warning("Falló el verificador semántico de evidencia; se conservó la decisión determinista.")
            rejected["semantic_verifier_failed"] = rejected.get("semantic_verifier_failed", 0) + 1

    rejected["insufficient_direct_evidence"] = (
        rejected.get("insufficient_direct_evidence", 0) + len(semantic_candidates)
    )
    direct_records.sort(key=lambda item: item[0], reverse=True)
    sources: list[EvidenceSource] = []
    covered: set[str] = set()
    sources_per_document: dict[str, int] = {}
    for _, record, coverage, verification_mode in direct_records:
        document_key = str(record.get("document_id") or record.get("id") or "")
        if sources_per_document.get(document_key, 0) >= MAX_CANDIDATES_PER_DOCUMENT:
            rejected["document_diversity"] = rejected.get("document_diversity", 0) + 1
            continue
        fragment, validated_coverage = (
            _v2_semantic_evidence(record, plan, coverage)
            if verification_mode == "semantic"
            else _v2_validated_evidence(record, plan)
        )
        if not fragment or not validated_coverage:
            rejected["fragment_without_direct_coverage"] = rejected.get("fragment_without_direct_coverage", 0) + 1
            continue
        if set(validated_coverage).issubset(covered):
            rejected["redundant_direct_evidence"] = rejected.get("redundant_direct_evidence", 0) + 1
            continue
        sources.append(
            EvidenceSource(
                tipo="sharepoint" if record.get("source_system") == "sharepoint" else "azure_ai_search",
                titulo=record.get("title") or "Documento sin título",
                ubicacion=record.get("source_url") or "Azure AI Search",
                fragmento=fragment,
                source_system=str(record.get("source_system") or ""),
                document_id=str(record.get("document_id") or ""),
                document_version=str(record.get("document_version") or ""),
                last_modified=str(record.get("last_modified") or ""),
                document_type=str(record.get("document_type") or ""),
                folder_path=str(record.get("folder_path") or ""),
                artifact_role=str(record.get("artifact_role") or ""),
                quality_status=str(record.get("quality_status") or ""),
                evidence_kind=str(record.get("evidence_kind") or ""),
                covered_requirements=validated_coverage,
                descripcion=_record_description(record),
            )
        )
        covered.update(validated_coverage)
        sources_per_document[document_key] = sources_per_document.get(document_key, 0) + 1
        if set(plan.requirement_ids).issubset(covered) or len(sources) >= 3:
            break
    trace = RetrievalTrace(
        sources=sources,
        query_hash=plan.query_hash,
        candidate_count=len(candidate_records),
        direct_evidence_count=len(sources),
        requirement_count=len(plan.requirements),
        covered_requirement_count=len(covered),
        rejected_reasons=rejected,
        stage_counts={
            "azure_union": raw_candidate_count,
            "post_hard_filters": len(candidate_records),
            "llm_candidate_pool": len(verifier_records) if getattr(config, "use_llm_evidence_verifier", False) else 0,
            "final_relevant": len(sources),
        },
    )
    logger.info(
        "retrieval_v2 query_hash=%s candidates=%s direct=%s requirements=%s covered=%s rejected=%s",
        trace.query_hash,
        trace.candidate_count,
        trace.direct_evidence_count,
        trace.requirement_count,
        trace.covered_requirement_count,
        ",".join(f"{key}:{value}" for key, value in sorted(trace.rejected_reasons.items())) or "none",
    )
    return trace


def retrieve_azure_search_evidence(
    user_message: str, config, client=None, return_trace: bool = False
) -> list[EvidenceSource] | RetrievalTrace:
    """Select the reversible retrieval strategy configured for this environment."""
    if getattr(config, "retrieval_strategy", "legacy") == "v2":
        trace = _retrieve_v2_azure_search_evidence(user_message, config, client=client)
        return trace if return_trace else trace.sources
    diagnostics: dict = {}
    sources = _retrieve_legacy_azure_search_evidence(
        user_message,
        config,
        client=client,
        diagnostics=diagnostics if return_trace else None,
    )
    if return_trace:
        return RetrievalTrace(
            sources=sources,
            candidate_count=int(diagnostics.get("candidate_count", 0)),
            direct_evidence_count=len(sources),
            rejected_reasons=dict(diagnostics.get("rejected_reasons", {})),
            stage_counts=dict(diagnostics.get("stage_counts", {})),
            requires_version_context=bool(diagnostics.get("requires_version_context", False)),
        )
    return sources


def _document_pages(document_path: Path) -> list[tuple[int | None, str]]:
    """Extract text by page so a search result keeps its original context."""
    if document_path.stat().st_size == 0:
        return []
    extension = document_path.suffix.lower()
    if extension == ".pdf":
        reader = PdfReader(str(document_path))
        return [
            (page_number, page.extract_text() or "")
            for page_number, page in enumerate(reader.pages, start=1)
        ]
    if extension == ".docx":
        from docx import Document

        document = Document(str(document_path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return [(None, "\n".join(parts))]
    if extension == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(str(document_path), read_only=True, data_only=True)
        parts: list[str] = []
        for worksheet in workbook.worksheets:
            parts.append(f"Hoja: {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    parts.append(" | ".join(values))
        workbook.close()
        return [(None, "\n".join(parts))]
    return [(None, document_path.read_text(encoding="utf-8", errors="replace"))]


def _chunks(
    text: str,
    size: int = 450,
    overlap: int = 75,
    max_characters: int = 6_000,
) -> Iterable[str]:
    """Split text by words and characters for safe embedding requests."""
    words = text.split()
    if not words:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = start
        character_count = 0
        while end < len(words) and end < start + size:
            word_length = len(words[end])
            separator_length = 1 if end > start else 0
            if end > start and character_count + separator_length + word_length > max_characters:
                break
            if end == start and word_length > max_characters:
                for offset in range(0, word_length, max_characters):
                    chunks.append(words[end][offset : offset + max_characters])
                end += 1
                break
            character_count += separator_length + word_length
            end += 1
        split_long_word = False
        if end == start:
            end += 1
        elif len(words[start]) > max_characters:
            split_long_word = True
        if not split_long_word:
            chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start = max(start + 1, end - overlap)
    return chunks


def _metadata_for(document_path: Path) -> dict:
    metadata_path = document_path.with_suffix(document_path.suffix + ".metadata.json")
    if not metadata_path.exists():
        return {}
    try:
        return json.loads(metadata_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def _review_metadata(metadata: dict) -> dict:
    """Read optional, reviewed functional metadata without inventing it."""
    reviewed = metadata.get("libras", {}) if isinstance(metadata, dict) else {}
    if not isinstance(reviewed, dict):
        reviewed = {}
    return {
        # Functional fields are accepted only from the explicit ``libras``
        # review block.  Graph/SharePoint technical metadata must never turn
        # into an unreviewed product, module or operation assertion.
        field: str(reviewed.get(field) or "").strip()
        for field in ("product", "module", "operation", "artifact_role", "version", "country", "quality_status")
    }


def _artifact_role(document_path: Path, reviewed: dict) -> str:
    return str(reviewed.get("artifact_role") or ARTIFACT_ROLE_BY_EXTENSION.get(document_path.suffix.lower(), "otro"))


def _sql_routine_identifiers(text: str) -> list[str]:
    return list(
        dict.fromkeys(
            match.group(1)
            for match in re.finditer(
                r"(?im)^\s*create\s+(?:or\s+alter\s+)?(?:procedure|proc|function|view|trigger)\s+([\[\]a-z0-9_.-]+)",
                text or "",
            )
        )
    )


def _chunks_for_document(document_path: Path, text: str) -> Iterable[str]:
    """Keep code declarations intact while retaining generic chunking elsewhere."""
    if document_path.suffix.lower() != ".sql":
        return _chunks(text)
    starts = [match.start() for match in re.finditer(
        r"(?im)^\s*create\s+(?:or\s+alter\s+)?(?:procedure|proc|function|view|trigger)\b",
        text or "",
    )]
    if not starts:
        return _chunks(text)
    starts.append(len(text))
    chunks: list[str] = []
    if starts[0] > 0 and text[: starts[0]].strip():
        chunks.extend(_chunks(text[: starts[0]]))
    for start, end in zip(starts, starts[1:]):
        declaration = text[start:end].strip()
        if len(declaration) <= 6_000:
            chunks.append(declaration)
        else:
            chunks.extend(_chunks(declaration))
    return chunks


def _is_navigation_text(text: str) -> bool:
    normalized = " ".join((text or "").casefold().split())
    page_references = re.findall(r",\s*\d{1,3}(?=\s|$)", normalized)
    return "tabla de contenido" in normalized or (
        len(page_references) >= 6
        and any(marker in normalized for marker in ("pagina", "página", "indice", "contenido", "seccion", "capitulo"))
    )


def _evidence_kind(title: str, text: str) -> str:
    """Classify structural non-factual artifacts without inferring domain meaning."""
    if _is_navigation_text(text):
        return "navigation"
    normalized_title = " ".join((title or "").casefold().split())
    if any(marker in normalized_title for marker in ("links de apoyo", "bibliografía", "bibliografia", "referencias externas")):
        return "reference"
    return "primary"


def _retrieval_metadata_text(
    title: str,
    readable_title: str,
    parent_context: str,
    artifact_role: str,
    reviewed: dict,
    routine_identifiers: list[str],
    description: str = "",
    dependency: str = "",
) -> tuple[str, str]:
    """Build retrieval-only text from source facts and reviewed metadata."""
    source_facts = [
        f"Archivo {title}",
        f"Terminos de archivo {readable_title}",
        f"Tipo de artefacto {artifact_role}",
    ]
    if parent_context:
        source_facts.append(f"Carpetas de origen {parent_context}")
    if routine_identifiers:
        source_facts.append(f"Declaraciones de codigo {' '.join(routine_identifiers)}")
    if description:
        source_facts.append(f"Descripción de la solución {description[:900]}")
    if dependency:
        source_facts.append(f"Dependencia {dependency[:300]}")
    for field in ("product", "module", "operation", "version", "country"):
        if reviewed.get(field):
            source_facts.append(f"{field} revisado {reviewed[field]}")
    retrieval_text = ". ".join(source_facts) + "."
    return retrieval_text, " ".join(concept_keys(retrieval_text))


def _add_runtime_sharepoint_parent_context(record: dict) -> dict:
    """Enrich an older indexed record with its folder trail for ranking.

    The legacy production index can be storage-full, so historical records
    cannot always be rewritten immediately. Their SharePoint URL is already a
    trusted indexed field, however, and preserves the parent folders. Derive
    that retrieval context in memory rather than losing a solution stored in a
    generically named file such as ``Indicaciones.txt``.
    """
    parent_context = _sharepoint_parent_context(
        str(record.get("source_url") or ""),
        str(record.get("folder_path") or ""),
    )
    if not parent_context:
        return record
    context = str(record.get(CONTEXT_FIELD) or "")
    marker = f"Carpetas de origen: {parent_context}."
    if parent_context.casefold() not in context.casefold():
        record[CONTEXT_FIELD] = _clean_text(f"{context} {marker}", limit=1_200)
    return record


def _document_records(
    source_dir: Path, document_ids: set[str] | None = None
) -> list[dict]:
    records: list[dict] = []
    documents = [
        path
        for path in sorted(source_dir.rglob("*"))
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    # Once the OneDrive sync has written metadata, index only those managed
    # copies. This prevents duplicate results from PDFs manually staged before
    # delegated access was approved.
    has_managed_documents = any(
        path.with_suffix(path.suffix + ".metadata.json").exists() for path in documents
    )
    sync_state_path = source_dir / SYNC_STATE_NAME
    synchronized_document_ids: set[str] | None = None
    if sync_state_path.exists():
        try:
            sync_state = json.loads(sync_state_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError as error:
            raise RuntimeError(
                f"El estado de sincronización no es JSON válido: {sync_state_path}"
            ) from error
        if not isinstance(sync_state, dict):
            raise RuntimeError(
                f"El estado de sincronización debe ser un objeto JSON: {sync_state_path}"
            )
        synchronized_documents = sync_state.get("documents", {})
        if isinstance(synchronized_documents, dict):
            synchronized_document_ids = set(synchronized_documents)
    for document_path in documents:
        if has_managed_documents and not document_path.with_suffix(
            document_path.suffix + ".metadata.json"
        ).exists():
            continue
        metadata = _metadata_for(document_path)
        source_url = metadata.get("web_url") or str(document_path.resolve())
        title = metadata.get("name") or document_path.stem
        source_system = metadata.get("source_system", "local")
        document_id = str(metadata.get("document_id") or source_url)
        if (
            synchronized_document_ids is not None
            and source_system == "sharepoint"
            and document_id not in synchronized_document_ids
        ):
            continue
        if document_ids is not None and document_id not in document_ids:
            continue
        pages = _document_pages(document_path)
        full_text = "\n".join(text for _, text in pages)
        if not full_text.strip():
            continue
        document_version = str(metadata.get("etag") or metadata.get("document_version") or "")
        last_modified = str(metadata.get("last_modified") or "")
        folder_path = str(metadata.get("folder_path") or "")
        drive_id = str(metadata.get("drive_id") or "")
        parent_context = _sharepoint_parent_context(source_url, folder_path)
        reviewed = _review_metadata(metadata)
        artifact_role = _artifact_role(document_path, reviewed)
        quality_status = reviewed.get("quality_status") or "pendiente"
        # Later pages frequently omit the country/product named on the cover.
        # Store compact document-level context with every chunk. The document
        # body is already in ``content``; duplicating it here wastes scarce
        # legacy-index storage and makes an incremental metadata correction
        # larger than the record it replaces.
        readable_title = _searchable_filename_terms(title)
        routine_identifiers = _sql_routine_identifiers(full_text) if artifact_role == "script" else []
        description = str(metadata.get("description") or metadata.get("detalle") or "").strip()
        dependency = str(metadata.get("dependency") or metadata.get("dependencia") or "").strip()
        retrieval_text, retrieval_concepts = _retrieval_metadata_text(
            title,
            readable_title,
            parent_context,
            artifact_role,
            reviewed,
            routine_identifiers,
            description,
            dependency,
        )
        document_context = _clean_text(
            f"Título del archivo: {title}. Términos del nombre: {readable_title}. "
            f"Carpetas de origen: {parent_context}."
            + (f" Descripción de la solución: {description}." if description else "")
            + (f" Dependencia: {dependency}." if dependency else ""),
            limit=1_200,
        )
        content_hash = hashlib.sha256(full_text.encode("utf-8")).hexdigest()
        document_key = hashlib.sha256(document_id.encode("utf-8")).hexdigest()
        sequence = 0
        for page_number, page_text in pages:
            for chunk in _chunks_for_document(document_path, page_text):
                page_label = f"Página {page_number}" if page_number else "Documento"
                records.append(
                    {
                        "id": f"{document_key}-{sequence}",
                        "title": f"{title} — {page_label}",
                        "content": f"{page_label}\n{chunk}",
                        CONTEXT_FIELD: document_context,
                        "source_url": source_url,
                        "source_system": source_system,
                        "document_id": document_id,
                        "document_version": document_version,
                        "last_modified": last_modified,
                        "content_hash": content_hash,
                        "document_type": document_path.suffix.lower().lstrip("."),
                        "folder_path": folder_path,
                        "drive_id": drive_id,
                        RETRIEVAL_TEXT_FIELD: retrieval_text,
                        RETRIEVAL_CONCEPTS_FIELD: retrieval_concepts,
                        "product": reviewed.get("product", ""),
                        "module": reviewed.get("module", ""),
                        "operation": reviewed.get("operation") or " ".join(routine_identifiers),
                        "artifact_role": artifact_role,
                        "version": reviewed.get("version", ""),
                        "country": reviewed.get("country", ""),
                        "quality_status": quality_status,
                        "evidence_kind": _evidence_kind(title, chunk),
                        "indexed_at": datetime.now(timezone.utc).isoformat(),
                        "chunk_number": sequence,
                "content_tokens": " ".join(
                            tokenize(
                                f"{title} {readable_title} {parent_context} "
                                f"{retrieval_text} {description} {dependency} {chunk}"
                            )
                        ),
                    }
                )
                sequence += 1
    return records


def _existing_record_ids(client: SearchClient, document_ids: set[str]) -> set[str]:
    """Find prior chunks so updates and deletions remove stale fragments."""
    existing_ids: set[str] = set()
    for document_id in document_ids:
        escaped_id = document_id.replace("'", "''")
        results = client.search(
            search_text="*",
            filter=f"document_id eq '{escaped_id}'",
            select=["id"],
            top=1_000,
        )
        existing_ids.update(str(result["id"]) for result in results)
    return existing_ids


def ensure_index(config) -> None:
    """Create the minimal text/semantic index when it does not exist."""
    index_client = SearchIndexClient(
        endpoint=config.azure_search_endpoint,
        credential=_credential(config),
    )
    try:
        index = index_client.get_index(config.azure_search_index_name)
        field_names = {field.name for field in index.fields}
        if "drive_id" not in field_names:
            raise RuntimeError(
                "El índice existente no contiene drive_id; ejecute la ingesta con --reset-index "
                "para migrar al alcance multi-biblioteca."
            )
        if getattr(config, "retrieval_strategy", "legacy") == "v2" and not {
            RETRIEVAL_TEXT_FIELD,
            RETRIEVAL_CONCEPTS_FIELD,
            "artifact_role",
            "quality_status",
            "evidence_kind",
        }.issubset(field_names):
            raise RuntimeError(
                "El índice no contiene los campos de calidad v2; ejecute una reconstrucción controlada."
            )
        return
    except ResourceNotFoundError:
        pass

    fields = [
        SimpleField(name="id", type=SearchFieldDataType.String, key=True, filterable=True),
        SearchableField(name="title", type=SearchFieldDataType.String, searchable=True),
        SearchableField(name=CONTENT_FIELD, type=SearchFieldDataType.String, searchable=True),
        SearchableField(name=CONTEXT_FIELD, type=SearchFieldDataType.String, searchable=True),
        SearchableField(name=RETRIEVAL_TEXT_FIELD, type=SearchFieldDataType.String, searchable=True),
        SearchableField(name=RETRIEVAL_CONCEPTS_FIELD, type=SearchFieldDataType.String, searchable=True),
        SearchField(
            name=CONTENT_VECTOR_FIELD,
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=config.openai_embedding_dimensions,
            vector_search_profile_name="content-vector-profile",
        ),
        SimpleField(name="source_url", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="source_system", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="document_id", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="document_version", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="last_modified", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="content_hash", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="document_type", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="folder_path", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="drive_id", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="indexed_at", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="chunk_number", type=SearchFieldDataType.Int32, filterable=True),
        SearchableField(name="content_tokens", type=SearchFieldDataType.String, searchable=True),
        SimpleField(name="product", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="module", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="operation", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="artifact_role", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="version", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="country", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="quality_status", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="evidence_kind", type=SearchFieldDataType.String, filterable=True),
    ]
    semantic_search = SemanticSearch(
        configurations=[
            SemanticConfiguration(
                name=config.azure_search_semantic_configuration,
                prioritized_fields=SemanticPrioritizedFields(
                    title_field=SemanticField(field_name="title"),
                    content_fields=[
                        SemanticField(field_name=RETRIEVAL_TEXT_FIELD),
                        SemanticField(field_name=CONTENT_FIELD),
                    ],
                ),
            )
        ]
    )
    index_client.create_index(
        SearchIndex(
            name=config.azure_search_index_name,
            fields=fields,
            semantic_search=semantic_search,
            vector_search=VectorSearch(
                algorithms=[HnswAlgorithmConfiguration(name="content-vector-hnsw")],
                profiles=[
                    VectorSearchProfile(
                        name="content-vector-profile",
                        algorithm_configuration_name="content-vector-hnsw",
                    )
                ],
            ),
        )
    )


def reset_index(config) -> None:
    """Delete and recreate only the explicitly configured pilot index."""
    index_client = SearchIndexClient(
        endpoint=config.azure_search_endpoint,
        credential=_credential(config),
    )
    try:
        index_client.delete_index(config.azure_search_index_name)
    except ResourceNotFoundError:
        pass
    ensure_index(config)


def _deletion_document_ids(source_dir: Path) -> set[str]:
    """Read idempotent SharePoint deletion notices written by the sync step."""
    manifest_path = source_dir / DELETION_MANIFEST_NAME
    if not manifest_path.exists():
        return set()
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        raise RuntimeError(f"El manifiesto de eliminaciones no es válido: {manifest_path}") from error
    deletions = payload.get("deleted_document_ids", [])
    if not isinstance(deletions, list) or not all(isinstance(item, str) for item in deletions):
        raise RuntimeError("El manifiesto de eliminaciones debe contener deleted_document_ids.")
    return {item for item in deletions if item}


def _changed_document_ids(source_dir: Path) -> set[str] | None:
    """Read pending SharePoint upserts, or ``None`` for legacy full ingests."""
    manifest_path = source_dir / CHANGE_MANIFEST_NAME
    if not manifest_path.exists():
        return None
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        raise RuntimeError(f"El manifiesto de cambios no es válido: {manifest_path}") from error
    changes = payload.get("changed_document_ids", [])
    if not isinstance(changes, list) or not all(isinstance(item, str) for item in changes):
        raise RuntimeError("El manifiesto de cambios debe contener changed_document_ids.")
    return {item for item in changes if item}


def _delete_record_ids(client: SearchClient, record_ids: set[str]) -> None:
    for offset in range(0, len(record_ids), 500):
        batch = list(record_ids)[offset : offset + 500]
        results = client.delete_documents(documents=[{"id": record_id} for record_id in batch])
        failures = [result.key for result in results if not result.succeeded]
        if failures:
            raise RuntimeError(f"Azure AI Search no eliminó {len(failures)} fragmentos.")


def _clear_deletion_manifest(source_dir: Path) -> None:
    manifest_path = source_dir / DELETION_MANIFEST_NAME
    if manifest_path.exists():
        manifest_path.write_text(
            json.dumps({"deleted_document_ids": []}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


def _clear_change_manifest(source_dir: Path) -> None:
    manifest_path = source_dir / CHANGE_MANIFEST_NAME
    if manifest_path.exists():
        manifest_path.write_text(
            json.dumps({"changed_document_ids": []}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


def _records_supported_by_index(records: list[dict], field_names: set[str]) -> list[dict]:
    """Drop additive fields when updating a legacy production index.

    The v2 candidate has extra retrieval fields, while the existing Free-tier
    pilot intentionally remains on its smaller legacy schema.  A focused
    upsert must update compatible fields instead of requiring a risky index
    reset just because a new optional field exists locally.
    """
    return [
        {name: value for name, value in record.items() if name in field_names}
        for record in records
    ]


def index_directory(source_dir: Path, config, create_index: bool = False) -> int:
    """Upload all legacy files or only pending SharePoint changes to Azure AI Search."""
    if not getattr(config, "azure_search_configured", False):
        raise RuntimeError("Falta configurar AZURE_SEARCH_ENDPOINT y AZURE_SEARCH_API_KEY.")
    if create_index:
        ensure_index(config)
    client = SearchClient(
        endpoint=config.azure_search_endpoint,
        index_name=config.azure_search_index_name,
        credential=_credential(config),
    )
    deleted_document_ids = _deletion_document_ids(source_dir)
    if deleted_document_ids:
        _delete_record_ids(client, _existing_record_ids(client, deleted_document_ids))
        _clear_deletion_manifest(source_dir)

    change_manifest_document_ids = _changed_document_ids(source_dir)
    pending_changes = None if create_index else change_manifest_document_ids
    records = _document_records(source_dir, document_ids=pending_changes)
    quality_metadata_enabled = bool(
        getattr(config, "index_quality_metadata_enabled", False)
    )
    if getattr(config, "retrieval_strategy", "legacy") == "v2" or quality_metadata_enabled:
        # A reviewed document can gain v2 metadata before the chat switches to
        # v2. Filter it against the live schema, keeping the update focused and
        # avoiding a full rebuild just to enrich one approved artifact.
        index_client = SearchIndexClient(
            endpoint=config.azure_search_endpoint,
            credential=_credential(config),
        )
        field_names = {
            field.name
            for field in index_client.get_index(config.azure_search_index_name).fields
            if field.name
        }
        records = _records_supported_by_index(records, field_names)
    elif records:
        # Production keeps the proven legacy schema until v2 promotion is
        # explicit. Do not send additive quality fields during ordinary
        # refreshes to an older index that would reject them.
        legacy_field_names = set(records[0]).difference(V2_ONLY_INDEX_FIELDS)
        records = _records_supported_by_index(records, legacy_field_names)
    target_document_ids = (
        pending_changes
        if pending_changes is not None
        else {str(record["document_id"]) for record in records}
    )
    previous_ids = _existing_record_ids(client, target_document_ids)

    if records:
        # A changed SharePoint document normally keeps its stable chunk id.
        # Update those records with ``merge`` and preserve their stored vector.
        # This is important for the legacy pilot index: it is at its storage
        # quota and its historical vectors can have a different configured
        # dimension than a local ingestion environment. New chunks still get
        # embeddings and use merge-or-upload as usual.
        existing_records = [record for record in records if record["id"] in previous_ids]
        new_records = [record for record in records if record["id"] not in previous_ids]

        for offset in range(0, len(existing_records), 500):
            results = client.merge_documents(documents=existing_records[offset : offset + 500])
            failures = [result.key for result in results if not result.succeeded]
            if failures:
                raise RuntimeError(f"Azure AI Search no actualizó {len(failures)} fragmentos.")

        if new_records:
            _attach_embeddings(new_records, config)
            for offset in range(0, len(new_records), 500):
                results = client.merge_or_upload_documents(documents=new_records[offset : offset + 500])
                failures = [result.key for result in results if not result.succeeded]
                if failures:
                    raise RuntimeError(f"Azure AI Search rechazó {len(failures)} fragmentos.")

    stale_ids = previous_ids.difference(str(record["id"]) for record in records)
    _delete_record_ids(client, stale_ids)
    if change_manifest_document_ids is not None:
        _clear_change_manifest(source_dir)
    return len(records)
